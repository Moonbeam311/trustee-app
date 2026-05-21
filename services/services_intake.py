from datetime import datetime
from database.db import get_connection, get_current_firm_id


INTAKE_LANES = {
    "new_planning": {
        "label": "I want to start organizing my estate, trust, or family plan.",
        "posture": "planning",
        "default_depth": "standard",
        "risk_posture": "unknown",
        "professional_review_recommended": False,
        "automation_limits": "standard",
        "next_screen": "universal_profile",
    },
    "document_review": {
        "label": "I already have documents and want them reviewed or updated.",
        "posture": "review",
        "default_depth": "document_focused",
        "risk_posture": "unknown",
        "professional_review_recommended": False,
        "automation_limits": "standard",
        "next_screen": "document_inventory",
    },
    "administration": {
        "label": "I am responsible for a trust, estate, or fiduciary role.",
        "posture": "fiduciary",
        "default_depth": "administrative",
        "risk_posture": "unknown",
        "professional_review_recommended": False,
        "automation_limits": "standard",
        "next_screen": "fiduciary_role_check",
    },
    "asset_funding": {
        "label": "I want to organize, transfer, or fund assets into a trust or structure.",
        "posture": "execution_preparation",
        "default_depth": "asset_focused",
        "risk_posture": "unknown",
        "professional_review_recommended": False,
        "automation_limits": "standard",
        "next_screen": "asset_snapshot",
    },
    "business_continuity": {
        "label": "I own or manage a business and want continuity or protection planning.",
        "posture": "business_owner",
        "default_depth": "business_focused",
        "risk_posture": "unknown",
        "professional_review_recommended": False,
        "automation_limits": "standard",
        "next_screen": "business_profile",
    },
    "urgent_triage": {
        "label": "Something urgent or complicated is happening.",
        "posture": "crisis_or_pressure",
        "default_depth": "triage",
        "risk_posture": "elevated",
        "professional_review_recommended": True,
        "automation_limits": "high",
        "next_screen": "triage_precheck",
    },
    "education": {
        "label": "I am just learning and want guidance.",
        "posture": "exploratory",
        "default_depth": "light",
        "risk_posture": "low",
        "professional_review_recommended": False,
        "automation_limits": "low",
        "next_screen": "guided_orientation",
    },
}


def ensure_intake_tables():
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT UNIQUE NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            client_id TEXT,
            intake_lane TEXT NOT NULL,
            user_posture TEXT,
            default_depth TEXT,
            risk_posture TEXT,
            professional_review_recommended INTEGER DEFAULT 0,
            automation_limits TEXT,
            next_screen TEXT,
            status TEXT DEFAULT 'lane_selected',
            created_at TEXT,
            updated_at TEXT,
            completed_at TEXT,
            created_by TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_lane_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            event_type TEXT,
            event_label TEXT,
            event_value TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def get_intake_lanes():
    return INTAKE_LANES


def get_lane_config(lane_key):
    return INTAKE_LANES.get(lane_key)


def _next_intake_id(cur):
    cur.execute("SELECT intake_id FROM intake_sessions ORDER BY id DESC LIMIT 1")
    row = cur.fetchone()
    if not row:
        return "INTAKE-0001"

    last = row["intake_id"] if hasattr(row, "keys") else row[0]
    try:
        number = int(str(last).split("-")[-1])
    except Exception:
        number = 0
    return f"INTAKE-{number + 1:04d}"


def create_intake_session(lane_key, client_id=None, created_by=None):
    ensure_intake_tables()

    lane = get_lane_config(lane_key)
    if not lane:
        raise ValueError(f"Invalid intake lane: {lane_key}")

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    conn.row_factory = None
    cur = conn.cursor()

    intake_id = _next_intake_id(cur)

    cur.execute("""
        INSERT INTO intake_sessions (
            intake_id, firm_id, client_id, intake_lane, user_posture,
            default_depth, risk_posture, professional_review_recommended,
            automation_limits, next_screen, status, created_at, updated_at,
            created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        client_id,
        lane_key,
        lane["posture"],
        lane["default_depth"],
        lane["risk_posture"],
        1 if lane["professional_review_recommended"] else 0,
        lane["automation_limits"],
        lane["next_screen"],
        "lane_selected",
        now,
        now,
        created_by,
    ))

    cur.execute("""
        INSERT INTO intake_lane_events (
            intake_id, firm_id, event_type, event_label, event_value,
            created_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        "lane_selected",
        "Intake lane selected",
        lane_key,
        now,
        created_by,
    ))

    conn.commit()
    conn.close()

    return {
        "intake_id": intake_id,
        "intake_lane": lane_key,
        "user_posture": lane["posture"],
        "default_depth": lane["default_depth"],
        "risk_posture": lane["risk_posture"],
        "professional_review_recommended": lane["professional_review_recommended"],
        "automation_limits": lane["automation_limits"],
        "next_screen": lane["next_screen"],
    }


# -------------------------------------------------------------------
# INT-1B — Intake Translation Map
# -------------------------------------------------------------------

UNIVERSAL_INTAKE_QUESTIONS = {
    "planning_context": {
        "label": "Who are we planning for?",
        "input_type": "single",
        "options": {
            "just_me": "Just me",
            "me_and_spouse": "Me and my spouse or partner",
            "family": "My family",
            "parent_or_elder": "A parent or elder",
            "business": "A business",
            "existing_trust_or_estate": "A trust or estate already created",
            "other": "Other / not sure",
        },
    },
    "main_goals": {
        "label": "What are your main goals?",
        "input_type": "multi",
        "options": {
            "protect_family": "Protect family",
            "avoid_probate": "Avoid probate or court confusion",
            "organize_assets": "Organize assets",
            "reduce_conflict": "Reduce confusion or family conflict",
            "choose_decision_makers": "Choose decision-makers",
            "business_continuity": "Protect or continue a business",
            "plan_for_children": "Plan for children",
            "incapacity_planning": "Plan for incapacity",
            "privacy": "Preserve privacy",
            "legacy_charitable": "Support legacy, charitable, or community goals",
            "review_documents": "Review existing documents",
            "other_goal": "Other / not listed",
        },
    },
    "asset_snapshot": {
        "label": "Which of these do you own, manage, or expect to inherit?",
        "input_type": "multi",
        "options": {
            "home": "Home",
            "rental_property": "Rental property",
            "land": "Land",
            "bank_accounts": "Bank accounts",
            "vehicles": "Vehicles",
            "business": "Business",
            "retirement_accounts": "Retirement accounts",
            "life_insurance": "Life insurance",
            "investments": "Stocks, investments, or brokerage accounts",
            "digital_assets": "Digital assets",
            "intellectual_property": "Intellectual property",
            "collectibles_heirlooms": "Collectibles, heirlooms, or heritage property",
            "trust_estate_assets": "Trust or estate assets",
            "other_asset": "Other / not listed",
            "none_not_sure": "None / not sure",
        },
    },
    "people_dependents": {
        "label": "Who depends on you or may need to be included?",
        "input_type": "multi",
        "options": {
            "spouse_partner": "Spouse or partner",
            "minor_children": "Minor children",
            "adult_children": "Adult children",
            "parent_elder": "Parent or elder",
            "special_needs_person": "Disabled or special-needs person",
            "business_partner": "Business partner",
            "charitable_beneficiary": "Charitable, community, or legacy beneficiary",
            "other_person": "Other / not listed",
            "no_dependents": "No dependents",
            "not_sure": "Not sure",
        },
    },
    "concerns": {
        "label": "Do any of these concerns apply?",
        "input_type": "multi",
        "options": {
            "family_conflict": "Family conflict",
            "lawsuit_concern": "Lawsuit concern",
            "debt_creditor": "Debt or creditor concern",
            "tax_concern": "Tax concern",
            "business_liability": "Business liability",
            "divorce_remarriage": "Divorce or remarriage concern",
            "multi_state_property": "Property in more than one state",
            "medical_incapacity": "Medical or incapacity concern",
            "no_trusted_successor": "No trusted successor",
            "missing_documents": "Missing documents",
            "none_apply": "None of these apply",
            "not_sure": "Not sure",
        },
    },
    "existing_documents": {
        "label": "Which documents do you already have?",
        "input_type": "multi",
        "options": {
            "will": "Will",
            "trust": "Trust",
            "power_of_attorney": "Power of attorney",
            "health_directive": "Health directive",
            "deed": "Deed",
            "mortgage_statement": "Mortgage statement",
            "business_documents": "Business documents",
            "insurance_policies": "Insurance policies",
            "beneficiary_forms": "Retirement / beneficiary forms",
            "tax_filings": "Tax filings",
            "court_documents": "Court documents",
            "other_document": "Other / not listed",
            "none": "None",
            "not_sure": "Not sure",
        },
    },
}


TRANSLATION_RULES = {
    # Planning context
    "planning_context.just_me": {
        "system_category": "PERSON_PROFILE",
        "system_meaning": "individual_planning",
        "module_triggers": ["foundational_profile"],
        "document_requests": [],
        "next_sessions": ["initial_structure_review"],
        "risk_flags": [],
    },
    "planning_context.me_and_spouse": {
        "system_category": "FAMILY_STRUCTURE",
        "system_meaning": "spousal_or_partner_planning",
        "module_triggers": ["spousal_profile", "beneficiary_review"],
        "document_requests": [],
        "next_sessions": ["family_structure_review"],
        "risk_flags": [],
    },
    "planning_context.family": {
        "system_category": "FAMILY_STRUCTURE",
        "system_meaning": "family_planning",
        "module_triggers": ["family_profile", "beneficiary_review"],
        "document_requests": [],
        "next_sessions": ["family_structure_review"],
        "risk_flags": [],
    },
    "planning_context.parent_or_elder": {
        "system_category": "DEPENDENCY_PROFILE",
        "system_meaning": "elder_or_parent_planning",
        "module_triggers": ["elder_planning", "authority_document_review"],
        "document_requests": ["power_of_attorney", "health_directive", "existing_care_documents"],
        "next_sessions": ["elder_authority_review"],
        "risk_flags": ["possible_incapacity_or_care_issue"],
    },
    "planning_context.business": {
        "system_category": "BUSINESS_PROFILE",
        "system_meaning": "business_planning_context",
        "module_triggers": ["business_continuity"],
        "document_requests": ["operating_agreement", "ein_letter", "business_license"],
        "next_sessions": ["business_continuity_review"],
        "risk_flags": ["business_continuity_needed"],
    },
    "planning_context.existing_trust_or_estate": {
        "system_category": "FIDUCIARY_CONTEXT",
        "system_meaning": "existing_trust_or_estate_matter",
        "module_triggers": ["trust_audit", "fiduciary_administration"],
        "document_requests": ["trust_document", "letters_testamentary_or_authority", "asset_inventory"],
        "next_sessions": ["fiduciary_authority_review"],
        "risk_flags": ["authority_review_needed"],
    },

    # Goals
    "main_goals.protect_family": {
        "system_category": "OBJECTIVE_PROFILE",
        "system_meaning": "family_protection_objective",
        "module_triggers": ["beneficiary_planning", "successor_planning"],
        "document_requests": [],
        "next_sessions": ["beneficiary_planning_review"],
        "risk_flags": [],
    },
    "main_goals.avoid_probate": {
        "system_category": "OBJECTIVE_PROFILE",
        "system_meaning": "probate_avoidance_objective",
        "module_triggers": ["trust_planning", "funding_checklist"],
        "document_requests": ["existing_will", "deed", "beneficiary_forms"],
        "next_sessions": ["probate_avoidance_review"],
        "risk_flags": [],
    },
    "main_goals.organize_assets": {
        "system_category": "OBJECTIVE_PROFILE",
        "system_meaning": "asset_organization_objective",
        "module_triggers": ["asset_inventory"],
        "document_requests": ["asset_statements", "titles", "deeds"],
        "next_sessions": ["asset_document_deep_dive"],
        "risk_flags": [],
    },
    "main_goals.reduce_conflict": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "conflict_reduction_objective",
        "module_triggers": ["governance_controls", "decision_authority_review"],
        "document_requests": ["existing_will", "existing_trust", "family_agreements"],
        "next_sessions": ["governance_conflict_review"],
        "risk_flags": ["family_conflict_risk"],
    },
    "main_goals.choose_decision_makers": {
        "system_category": "FIDUCIARY_READINESS",
        "system_meaning": "decision_maker_selection_needed",
        "module_triggers": ["trustee_selection", "agent_selection"],
        "document_requests": [],
        "next_sessions": ["fiduciary_selection_review"],
        "risk_flags": [],
    },
    "main_goals.business_continuity": {
        "system_category": "BUSINESS_PROFILE",
        "system_meaning": "business_continuity_objective",
        "module_triggers": ["business_continuity", "succession_planning"],
        "document_requests": ["operating_agreement", "business_license", "bank_authority_records"],
        "next_sessions": ["business_continuity_review"],
        "risk_flags": ["business_continuity_needed"],
    },
    "main_goals.plan_for_children": {
        "system_category": "BENEFICIARY_PROFILE",
        "system_meaning": "children_planning_objective",
        "module_triggers": ["beneficiary_planning", "guardian_review"],
        "document_requests": [],
        "next_sessions": ["children_guardian_review"],
        "risk_flags": [],
    },
    "main_goals.incapacity_planning": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "incapacity_planning_objective",
        "module_triggers": ["poa_review", "health_directive_review"],
        "document_requests": ["power_of_attorney", "health_directive"],
        "next_sessions": ["incapacity_authority_review"],
        "risk_flags": ["incapacity_planning_needed"],
    },
    "main_goals.privacy": {
        "system_category": "OBJECTIVE_PROFILE",
        "system_meaning": "privacy_preference",
        "module_triggers": ["privacy_review", "disclosure_control"],
        "document_requests": [],
        "next_sessions": ["privacy_preferences_review"],
        "risk_flags": [],
    },
    "main_goals.legacy_charitable": {
        "system_category": "LEGACY_PROFILE",
        "system_meaning": "legacy_or_charitable_objective",
        "module_triggers": ["legacy_planning", "charitable_intent_review"],
        "document_requests": [],
        "next_sessions": ["legacy_objectives_review"],
        "risk_flags": [],
    },
    "main_goals.other_goal": {
        "system_category": "OBJECTIVE_PROFILE",
        "system_meaning": "other_goal_not_listed",
        "module_triggers": ["planning_objective_review"],
        "document_requests": ["goal_description"],
        "next_sessions": ["initial_structure_review"],
        "risk_flags": [],
    },

    "main_goals.review_documents": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "document_review_objective",
        "module_triggers": ["document_audit"],
        "document_requests": ["existing_documents"],
        "next_sessions": ["document_audit_session"],
        "risk_flags": [],
    },

    # Assets
    "asset_snapshot.home": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "primary_residence",
        "module_triggers": ["real_property_review", "funding_checklist"],
        "document_requests": ["deed", "mortgage_statement", "property_tax_bill", "homeowners_insurance"],
        "next_sessions": ["real_property_deep_dive"],
        "risk_flags": [],
    },
    "asset_snapshot.rental_property": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "rental_real_property",
        "module_triggers": ["real_property_review", "liability_review", "income_property_review"],
        "document_requests": ["deed", "lease", "insurance", "mortgage_statement"],
        "next_sessions": ["real_property_deep_dive"],
        "risk_flags": ["liability_exposure_possible"],
    },
    "asset_snapshot.land": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "land_or_real_estate",
        "module_triggers": ["real_property_review"],
        "document_requests": ["deed", "tax_bill", "survey_if_available"],
        "next_sessions": ["real_property_deep_dive"],
        "risk_flags": [],
    },
    "asset_snapshot.bank_accounts": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "financial_accounts",
        "module_triggers": ["account_inventory", "beneficiary_review"],
        "document_requests": ["bank_statement", "account_registration_info"],
        "next_sessions": ["financial_account_review"],
        "risk_flags": [],
    },
    "asset_snapshot.vehicles": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "vehicle_assets",
        "module_triggers": ["title_review", "insurance_review"],
        "document_requests": ["vehicle_title", "registration", "insurance"],
        "next_sessions": ["vehicle_title_review"],
        "risk_flags": ["liability_exposure_possible"],
    },
    "asset_snapshot.business": {
        "system_category": "BUSINESS_PROFILE",
        "system_meaning": "business_asset_or_operating_entity",
        "module_triggers": ["business_continuity", "entity_review"],
        "document_requests": ["operating_agreement", "ein_letter", "business_license", "bank_authority_records"],
        "next_sessions": ["business_continuity_review"],
        "risk_flags": ["business_liability_possible"],
    },
    "asset_snapshot.retirement_accounts": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "retirement_assets",
        "module_triggers": ["beneficiary_designation_review"],
        "document_requests": ["retirement_statement", "beneficiary_designation_form"],
        "next_sessions": ["beneficiary_designation_review"],
        "risk_flags": ["transfer_restriction_review_needed"],
    },
    "asset_snapshot.life_insurance": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "life_insurance_policy",
        "module_triggers": ["beneficiary_designation_review", "insurance_review"],
        "document_requests": ["policy_declaration_page", "beneficiary_designation_form"],
        "next_sessions": ["insurance_beneficiary_review"],
        "risk_flags": [],
    },
    "asset_snapshot.investments": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "investment_assets",
        "module_triggers": ["investment_account_review", "beneficiary_review"],
        "document_requests": ["brokerage_statement", "account_registration_info"],
        "next_sessions": ["investment_account_review"],
        "risk_flags": [],
    },
    "asset_snapshot.digital_assets": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "digital_assets",
        "module_triggers": ["digital_asset_inventory"],
        "document_requests": ["digital_asset_list"],
        "next_sessions": ["digital_asset_review"],
        "risk_flags": [],
    },
    "asset_snapshot.intellectual_property": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "intellectual_property",
        "module_triggers": ["ip_inventory"],
        "document_requests": ["ip_registration_records", "licensing_agreements"],
        "next_sessions": ["intellectual_property_review"],
        "risk_flags": [],
    },
    "asset_snapshot.collectibles_heirlooms": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "collectibles_heirlooms_heritage_property",
        "module_triggers": ["heritage_asset_ledger", "special_property_inventory"],
        "document_requests": ["photos", "appraisals", "provenance_records"],
        "next_sessions": ["heritage_asset_review"],
        "risk_flags": ["special_custody_or_heritage_flag"],
    },
    "asset_snapshot.other_asset": {
        "system_category": "ASSET_PROFILE",
        "system_meaning": "other_asset_not_listed",
        "module_triggers": ["general_asset_review"],
        "document_requests": ["asset_description"],
        "next_sessions": ["asset_document_deep_dive"],
        "risk_flags": [],
    },

    "asset_snapshot.trust_estate_assets": {
        "system_category": "FIDUCIARY_CONTEXT",
        "system_meaning": "existing_trust_or_estate_assets",
        "module_triggers": ["fiduciary_inventory", "ledger_readiness"],
        "document_requests": ["trust_document", "estate_authority_document", "asset_inventory"],
        "next_sessions": ["fiduciary_inventory_review"],
        "risk_flags": ["authority_review_needed"],
    },

    # People
    "people_dependents.spouse_partner": {
        "system_category": "FAMILY_STRUCTURE",
        "system_meaning": "spouse_or_partner_involved",
        "module_triggers": ["spousal_planning"],
        "document_requests": [],
        "next_sessions": ["family_structure_review"],
        "risk_flags": [],
    },
    "people_dependents.minor_children": {
        "system_category": "BENEFICIARY_PROFILE",
        "system_meaning": "minor_children_involved",
        "module_triggers": ["guardian_review", "minor_beneficiary_controls"],
        "document_requests": [],
        "next_sessions": ["children_guardian_review"],
        "risk_flags": ["minor_children_flag"],
    },
    "people_dependents.adult_children": {
        "system_category": "BENEFICIARY_PROFILE",
        "system_meaning": "adult_children_involved",
        "module_triggers": ["beneficiary_planning"],
        "document_requests": [],
        "next_sessions": ["beneficiary_planning_review"],
        "risk_flags": [],
    },
    "people_dependents.parent_elder": {
        "system_category": "DEPENDENCY_PROFILE",
        "system_meaning": "elder_dependency",
        "module_triggers": ["elder_planning"],
        "document_requests": ["poa_if_available", "health_directive_if_available"],
        "next_sessions": ["elder_authority_review"],
        "risk_flags": ["elder_dependency_flag"],
    },
    "people_dependents.special_needs_person": {
        "system_category": "BENEFICIARY_PROFILE",
        "system_meaning": "special_needs_person_involved",
        "module_triggers": ["special_needs_review"],
        "document_requests": ["benefits_information_if_available"],
        "next_sessions": ["special_needs_planning_review"],
        "risk_flags": ["special_needs_flag"],
    },
    "people_dependents.business_partner": {
        "system_category": "BUSINESS_PROFILE",
        "system_meaning": "business_partner_involved",
        "module_triggers": ["business_governance_review"],
        "document_requests": ["operating_agreement", "partnership_agreement"],
        "next_sessions": ["business_governance_review"],
        "risk_flags": ["co_owner_or_partner_flag"],
    },
    "people_dependents.other_person": {
        "system_category": "BENEFICIARY_PROFILE",
        "system_meaning": "other_person_or_dependent_not_listed",
        "module_triggers": ["beneficiary_planning"],
        "document_requests": ["person_or_dependent_description"],
        "next_sessions": ["beneficiary_planning_review"],
        "risk_flags": [],
    },

    "people_dependents.charitable_beneficiary": {
        "system_category": "LEGACY_PROFILE",
        "system_meaning": "charitable_or_legacy_beneficiary",
        "module_triggers": ["legacy_planning", "charitable_intent_review"],
        "document_requests": [],
        "next_sessions": ["legacy_objectives_review"],
        "risk_flags": [],
    },

    # Concerns
    "concerns.family_conflict": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "family_conflict_concern",
        "module_triggers": ["governance_controls", "conflict_review"],
        "document_requests": ["existing_will", "existing_trust", "family_agreements"],
        "next_sessions": ["governance_conflict_review"],
        "risk_flags": ["family_conflict_risk"],
    },
    "concerns.lawsuit_concern": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "lawsuit_concern",
        "module_triggers": ["risk_review"],
        "document_requests": ["court_documents", "claim_letters_if_available"],
        "next_sessions": ["risk_triage_review"],
        "risk_flags": ["urgent_or_legal_review_flag"],
    },
    "concerns.debt_creditor": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "debt_or_creditor_concern",
        "module_triggers": ["liability_review"],
        "document_requests": ["debt_statements", "creditor_letters_if_available"],
        "next_sessions": ["liability_review"],
        "risk_flags": ["creditor_pressure_flag"],
    },
    "concerns.tax_concern": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "tax_concern",
        "module_triggers": ["tax_professional_review"],
        "document_requests": ["tax_notices", "tax_filings"],
        "next_sessions": ["tax_review_referral"],
        "risk_flags": ["tax_review_flag"],
    },
    "concerns.business_liability": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "business_liability_concern",
        "module_triggers": ["business_liability_review"],
        "document_requests": ["operating_agreement", "insurance_policies"],
        "next_sessions": ["business_liability_review"],
        "risk_flags": ["business_liability_possible"],
    },
    "concerns.divorce_remarriage": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "divorce_or_remarriage_concern",
        "module_triggers": ["family_structure_review", "beneficiary_review"],
        "document_requests": ["divorce_decree_if_available", "prenuptial_or_postnuptial_if_available"],
        "next_sessions": ["family_structure_review"],
        "risk_flags": ["family_status_complexity"],
    },
    "concerns.multi_state_property": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "multi_state_property_concern",
        "module_triggers": ["multi_jurisdiction_review", "real_property_review"],
        "document_requests": ["deeds", "property_tax_bills"],
        "next_sessions": ["multi_jurisdiction_property_review"],
        "risk_flags": ["multi_jurisdiction_flag"],
    },
    "concerns.medical_incapacity": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "medical_or_incapacity_concern",
        "module_triggers": ["poa_review", "health_directive_review"],
        "document_requests": ["power_of_attorney", "health_directive", "medical_authority_documents"],
        "next_sessions": ["incapacity_authority_review"],
        "risk_flags": ["incapacity_planning_needed"],
    },
    "concerns.no_trusted_successor": {
        "system_category": "FIDUCIARY_READINESS",
        "system_meaning": "no_trusted_successor_identified",
        "module_triggers": ["successor_selection"],
        "document_requests": [],
        "next_sessions": ["fiduciary_selection_review"],
        "risk_flags": ["successor_gap_flag"],
    },
    "concerns.none_apply": {
        "system_category": "RISK_PROFILE",
        "system_meaning": "no_initial_concerns_reported",
        "module_triggers": [],
        "document_requests": [],
        "next_sessions": [],
        "risk_flags": [],
    },

    "concerns.missing_documents": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "missing_documents_concern",
        "module_triggers": ["document_collection"],
        "document_requests": ["document_checklist"],
        "next_sessions": ["document_collection_review"],
        "risk_flags": ["documentation_gap"],
    },

    # Existing docs
    "existing_documents.will": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "will_exists",
        "module_triggers": ["document_audit"],
        "document_requests": ["will"],
        "next_sessions": ["document_audit_session"],
        "risk_flags": [],
    },
    "existing_documents.trust": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "trust_exists",
        "module_triggers": ["trust_audit"],
        "document_requests": ["trust_document"],
        "next_sessions": ["trust_document_review"],
        "risk_flags": [],
    },
    "existing_documents.power_of_attorney": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "poa_exists",
        "module_triggers": ["authority_document_review"],
        "document_requests": ["power_of_attorney"],
        "next_sessions": ["authority_document_review"],
        "risk_flags": [],
    },
    "existing_documents.health_directive": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "health_directive_exists",
        "module_triggers": ["health_directive_review"],
        "document_requests": ["health_directive"],
        "next_sessions": ["authority_document_review"],
        "risk_flags": [],
    },
    "existing_documents.deed": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "deed_available",
        "module_triggers": ["real_property_review"],
        "document_requests": ["deed"],
        "next_sessions": ["real_property_deep_dive"],
        "risk_flags": [],
    },
    "existing_documents.mortgage_statement": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "mortgage_statement_available",
        "module_triggers": ["real_property_review"],
        "document_requests": ["mortgage_statement"],
        "next_sessions": ["real_property_deep_dive"],
        "risk_flags": [],
    },
    "existing_documents.business_documents": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "business_documents_exist",
        "module_triggers": ["business_document_audit"],
        "document_requests": ["business_documents"],
        "next_sessions": ["business_continuity_review"],
        "risk_flags": [],
    },
    "existing_documents.insurance_policies": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "insurance_documents_exist",
        "module_triggers": ["insurance_review"],
        "document_requests": ["insurance_policies"],
        "next_sessions": ["insurance_beneficiary_review"],
        "risk_flags": [],
    },
    "existing_documents.beneficiary_forms": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "beneficiary_forms_exist",
        "module_triggers": ["beneficiary_designation_review"],
        "document_requests": ["beneficiary_forms"],
        "next_sessions": ["beneficiary_designation_review"],
        "risk_flags": [],
    },
    "existing_documents.tax_filings": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "tax_filings_available",
        "module_triggers": ["tax_professional_review"],
        "document_requests": ["tax_filings"],
        "next_sessions": ["tax_review_referral"],
        "risk_flags": ["tax_review_flag"],
    },
    "existing_documents.court_documents": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "court_documents_available",
        "module_triggers": ["risk_review"],
        "document_requests": ["court_documents"],
        "next_sessions": ["risk_triage_review"],
        "risk_flags": ["urgent_or_legal_review_flag"],
    },
    "existing_documents.other_document": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "other_document_available",
        "module_triggers": ["document_audit"],
        "document_requests": ["other_document_description"],
        "next_sessions": ["document_audit_session"],
        "risk_flags": [],
    },

    "existing_documents.none": {
        "system_category": "DOCUMENT_STATUS",
        "system_meaning": "no_documents_available",
        "module_triggers": ["foundational_estate_package", "document_collection"],
        "document_requests": ["document_checklist"],
        "next_sessions": ["foundational_planning_review"],
        "risk_flags": ["documentation_gap"],
    },
}


def ensure_intake_translation_tables():
    ensure_intake_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            question_key TEXT NOT NULL,
            answer_key TEXT NOT NULL,
            answer_label TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_translations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            source_key TEXT NOT NULL,
            system_category TEXT,
            system_meaning TEXT,
            module_trigger TEXT,
            document_request TEXT,
            next_session TEXT,
            risk_flag TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def get_universal_intake_questions():
    return UNIVERSAL_INTAKE_QUESTIONS


def get_translation_rules():
    return TRANSLATION_RULES


def _normalize_answer_values(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v) for v in value if str(v).strip()]
    if isinstance(value, tuple):
        return [str(v) for v in value if str(v).strip()]
    if str(value).strip():
        return [str(value).strip()]
    return []


def translate_answer(question_key, answer_key):
    source_key = f"{question_key}.{answer_key}"
    return TRANSLATION_RULES.get(source_key)


def save_universal_profile_answers(intake_id, form_data, created_by=None):
    ensure_intake_translation_tables()

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    saved_answers = []
    translations = []

    for question_key, question in UNIVERSAL_INTAKE_QUESTIONS.items():
        if question["input_type"] == "multi":
            values = form_data.getlist(question_key) if hasattr(form_data, "getlist") else _normalize_answer_values(form_data.get(question_key))
        else:
            values = _normalize_answer_values(form_data.get(question_key))

        for answer_key in values:
            answer_label = question["options"].get(answer_key, answer_key)
            source_key = f"{question_key}.{answer_key}"
            rule = TRANSLATION_RULES.get(source_key)

            cur.execute("""
                INSERT INTO intake_answers (
                    intake_id, firm_id, question_key, answer_key, answer_label,
                    created_at, created_by
                )
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                intake_id,
                firm_id,
                question_key,
                answer_key,
                answer_label,
                now,
                created_by,
            ))

            saved_answers.append({
                "question_key": question_key,
                "answer_key": answer_key,
                "answer_label": answer_label,
            })

            if rule:
                module_triggers = rule.get("module_triggers") or [None]
                document_requests = rule.get("document_requests") or [None]
                next_sessions = rule.get("next_sessions") or [None]
                risk_flags = rule.get("risk_flags") or [None]

                max_len = max(len(module_triggers), len(document_requests), len(next_sessions), len(risk_flags))

                for idx in range(max_len):
                    module_trigger = module_triggers[idx] if idx < len(module_triggers) else None
                    document_request = document_requests[idx] if idx < len(document_requests) else None
                    next_session = next_sessions[idx] if idx < len(next_sessions) else None
                    risk_flag = risk_flags[idx] if idx < len(risk_flags) else None

                    cur.execute("""
                        INSERT INTO intake_translations (
                            intake_id, firm_id, source_key, system_category,
                            system_meaning, module_trigger, document_request,
                            next_session, risk_flag, created_at, created_by
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        intake_id,
                        firm_id,
                        source_key,
                        rule.get("system_category"),
                        rule.get("system_meaning"),
                        module_trigger,
                        document_request,
                        next_session,
                        risk_flag,
                        now,
                        created_by,
                    ))

                    translations.append({
                        "source_key": source_key,
                        "system_category": rule.get("system_category"),
                        "system_meaning": rule.get("system_meaning"),
                        "module_trigger": module_trigger,
                        "document_request": document_request,
                        "next_session": next_session,
                        "risk_flag": risk_flag,
                    })

    cur.execute("""
        UPDATE intake_sessions
        SET status = ?, updated_at = ?
        WHERE intake_id = ?
    """, ("universal_profile_completed", now, intake_id))

    conn.commit()
    conn.close()

    scores = score_and_save_intake(intake_id, translations, created_by=created_by)

    return {
        "intake_id": intake_id,
        "answers": saved_answers,
        "translations": translations,
        "summary": summarize_intake_translations(translations),
        "scores": scores,
    }


def summarize_intake_translations(translations):
    def ordered_unique(values):
        seen = set()
        output = []
        for value in values:
            if value and value not in seen:
                seen.add(value)
                output.append(value)
        return output

    return {
        "system_categories": ordered_unique(t.get("system_category") for t in translations),
        "system_meanings": ordered_unique(t.get("system_meaning") for t in translations),
        "module_triggers": ordered_unique(t.get("module_trigger") for t in translations),
        "document_requests": ordered_unique(t.get("document_request") for t in translations),
        "next_sessions": ordered_unique(t.get("next_session") for t in translations),
        "risk_flags": ordered_unique(t.get("risk_flag") for t in translations),
    }


def get_intake_session(intake_id):
    ensure_intake_tables()
    conn = get_connection()
    conn.row_factory = None
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, intake_lane, user_posture, default_depth, risk_posture,
               professional_review_recommended, automation_limits, next_screen,
               status, created_at, updated_at
        FROM intake_sessions
        WHERE intake_id = ?
    """, (intake_id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        return None

    keys = [
        "intake_id", "intake_lane", "user_posture", "default_depth",
        "risk_posture", "professional_review_recommended", "automation_limits",
        "next_screen", "status", "created_at", "updated_at"
    ]
    return dict(zip(keys, row))


# -------------------------------------------------------------------
# INT-1C — Intake Scoring Engine
# -------------------------------------------------------------------

COMPLEXITY_WEIGHTS = {
    # Asset / structure complexity
    "real_property_review": 2,
    "funding_checklist": 2,
    "income_property_review": 3,
    "business_continuity": 3,
    "entity_review": 3,
    "business_governance_review": 3,
    "multi_jurisdiction_review": 4,
    "fiduciary_administration": 3,
    "trust_audit": 2,
    "fiduciary_inventory": 3,
    "ledger_readiness": 2,
    "special_needs_review": 4,
    "heritage_asset_ledger": 2,
    "ip_inventory": 2,
    "digital_asset_inventory": 1,

    # Family/governance complexity
    "guardian_review": 2,
    "minor_beneficiary_controls": 2,
    "governance_controls": 3,
    "conflict_review": 3,
    "successor_selection": 2,
    "elder_planning": 3,
}

URGENCY_WEIGHTS = {
    # Strong urgency flags
    "urgent_or_legal_review_flag": 5,
    "creditor_pressure_flag": 4,
    "tax_review_flag": 4,
    "incapacity_planning_needed": 4,
    "authority_review_needed": 3,
    "family_conflict_risk": 3,

    # Moderate urgency flags
    "business_continuity_needed": 3,
    "business_liability_possible": 3,
    "successor_gap_flag": 3,
    "documentation_gap": 3,
    "minor_children_flag": 2,
    "special_needs_flag": 4,
    "elder_dependency_flag": 3,
    "multi_jurisdiction_flag": 3,
    "liability_exposure_possible": 2,
    "transfer_restriction_review_needed": 2,
    "co_owner_or_partner_flag": 2,
    "family_status_complexity": 2,
    "special_custody_or_heritage_flag": 1,
}

READINESS_WEIGHTS = {
    # Existing documents increase readiness
    "will_exists": 1,
    "trust_exists": 2,
    "poa_exists": 1,
    "health_directive_exists": 1,
    "deed_available": 2,
    "mortgage_statement_available": 1,
    "business_documents_exist": 2,
    "insurance_documents_exist": 1,
    "beneficiary_forms_exist": 1,
    "tax_filings_available": 1,
    "court_documents_available": 1,

    # Available document requests can also indicate partial readiness
    "will": 1,
    "trust_document": 2,
    "power_of_attorney": 1,
    "health_directive": 1,
    "deed": 2,
    "mortgage_statement": 1,
    "business_documents": 2,
    "insurance_policies": 1,
    "beneficiary_forms": 1,
    "tax_filings": 1,
    "court_documents": 1,
}


def ensure_intake_scoring_tables():
    ensure_intake_translation_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_scores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            complexity_score INTEGER DEFAULT 0,
            complexity_level TEXT,
            urgency_score INTEGER DEFAULT 0,
            urgency_level TEXT,
            readiness_score INTEGER DEFAULT 0,
            readiness_level TEXT,
            scoring_notes TEXT,
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def classify_complexity(score):
    if score <= 4:
        return "Simple"
    if score <= 9:
        return "Moderate"
    if score <= 15:
        return "Advanced"
    return "Complex"


def classify_urgency(score):
    if score <= 3:
        return "Low"
    if score <= 7:
        return "Medium"
    return "High"


def classify_readiness(score):
    if score <= 2:
        return "Not Ready"
    if score <= 6:
        return "Partially Ready"
    return "Ready for Deep Review"


def calculate_intake_scores(translations):
    complexity_score = 0
    urgency_score = 0
    readiness_score = 0

    complexity_hits = []
    urgency_hits = []
    readiness_hits = []

    for item in translations:
        module_trigger = item.get("module_trigger")
        risk_flag = item.get("risk_flag")
        system_meaning = item.get("system_meaning")
        document_request = item.get("document_request")

        if module_trigger in COMPLEXITY_WEIGHTS:
            value = COMPLEXITY_WEIGHTS[module_trigger]
            complexity_score += value
            complexity_hits.append(f"{module_trigger}+{value}")

        if risk_flag in URGENCY_WEIGHTS:
            value = URGENCY_WEIGHTS[risk_flag]
            urgency_score += value
            urgency_hits.append(f"{risk_flag}+{value}")

        if system_meaning in READINESS_WEIGHTS:
            value = READINESS_WEIGHTS[system_meaning]
            readiness_score += value
            readiness_hits.append(f"{system_meaning}+{value}")

        if document_request in READINESS_WEIGHTS:
            value = READINESS_WEIGHTS[document_request]
            readiness_score += value
            readiness_hits.append(f"{document_request}+{value}")

    return {
        "complexity_score": complexity_score,
        "complexity_level": classify_complexity(complexity_score),
        "urgency_score": urgency_score,
        "urgency_level": classify_urgency(urgency_score),
        "readiness_score": readiness_score,
        "readiness_level": classify_readiness(readiness_score),
        "scoring_notes": {
            "complexity_hits": complexity_hits,
            "urgency_hits": urgency_hits,
            "readiness_hits": readiness_hits,
        }
    }


def save_intake_scores(intake_id, scores, created_by=None):
    ensure_intake_scoring_tables()

    import json
    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_scores (
            intake_id, firm_id, complexity_score, complexity_level,
            urgency_score, urgency_level, readiness_score, readiness_level,
            scoring_notes, created_at, updated_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        scores["complexity_score"],
        scores["complexity_level"],
        scores["urgency_score"],
        scores["urgency_level"],
        scores["readiness_score"],
        scores["readiness_level"],
        json.dumps(scores.get("scoring_notes", {})),
        now,
        now,
        created_by,
    ))

    cur.execute("""
        UPDATE intake_sessions
        SET status = ?, updated_at = ?
        WHERE intake_id = ?
    """, ("scored", now, intake_id))

    conn.commit()
    conn.close()

    return scores


def score_and_save_intake(intake_id, translations, created_by=None):
    scores = calculate_intake_scores(translations)
    save_intake_scores(intake_id, scores, created_by=created_by)
    return scores


# -------------------------------------------------------------------
# INT-1D — Client-Facing Initial Fiduciary Snapshot
# -------------------------------------------------------------------

CLIENT_PRIORITY_LABELS = {
    "family_structure_review": "Review family structure and decision-maker roles",
    "initial_structure_review": "Review the stated planning goal and place it into the right follow-up path",
    "fiduciary_selection_review": "Choose or confirm trusted decision-makers",
    "real_property_deep_dive": "Review real property ownership and supporting documents",
    "financial_account_review": "Review financial accounts and how they are titled",
    "business_continuity_review": "Review business continuity and operating authority",
    "beneficiary_planning_review": "Review beneficiary planning and distribution intentions",
    "business_governance_review": "Review business ownership, partners, and governance",
    "governance_conflict_review": "Address family conflict or governance concerns early",
    "risk_triage_review": "Complete a risk-focused review before taking major action",
    "tax_review_referral": "Gather tax records for qualified tax review",
    "foundational_planning_review": "Begin foundational estate and trust planning",
    "children_guardian_review": "Review minor child planning and guardian choices",
    "elder_authority_review": "Review elder, parent, or incapacity authority documents",
    "special_needs_planning_review": "Review special-needs planning before final documents",
    "insurance_beneficiary_review": "Review insurance ownership and beneficiary designations",
    "beneficiary_designation_review": "Review beneficiary designations on non-trust assets",
    "investment_account_review": "Review investment accounts and registration details",
    "digital_asset_review": "Prepare a digital asset inventory",
    "intellectual_property_review": "Review intellectual property ownership and records",
    "heritage_asset_review": "Document heritage, heirloom, or special family property",
    "fiduciary_inventory_review": "Build or verify trust/estate inventory records",
    "document_collection_review": "Gather missing documents before deeper review",
    "authority_document_review": "Review authority documents such as POA or health directive",
    "privacy_preferences_review": "Clarify privacy and disclosure preferences",
    "legacy_objectives_review": "Clarify legacy, charitable, or community goals",
    "probate_avoidance_review": "Review probate-avoidance objectives and document gaps",
    "asset_document_deep_dive": "Organize assets and supporting records",
    "general_asset_review": "Review other assets not listed in the standard intake",
    "planning_objective_review": "Review other planning goals not listed in the standard intake",
    "multi_jurisdiction_property_review": "Review property located in more than one state",
    "business_liability_review": "Review business liability and insurance concerns",
    "liability_review": "Review debt, creditor, or liability concerns",
    "incapacity_authority_review": "Review incapacity planning and authority documents",
    "document_audit_session": "Audit existing estate, trust, or authority documents",
    "trust_document_review": "Review existing trust documents",
}

CLIENT_DOCUMENT_LABELS = {
    "deed": "Deed or title record",
    "deeds": "Deeds or title records",
    "mortgage_statement": "Mortgage statement",
    "property_tax_bill": "Property tax bill",
    "tax_bill": "Property tax bill",
    "survey_if_available": "Survey, if available",
    "homeowners_insurance": "Homeowners insurance declarations",
    "insurance": "Insurance policy or declarations",
    "insurance_policies": "Insurance policies",
    "policy_declaration_page": "Insurance policy declaration page",
    "beneficiary_designation_form": "Beneficiary designation form",
    "beneficiary_forms": "Beneficiary designation forms",
    "bank_statement": "Bank or financial account statement",
    "asset_statements": "Asset statements",
    "account_registration_info": "Account registration information",
    "retirement_statement": "Retirement account statement",
    "brokerage_statement": "Brokerage or investment statement",
    "operating_agreement": "Operating agreement",
    "partnership_agreement": "Partnership agreement",
    "ein_letter": "EIN letter",
    "business_license": "Business license",
    "business_documents": "Business documents",
    "bank_authority_records": "Business bank authority records",
    "existing_will": "Existing will",
    "will": "Will",
    "existing_trust": "Existing trust",
    "trust_document": "Trust document",
    "power_of_attorney": "Power of attorney",
    "poa_if_available": "Power of attorney, if available",
    "health_directive": "Health directive",
    "health_directive_if_available": "Health directive, if available",
    "medical_authority_documents": "Medical authority documents",
    "family_agreements": "Family agreements, if any",
    "court_documents": "Court documents",
    "claim_letters_if_available": "Claim letters, if available",
    "tax_notices": "Tax notices",
    "tax_filings": "Tax filings",
    "debt_statements": "Debt statements",
    "creditor_letters_if_available": "Creditor letters, if available",
    "divorce_decree_if_available": "Divorce decree, if available",
    "prenuptial_or_postnuptial_if_available": "Prenuptial or postnuptial agreement, if available",
    "document_checklist": "Basic document checklist",
    "goal_description": "Description of other planning goal not listed",
    "asset_description": "Description of other asset not listed",
    "person_or_dependent_description": "Description of other person, dependent, or role not listed",
    "other_document_description": "Description or copy of other document not listed",
    "digital_asset_list": "Digital asset list",
    "ip_registration_records": "IP registration records",
    "licensing_agreements": "Licensing agreements",
    "photos": "Photos or visual inventory",
    "appraisals": "Appraisals",
    "provenance_records": "Provenance or history records",
    "trust_document": "Trust document",
    "estate_authority_document": "Estate authority document",
    "letters_testamentary_or_authority": "Letters testamentary or authority document",
    "asset_inventory": "Asset inventory",
    "existing_documents": "Existing documents",
    "existing_care_documents": "Existing care documents",
    "benefits_information_if_available": "Benefits information, if available",
}

CLIENT_RISK_LABELS = {
    "urgent_or_legal_review_flag": "Legal, court, or urgent review may be needed",
    "creditor_pressure_flag": "Debt or creditor pressure may require careful review",
    "tax_review_flag": "Tax records should be reviewed by a qualified professional",
    "incapacity_planning_needed": "Incapacity planning may need attention",
    "authority_review_needed": "Authority documents should be verified",
    "family_conflict_risk": "Family conflict risk should be addressed early",
    "business_continuity_needed": "Business continuity planning should be prioritized",
    "business_liability_possible": "Business liability should be reviewed",
    "successor_gap_flag": "A successor decision-maker gap may exist",
    "documentation_gap": "Important documents may be missing",
    "minor_children_flag": "Minor child planning should be reviewed",
    "special_needs_flag": "Special-needs planning should be reviewed carefully",
    "elder_dependency_flag": "Elder or parent support planning may be involved",
    "multi_jurisdiction_flag": "Property in multiple jurisdictions may need extra review",
    "liability_exposure_possible": "Possible liability exposure should be reviewed",
    "transfer_restriction_review_needed": "Some assets may have transfer restrictions",
    "co_owner_or_partner_flag": "Co-owner or partner interests should be reviewed",
    "family_status_complexity": "Family status changes may affect planning",
    "special_custody_or_heritage_flag": "Special custody or heritage property should be documented",
}

PLANNING_TYPE_LABELS = {
    "FAMILY_STRUCTURE": "Family Planning",
    "ASSET_PROFILE": "Asset Organization",
    "BUSINESS_PROFILE": "Business Continuity",
    "BENEFICIARY_PROFILE": "Beneficiary Planning",
    "FIDUCIARY_READINESS": "Decision-Maker Planning",
    "FIDUCIARY_CONTEXT": "Trust or Estate Administration",
    "RISK_PROFILE": "Risk Review",
    "DOCUMENT_STATUS": "Document Readiness",
    "DEPENDENCY_PROFILE": "Elder / Dependency Planning",
    "LEGACY_PROFILE": "Legacy or Charitable Planning",
    "PERSON_PROFILE": "Individual Planning",
    "OBJECTIVE_PROFILE": "Planning Objectives",
}


def _client_unique(values, limit=None):
    seen = set()
    output = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            output.append(value)
    if limit:
        return output[:limit]
    return output


def _label_items(values, labels, limit=None):
    labeled = []
    for value in values:
        if not value:
            continue
        labeled.append(labels.get(value, value.replace("_", " ").title()))
    return _client_unique(labeled, limit=limit)


def determine_primary_next_session(summary):
    sessions = summary.get("next_sessions", []) or []
    priority_order = [
        "risk_triage_review",
        "tax_review_referral",
        "governance_conflict_review",
        "business_continuity_review",
        "real_property_deep_dive",
        "asset_document_deep_dive",
        "fiduciary_selection_review",
        "children_guardian_review",
        "beneficiary_planning_review",
        "document_collection_review",
        "foundational_planning_review",
    ]

    for item in priority_order:
        if item in sessions:
            return CLIENT_PRIORITY_LABELS.get(item, item.replace("_", " ").title())

    if sessions:
        first = sessions[0]
        return CLIENT_PRIORITY_LABELS.get(first, first.replace("_", " ").title())

    return "Initial structure review"


def build_client_snapshot(result):
    summary = result.get("summary", {}) or {}
    scores = result.get("scores", {}) or {}

    categories = summary.get("system_categories", []) or []
    next_sessions = summary.get("next_sessions", []) or []
    document_requests = summary.get("document_requests", []) or []
    risk_flags = summary.get("risk_flags", []) or []

    planning_types = _label_items(categories, PLANNING_TYPE_LABELS, limit=5)
    top_priorities = _label_items(next_sessions, CLIENT_PRIORITY_LABELS, limit=6)
    documents_to_gather = _label_items(document_requests, CLIENT_DOCUMENT_LABELS, limit=12)
    review_flags = _label_items(risk_flags, CLIENT_RISK_LABELS, limit=6)

    if not top_priorities:
        top_priorities = ["Complete a deeper review of family, assets, documents, and decision-makers"]

    if not documents_to_gather:
        documents_to_gather = ["Any existing estate, trust, property, business, insurance, or account documents"]

    next_session = determine_primary_next_session(summary)

    complexity_level = scores.get("complexity_level", "Not Yet Rated")
    urgency_level = scores.get("urgency_level", "Not Yet Rated")
    readiness_level = scores.get("readiness_level", "Not Yet Rated")

    if urgency_level == "High":
        review_priority = "High"
    elif complexity_level in ["Advanced", "Complex"]:
        review_priority = "Elevated"
    elif urgency_level == "Medium":
        review_priority = "Moderate"
    else:
        review_priority = "Standard"

    return {
        "intake_id": result.get("intake_id"),
        "planning_types": planning_types,
        "complexity_level": complexity_level,
        "urgency_level": urgency_level,
        "readiness_level": readiness_level,
        "review_priority": review_priority,
        "top_priorities": top_priorities,
        "documents_to_gather": documents_to_gather,
        "review_flags": review_flags,
        "recommended_next_session": next_session,
        "technical_result": result,
    }


# -------------------------------------------------------------------
# INT-1E — Save Snapshot + Intake Dashboard List
# -------------------------------------------------------------------

def ensure_intake_snapshot_tables():
    ensure_intake_scoring_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_snapshots (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT UNIQUE NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            snapshot_json TEXT,
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def save_client_snapshot(intake_id, snapshot, created_by=None):
    ensure_intake_snapshot_tables()

    import json
    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    # Keep stored JSON compact and avoid recursive technical payload bloat.
    stored_snapshot = dict(snapshot)
    stored_snapshot.pop("technical_result", None)

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_snapshots (
            intake_id, firm_id, snapshot_json, created_at, updated_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?)
        ON CONFLICT(intake_id) DO UPDATE SET
            snapshot_json = excluded.snapshot_json,
            updated_at = excluded.updated_at,
            created_by = excluded.created_by
    """, (
        intake_id,
        firm_id,
        json.dumps(stored_snapshot),
        now,
        now,
        created_by,
    ))

    cur.execute("""
        UPDATE intake_sessions
        SET status = ?, updated_at = ?
        WHERE intake_id = ?
    """, ("snapshot_saved", now, intake_id))

    conn.commit()
    conn.close()

    return stored_snapshot


def get_latest_intake_score_map():
    ensure_intake_scoring_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT s1.intake_id,
               s1.complexity_score, s1.complexity_level,
               s1.urgency_score, s1.urgency_level,
               s1.readiness_score, s1.readiness_level
        FROM intake_scores s1
        INNER JOIN (
            SELECT intake_id, MAX(id) AS max_id
            FROM intake_scores
            GROUP BY intake_id
        ) latest
        ON s1.id = latest.max_id
    """)

    rows = cur.fetchall()
    conn.close()

    score_map = {}
    for row in rows:
        score_map[row[0]] = {
            "complexity_score": row[1],
            "complexity_level": row[2],
            "urgency_score": row[3],
            "urgency_level": row[4],
            "readiness_score": row[5],
            "readiness_level": row[6],
        }
    return score_map


def list_intake_dashboard(limit=50):
    ensure_intake_snapshot_tables()
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, intake_lane, user_posture, default_depth, risk_posture,
               professional_review_recommended, automation_limits, next_screen,
               status, created_at, updated_at
        FROM intake_sessions
        WHERE firm_id = ?
        ORDER BY id DESC
        LIMIT ?
    """, (firm_id, limit))

    rows = cur.fetchall()
    conn.close()

    scores = get_latest_intake_score_map()

    items = []
    for row in rows:
        intake_id = row[0]
        score = scores.get(intake_id, {})

        items.append({
            "intake_id": intake_id,
            "intake_lane": row[1],
            "user_posture": row[2],
            "default_depth": row[3],
            "risk_posture": row[4],
            "professional_review_recommended": bool(row[5]),
            "automation_limits": row[6],
            "next_screen": row[7],
            "status": row[8],
            "created_at": row[9],
            "updated_at": row[10],
            "complexity_level": score.get("complexity_level", "—"),
            "urgency_level": score.get("urgency_level", "—"),
            "readiness_level": score.get("readiness_level", "—"),
        })

    return items


def get_intake_answers(intake_id):
    ensure_intake_translation_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT question_key, answer_key, answer_label
        FROM intake_answers
        WHERE intake_id = ?
        ORDER BY id ASC
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "question_key": row[0],
            "answer_key": row[1],
            "answer_label": row[2],
        }
        for row in rows
    ]


def get_intake_translations_for_snapshot(intake_id):
    ensure_intake_translation_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT source_key, system_category, system_meaning,
               module_trigger, document_request, next_session, risk_flag
        FROM intake_translations
        WHERE intake_id = ?
        ORDER BY id ASC
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "source_key": row[0],
            "system_category": row[1],
            "system_meaning": row[2],
            "module_trigger": row[3],
            "document_request": row[4],
            "next_session": row[5],
            "risk_flag": row[6],
        }
        for row in rows
    ]


def get_latest_intake_scores(intake_id):
    ensure_intake_scoring_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT complexity_score, complexity_level,
               urgency_score, urgency_level,
               readiness_score, readiness_level,
               scoring_notes
        FROM intake_scores
        WHERE intake_id = ?
        ORDER BY id DESC
        LIMIT 1
    """, (intake_id,))

    row = cur.fetchone()
    conn.close()

    if not row:
        return {}

    return {
        "complexity_score": row[0],
        "complexity_level": row[1],
        "urgency_score": row[2],
        "urgency_level": row[3],
        "readiness_score": row[4],
        "readiness_level": row[5],
        "scoring_notes": row[6],
    }


def rebuild_intake_result(intake_id):
    translations = get_intake_translations_for_snapshot(intake_id)
    scores = get_latest_intake_scores(intake_id)

    return {
        "intake_id": intake_id,
        "answers": get_intake_answers(intake_id),
        "translations": translations,
        "summary": summarize_intake_translations(translations),
        "scores": scores,
    }


def get_saved_client_snapshot(intake_id):
    ensure_intake_snapshot_tables()

    import json

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT snapshot_json
        FROM intake_snapshots
        WHERE intake_id = ?
        LIMIT 1
    """, (intake_id,))

    row = cur.fetchone()
    conn.close()

    if row and row[0]:
        saved = json.loads(row[0])
        technical_result = rebuild_intake_result(intake_id)
        saved["technical_result"] = technical_result
        return saved, technical_result

    # Fallback: rebuild from stored translations/scores if no snapshot row exists.
    technical_result = rebuild_intake_result(intake_id)
    if not technical_result.get("translations"):
        return None, None

    snapshot = build_client_snapshot(technical_result)
    snapshot["technical_result"] = technical_result
    return snapshot, technical_result


# -------------------------------------------------------------------
# INT-1F — Intake Resume + Completion Status Controls
# -------------------------------------------------------------------

COMPLETED_INTAKE_STATUSES = {
    "scored",
    "snapshot_saved",
    "completed",
}

INCOMPLETE_INTAKE_STATUSES = {
    "lane_selected",
    "started",
    "universal_profile_started",
}


def intake_is_completed(status):
    return status in COMPLETED_INTAKE_STATUSES


def intake_is_incomplete(status):
    return not intake_is_completed(status)


def get_intake_resume_target(intake_id):
    """
    Decide where an intake should resume.
    V1 resumes incomplete lane-selected sessions at universal profile.
    Later versions can resume at lane-specific screens.
    """
    intake = get_intake_session(intake_id)
    if not intake:
        return None

    status = intake.get("status")

    if intake_is_completed(status):
        return {
            "route": "intake_saved_snapshot",
            "intake_id": intake_id,
            "label": "View snapshot",
        }

    return {
        "route": "intake_universal_profile",
        "intake_id": intake_id,
        "label": "Continue intake",
    }


def list_intake_dashboard_with_controls(limit=100):
    items = list_intake_dashboard(limit=limit)

    for item in items:
        status = item.get("status")
        completed = intake_is_completed(status)

        item["is_completed"] = completed
        item["is_incomplete"] = not completed
        item["status_label"] = "Completed" if completed else "Incomplete"

        if completed:
            item["primary_action"] = "View snapshot"
            item["primary_route"] = "snapshot"
        else:
            item["primary_action"] = "Continue intake"
            item["primary_route"] = "continue"

        if item.get("complexity_level") == "—" and not completed:
            item["complexity_level"] = "Pending"

        if item.get("urgency_level") == "—" and not completed:
            item["urgency_level"] = "Pending"

        if item.get("readiness_level") == "—" and not completed:
            item["readiness_level"] = "Pending"

    return items


# -------------------------------------------------------------------
# INT-1G — Dashboard Polish + Snapshot Print/Export Prep
# -------------------------------------------------------------------

def format_intake_timestamp(value):
    if not value:
        return "—"

    try:
        raw = str(value).replace("T", " ")
        if "." in raw:
            raw = raw.split(".")[0]
        return raw
    except Exception:
        return str(value)


def decorate_intake_dashboard_items(items, status_filter="all"):
    decorated = []

    for item in items:
        item = dict(item)

        item["updated_display"] = format_intake_timestamp(
            item.get("updated_at") or item.get("created_at")
        )

        urgency = item.get("urgency_level") or "Pending"
        readiness = item.get("readiness_level") or "Pending"
        complexity = item.get("complexity_level") or "Pending"

        item["urgency_badge"] = urgency
        item["readiness_badge"] = readiness
        item["complexity_badge"] = complexity

        item["urgency_class"] = "badge-high" if urgency == "High" else "badge-medium" if urgency == "Medium" else "badge-low" if urgency == "Low" else "badge-pending"
        item["readiness_class"] = "badge-ready" if readiness == "Ready for Deep Review" else "badge-medium" if readiness == "Partially Ready" else "badge-pending"
        item["complexity_class"] = "badge-high" if complexity in ["Advanced", "Complex"] else "badge-medium" if complexity == "Moderate" else "badge-low" if complexity == "Simple" else "badge-pending"

        if status_filter == "completed" and not item.get("is_completed"):
            continue
        if status_filter == "incomplete" and not item.get("is_incomplete"):
            continue
        if status_filter == "high_urgency" and item.get("urgency_level") != "High":
            continue

        decorated.append(item)

    return decorated


def list_intake_dashboard_polished(limit=100, status_filter="all"):
    items = list_intake_dashboard_with_controls(limit=limit)
    return decorate_intake_dashboard_items(items, status_filter=status_filter)


def prepare_snapshot_export_metadata(intake_id):
    snapshot, result = get_saved_client_snapshot(intake_id)
    if not snapshot:
        return None

    return {
        "intake_id": intake_id,
        "snapshot": snapshot,
        "result": result,
        "export_ready": True,
        "available_exports": [
            "print_view",
            "future_pdf_export",
            "future_docx_export",
        ],
        "note": "Export preparation is ready. PDF/DOCX generation is intentionally not activated in INT-1G.",
    }


# -------------------------------------------------------------------
# INT-1H — Intake Review Notes + Internal Staff Comments
# -------------------------------------------------------------------

VALID_REVIEW_NOTE_TYPES = {
    "general": "General Note",
    "risk": "Risk / Review Concern",
    "document": "Document Follow-Up",
    "client_followup": "Client Follow-Up",
    "professional_review": "Professional Review",
    "admin": "Administrative Note",
}

VALID_REVIEW_PRIORITIES = {
    "low": "Low",
    "normal": "Normal",
    "high": "High",
    "urgent": "Urgent",
}

VALID_FOLLOWUP_STATUSES = {
    "open": "Open",
    "pending_client": "Pending Client",
    "pending_review": "Pending Review",
    "resolved": "Resolved",
}


def ensure_intake_review_note_tables():
    ensure_intake_snapshot_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_review_notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            note_type TEXT DEFAULT 'general',
            priority TEXT DEFAULT 'normal',
            followup_status TEXT DEFAULT 'open',
            note_body TEXT NOT NULL,
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def create_intake_review_note(
    intake_id,
    note_body,
    note_type="general",
    priority="normal",
    followup_status="open",
    created_by=None
):
    ensure_intake_review_note_tables()

    note_body = (note_body or "").strip()
    if not note_body:
        raise ValueError("Review note cannot be blank.")

    if note_type not in VALID_REVIEW_NOTE_TYPES:
        note_type = "general"

    if priority not in VALID_REVIEW_PRIORITIES:
        priority = "normal"

    if followup_status not in VALID_FOLLOWUP_STATUSES:
        followup_status = "open"

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_review_notes (
            intake_id, firm_id, note_type, priority, followup_status,
            note_body, created_at, updated_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        note_type,
        priority,
        followup_status,
        note_body,
        now,
        now,
        created_by,
    ))

    cur.execute("""
        UPDATE intake_sessions
        SET updated_at = ?
        WHERE intake_id = ?
    """, (now, intake_id))

    conn.commit()
    conn.close()

    return {
        "intake_id": intake_id,
        "note_type": note_type,
        "priority": priority,
        "followup_status": followup_status,
        "note_body": note_body,
        "created_at": now,
        "created_by": created_by,
    }


def list_intake_review_notes(intake_id):
    ensure_intake_review_note_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT id, intake_id, note_type, priority, followup_status,
               note_body, created_at, updated_at, created_by
        FROM intake_review_notes
        WHERE intake_id = ?
        ORDER BY id DESC
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "id": row[0],
            "intake_id": row[1],
            "note_type": row[2],
            "note_type_label": VALID_REVIEW_NOTE_TYPES.get(row[2], row[2]),
            "priority": row[3],
            "priority_label": VALID_REVIEW_PRIORITIES.get(row[3], row[3]),
            "followup_status": row[4],
            "followup_status_label": VALID_FOLLOWUP_STATUSES.get(row[4], row[4]),
            "note_body": row[5],
            "created_at": format_intake_timestamp(row[6]),
            "updated_at": format_intake_timestamp(row[7]),
            "created_by": row[8] or "—",
        }
        for row in rows
    ]


def get_intake_review_note_counts():
    ensure_intake_review_note_tables()
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, COUNT(*) AS note_count
        FROM intake_review_notes
        WHERE firm_id = ?
        GROUP BY intake_id
    """, (firm_id,))

    rows = cur.fetchall()
    conn.close()

    return {row[0]: row[1] for row in rows}


def list_intake_dashboard_with_review_notes(limit=100, status_filter="all"):
    items = list_intake_dashboard_polished(limit=limit, status_filter=status_filter)
    counts = get_intake_review_note_counts()

    for item in items:
        count = counts.get(item["intake_id"], 0)
        item["review_note_count"] = count
        item["has_review_notes"] = count > 0

    return items


def get_review_note_form_options():
    return {
        "note_types": VALID_REVIEW_NOTE_TYPES,
        "priorities": VALID_REVIEW_PRIORITIES,
        "followup_statuses": VALID_FOLLOWUP_STATUSES,
    }


# -------------------------------------------------------------------
# INT-1I — Intake Follow-Up Task Builder
# -------------------------------------------------------------------

VALID_FOLLOWUP_TASK_TYPES = {
    "document": "Document Request",
    "client_followup": "Client Follow-Up",
    "professional_review": "Professional Review",
    "staff_action": "Staff Action",
    "next_session": "Next Session Prep",
}

VALID_FOLLOWUP_TASK_PRIORITIES = {
    "low": "Low",
    "normal": "Normal",
    "high": "High",
    "urgent": "Urgent",
}

VALID_FOLLOWUP_TASK_STATUSES = {
    "open": "Open",
    "pending_client": "Pending Client",
    "pending_staff": "Pending Staff",
    "pending_professional": "Pending Professional Review",
    "completed": "Completed",
    "deferred": "Deferred",
}


def ensure_intake_followup_task_tables():
    ensure_intake_review_note_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_followup_tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            task_type TEXT DEFAULT 'staff_action',
            priority TEXT DEFAULT 'normal',
            status TEXT DEFAULT 'open',
            title TEXT NOT NULL,
            description TEXT,
            source TEXT DEFAULT 'manual',
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT,
            completed_at TEXT,
            completed_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def create_intake_followup_task(
    intake_id,
    title,
    description="",
    task_type="staff_action",
    priority="normal",
    status="open",
    source="manual",
    created_by=None
):
    ensure_intake_followup_task_tables()

    title = (title or "").strip()
    description = (description or "").strip()

    if not title:
        raise ValueError("Task title cannot be blank.")

    if task_type not in VALID_FOLLOWUP_TASK_TYPES:
        task_type = "staff_action"

    if priority not in VALID_FOLLOWUP_TASK_PRIORITIES:
        priority = "normal"

    if status not in VALID_FOLLOWUP_TASK_STATUSES:
        status = "open"

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_followup_tasks (
            intake_id, firm_id, task_type, priority, status, title,
            description, source, created_at, updated_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        task_type,
        priority,
        status,
        title,
        description,
        source,
        now,
        now,
        created_by,
    ))

    cur.execute("""
        UPDATE intake_sessions
        SET updated_at = ?
        WHERE intake_id = ?
    """, (now, intake_id))

    conn.commit()
    conn.close()

    return {
        "intake_id": intake_id,
        "task_type": task_type,
        "priority": priority,
        "status": status,
        "title": title,
        "description": description,
        "source": source,
        "created_at": now,
        "created_by": created_by,
    }


def list_intake_followup_tasks(intake_id):
    ensure_intake_followup_task_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT id, intake_id, task_type, priority, status, title,
               description, source, created_at, updated_at, created_by,
               completed_at, completed_by
        FROM intake_followup_tasks
        WHERE intake_id = ?
        ORDER BY
            CASE status
                WHEN 'open' THEN 1
                WHEN 'pending_client' THEN 2
                WHEN 'pending_staff' THEN 3
                WHEN 'pending_professional' THEN 4
                WHEN 'deferred' THEN 5
                WHEN 'completed' THEN 6
                ELSE 7
            END,
            CASE priority
                WHEN 'urgent' THEN 1
                WHEN 'high' THEN 2
                WHEN 'normal' THEN 3
                WHEN 'low' THEN 4
                ELSE 5
            END,
            id ASC
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "id": row[0],
            "intake_id": row[1],
            "task_type": row[2],
            "task_type_label": VALID_FOLLOWUP_TASK_TYPES.get(row[2], row[2]),
            "priority": row[3],
            "priority_label": VALID_FOLLOWUP_TASK_PRIORITIES.get(row[3], row[3]),
            "status": row[4],
            "status_label": VALID_FOLLOWUP_TASK_STATUSES.get(row[4], row[4]),
            "title": row[5],
            "description": row[6],
            "source": row[7],
            "created_at": format_intake_timestamp(row[8]),
            "updated_at": format_intake_timestamp(row[9]),
            "created_by": row[10] or "—",
            "completed_at": format_intake_timestamp(row[11]) if row[11] else "",
            "completed_by": row[12] or "",
        }
        for row in rows
    ]


def get_intake_followup_task_counts():
    ensure_intake_followup_task_tables()
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id,
               COUNT(*) AS total_count,
               SUM(CASE WHEN status = 'completed' THEN 1 ELSE 0 END) AS completed_count,
               SUM(CASE WHEN status != 'completed' THEN 1 ELSE 0 END) AS open_count
        FROM intake_followup_tasks
        WHERE firm_id = ?
        GROUP BY intake_id
    """, (firm_id,))

    rows = cur.fetchall()
    conn.close()

    output = {}
    for row in rows:
        output[row[0]] = {
            "task_count": row[1] or 0,
            "completed_task_count": row[2] or 0,
            "open_task_count": row[3] or 0,
        }
    return output


def task_exists(intake_id, title, task_type=None, source=None):
    ensure_intake_followup_task_tables()

    conn = get_connection()
    cur = conn.cursor()

    query = """
        SELECT COUNT(*)
        FROM intake_followup_tasks
        WHERE intake_id = ? AND title = ?
    """
    params = [intake_id, title]

    if task_type:
        query += " AND task_type = ?"
        params.append(task_type)

    if source:
        query += " AND source = ?"
        params.append(source)

    cur.execute(query, tuple(params))
    row = cur.fetchone()
    conn.close()

    return bool(row and row[0])


def auto_generate_followup_tasks_from_snapshot(intake_id, snapshot, created_by=None):
    """
    Idempotently generates follow-up tasks from the client snapshot.
    It will not duplicate the same generated task title/source pair.
    """
    ensure_intake_followup_task_tables()

    created = []

    # Document requests become client-facing document tasks.
    for doc in snapshot.get("documents_to_gather", []) or []:
        title = f"Gather document: {doc}"
        if not task_exists(intake_id, title, task_type="document", source="auto_snapshot"):
            created.append(create_intake_followup_task(
                intake_id=intake_id,
                title=title,
                description="Client or staff should gather this item before the deeper review session.",
                task_type="document",
                priority="normal",
                status="pending_client",
                source="auto_snapshot",
                created_by=created_by,
            ))

    # Review flags become review tasks.
    for flag in snapshot.get("review_flags", []) or []:
        title = f"Review flag: {flag}"
        priority = "high" if "tax" in flag.lower() or "legal" in flag.lower() or "liability" in flag.lower() else "normal"
        status = "pending_professional" if "tax" in flag.lower() or "legal" in flag.lower() else "pending_staff"

        if not task_exists(intake_id, title, source="auto_snapshot"):
            created.append(create_intake_followup_task(
                intake_id=intake_id,
                title=title,
                description="This item was flagged by the intake translation/scoring engine and should be reviewed before final action.",
                task_type="professional_review" if status == "pending_professional" else "staff_action",
                priority=priority,
                status=status,
                source="auto_snapshot",
                created_by=created_by,
            ))

    # Recommended next session becomes a next-session prep task.
    next_session = snapshot.get("recommended_next_session")
    if next_session:
        title = f"Prepare next session: {next_session}"
        if not task_exists(intake_id, title, task_type="next_session", source="auto_snapshot"):
            created.append(create_intake_followup_task(
                intake_id=intake_id,
                title=title,
                description="Prepare agenda, documents, and follow-up questions for the recommended next review session.",
                task_type="next_session",
                priority="high" if snapshot.get("review_priority") in ["High", "Elevated"] else "normal",
                status="pending_staff",
                source="auto_snapshot",
                created_by=created_by,
            ))

    return created


def update_intake_followup_task_status(task_id, status, updated_by=None):
    ensure_intake_followup_task_tables()

    if status not in VALID_FOLLOWUP_TASK_STATUSES:
        raise ValueError("Invalid task status.")

    now = datetime.utcnow().isoformat(timespec="seconds")
    completed_at = now if status == "completed" else None
    completed_by = updated_by if status == "completed" else None

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        UPDATE intake_followup_tasks
        SET status = ?,
            updated_at = ?,
            completed_at = ?,
            completed_by = ?
        WHERE id = ?
    """, (
        status,
        now,
        completed_at,
        completed_by,
        task_id,
    ))

    conn.commit()
    conn.close()


def get_followup_task_form_options():
    return {
        "task_types": VALID_FOLLOWUP_TASK_TYPES,
        "priorities": VALID_FOLLOWUP_TASK_PRIORITIES,
        "statuses": VALID_FOLLOWUP_TASK_STATUSES,
    }


def list_intake_dashboard_with_tasks(limit=100, status_filter="all"):
    items = list_intake_dashboard_with_review_notes(limit=limit, status_filter=status_filter)
    counts = get_intake_followup_task_counts()

    for item in items:
        data = counts.get(item["intake_id"], {})
        item["task_count"] = data.get("task_count", 0)
        item["open_task_count"] = data.get("open_task_count", 0)
        item["completed_task_count"] = data.get("completed_task_count", 0)
        item["has_tasks"] = item["task_count"] > 0

    return items


# -------------------------------------------------------------------
# INT-1J — Intake Task Filters + Follow-Up Workflow Polish
# -------------------------------------------------------------------

def summarize_followup_tasks(tasks):
    summary = {
        "total": len(tasks or []),
        "open": 0,
        "pending_client": 0,
        "pending_staff": 0,
        "pending_professional": 0,
        "completed": 0,
        "deferred": 0,
    }

    for task in tasks or []:
        status = task.get("status")
        if status in summary:
            summary[status] += 1

        if status != "completed":
            summary["open"] += 1

    return summary


def group_followup_tasks(tasks):
    groups = {
        "pending_client": [],
        "pending_staff": [],
        "pending_professional": [],
        "open": [],
        "deferred": [],
        "completed": [],
    }

    for task in tasks or []:
        status = task.get("status") or "open"
        if status in groups:
            groups[status].append(task)
        else:
            groups["open"].append(task)

    return groups


# -------------------------------------------------------------------
# INT-1K — Intake Follow-Up Packet / Checklist Export Prep
# -------------------------------------------------------------------

def build_intake_followup_packet(intake_id):
    snapshot, result = get_saved_client_snapshot(intake_id)
    if not snapshot:
        return None

    tasks = list_intake_followup_tasks(intake_id)
    notes = list_intake_review_notes(intake_id)
    task_summary = summarize_followup_tasks(tasks)
    task_groups = group_followup_tasks(tasks)

    documents = snapshot.get("documents_to_gather", []) or []
    priorities = snapshot.get("top_priorities", []) or []
    review_flags = snapshot.get("review_flags", []) or []

    packet = {
        "intake_id": intake_id,
        "snapshot": snapshot,
        "result": result,
        "documents": documents,
        "priorities": priorities,
        "review_flags": review_flags,
        "recommended_next_session": snapshot.get("recommended_next_session"),
        "tasks": tasks,
        "task_summary": task_summary,
        "task_groups": task_groups,
        "notes": notes,
        "packet_status": "Prepared",
        "packet_type": "Initial Follow-Up Packet",
        "export_ready": True,
    }

    return packet


def get_packet_readiness_label(packet):
    if not packet:
        return "Not Available"

    task_summary = packet.get("task_summary", {}) or {}
    open_count = task_summary.get("open", 0)
    documents = packet.get("documents", []) or []

    if open_count == 0 and documents:
        return "Ready for Review"
    if open_count > 0:
        return "Action Items Open"
    return "Prepared"


# -------------------------------------------------------------------
# INT-1L — Follow-Up Packet PDF/DOCX Export
# -------------------------------------------------------------------

def ensure_intake_export_dir():
    from pathlib import Path

    export_dir = Path("exports/intake_packets")
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def safe_export_filename(value):
    value = str(value or "intake").strip()
    keep = []
    for ch in value:
        if ch.isalnum() or ch in ["-", "_"]:
            keep.append(ch)
        else:
            keep.append("_")
    return "".join(keep)


def generate_followup_packet_docx(intake_id):
    from pathlib import Path
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    packet = build_intake_followup_packet(intake_id)
    if not packet:
        return None

    export_dir = ensure_intake_export_dir()
    filename = f"{safe_export_filename(intake_id)}_Follow_Up_Packet.docx"
    out_path = export_dir / filename

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Initial Intake Follow-Up Packet")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Intake ID: {packet['intake_id']} | Packet Status: {packet['packet_status']}")

    doc.add_paragraph("")

    def add_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Pt(18)
        p.add_run(f"- {text}")

    add_heading("Snapshot Summary")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Planning Type", " + ".join(packet["snapshot"].get("planning_types", [])) or "Initial Planning"),
        ("Complexity", packet["snapshot"].get("complexity_level", "")),
        ("Priority", packet["snapshot"].get("review_priority", "")),
        ("Readiness", packet["snapshot"].get("readiness_level", "")),
        ("Recommended Next Session", packet.get("recommended_next_session") or "Initial structure review"),
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    doc.add_paragraph("")

    add_heading("Top Priorities")
    for idx, item in enumerate(packet.get("priorities", []) or [], start=1):
        add_bullet(f"{idx}. {item}")
    if not packet.get("priorities"):
        add_bullet("No priorities generated.")

    add_heading("Document Checklist")
    for item in packet.get("documents", []) or []:
        add_bullet(f"[ ] {item}")
    if not packet.get("documents"):
        add_bullet("No document checklist generated.")

    if packet.get("review_flags"):
        add_heading("Review Flags")
        for item in packet.get("review_flags", []):
            add_bullet(item)

    add_heading("Follow-Up Task Summary")
    task_summary = packet.get("task_summary", {}) or {}
    summary_rows = [
        ("Total", task_summary.get("total", 0)),
        ("Open", task_summary.get("open", 0)),
        ("Pending Client", task_summary.get("pending_client", 0)),
        ("Pending Staff", task_summary.get("pending_staff", 0)),
        ("Pending Professional", task_summary.get("pending_professional", 0)),
        ("Completed", task_summary.get("completed", 0)),
    ]

    task_table = doc.add_table(rows=0, cols=2)
    task_table.style = "Table Grid"

    for label, value in summary_rows:
        cells = task_table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    doc.add_paragraph("")

    add_heading("Follow-Up Task Checklist")
    for task in packet.get("tasks", []) or []:
        check = "[x]" if task.get("status") == "completed" else "[ ]"
        add_bullet(f"{check} {task.get('title')}")
        meta = f"{task.get('task_type_label')} | Priority: {task.get('priority_label')} | Status: {task.get('status_label')} | Source: {task.get('source')}"
        add_bullet(f"    {meta}")
        if task.get("description"):
            add_bullet(f"    {task.get('description')}")

    if not packet.get("tasks"):
        add_bullet("No follow-up tasks generated.")

    add_heading("Internal Review Notes")
    for note in packet.get("notes", []) or []:
        meta = f"{note.get('note_type_label')} | Priority: {note.get('priority_label')} | Status: {note.get('followup_status_label')} | By: {note.get('created_by')} | {note.get('created_at')}"
        add_bullet(meta)
        add_bullet(f"    {note.get('note_body')}")
    if not packet.get("notes"):
        add_bullet("No internal review notes added.")

    add_heading("Important Notice")
    doc.add_paragraph(
        "This follow-up packet is an intake preparation tool. It does not finalize legal, tax, fiduciary, "
        "or asset-transfer decisions. Additional review may be needed before documents are signed, assets are "
        "transferred, or formal action is taken."
    )

    doc.save(out_path)
    return str(out_path)


def generate_followup_packet_pdf(intake_id):
    """
    Generate a PDF version of the follow-up packet.

    Windows-safe fallback order:
    1. Generate DOCX.
    2. Try LibreOffice/soffice if installed.
    3. Try Microsoft Word COM automation if pywin32 + Word are available.
    4. Return None if unavailable so the route can fail gracefully.
    """
    import os
    import shutil
    import subprocess
    from pathlib import Path

    docx_path = generate_followup_packet_docx(intake_id)
    if not docx_path:
        return None

    docx_path = Path(docx_path).resolve()
    export_dir = docx_path.parent
    pdf_path = export_dir / f"{docx_path.stem}.pdf"

    try:
        if pdf_path.exists():
            pdf_path.unlink()
    except Exception:
        pass

    soffice_candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for soffice in soffice_candidates:
        if not soffice:
            continue

        try:
            soffice_path = Path(soffice)
            if not soffice_path.exists() and not shutil.which(str(soffice)):
                continue

            cmd = [
                str(soffice),
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(export_dir),
                str(docx_path),
            ]

            subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )

            if pdf_path.exists():
                return str(pdf_path)

        except Exception:
            continue

    if os.name == "nt":
        try:
            import pythoncom
            import win32com.client

            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            try:
                doc = word.Documents.Open(str(docx_path))
                doc.ExportAsFixedFormat(str(pdf_path), 17)
                doc.Close(False)
            finally:
                word.Quit()
                pythoncom.CoUninitialize()

            if pdf_path.exists():
                return str(pdf_path)

        except Exception:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            return None

    return None


# -------------------------------------------------------------------
# INT-1M — Export Folder Gitignore + Export Audit Record
# -------------------------------------------------------------------

def ensure_intake_export_log_tables():
    ensure_intake_snapshot_tables()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_export_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            export_type TEXT,
            export_status TEXT,
            file_path TEXT,
            message TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def log_intake_export(intake_id, export_type, export_status, file_path=None, message=None, created_by=None):
    ensure_intake_export_log_tables()

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_export_logs (
            intake_id, firm_id, export_type, export_status,
            file_path, message, created_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        export_type,
        export_status,
        file_path,
        message,
        now,
        created_by,
    ))

    conn.commit()
    conn.close()

    return {
        "intake_id": intake_id,
        "export_type": export_type,
        "export_status": export_status,
        "file_path": file_path,
        "message": message,
        "created_at": now,
        "created_by": created_by,
    }


def list_intake_export_logs(intake_id):
    ensure_intake_export_log_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT export_type, export_status, file_path, message, created_at, created_by
        FROM intake_export_logs
        WHERE intake_id = ?
        ORDER BY id DESC
        LIMIT 25
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "export_type": row[0],
            "export_status": row[1],
            "file_path": row[2],
            "message": row[3],
            "created_at": format_intake_timestamp(row[4]),
            "created_by": row[5] or "—",
        }
        for row in rows
    ]


def generate_followup_packet_docx_logged(intake_id, created_by=None):
    try:
        path = generate_followup_packet_docx(intake_id)
        if path:
            log_intake_export(
                intake_id=intake_id,
                export_type="docx",
                export_status="success",
                file_path=path,
                message="DOCX follow-up packet generated successfully.",
                created_by=created_by,
            )
            return path

        log_intake_export(
            intake_id=intake_id,
            export_type="docx",
            export_status="failed",
            file_path=None,
            message="DOCX follow-up packet could not be generated.",
            created_by=created_by,
        )
        return None

    except Exception as exc:
        log_intake_export(
            intake_id=intake_id,
            export_type="docx",
            export_status="error",
            file_path=None,
            message=str(exc),
            created_by=created_by,
        )
        return None


def generate_followup_packet_pdf_logged(intake_id, created_by=None):
    try:
        path = generate_followup_packet_pdf(intake_id)
        if path:
            log_intake_export(
                intake_id=intake_id,
                export_type="pdf",
                export_status="success",
                file_path=path,
                message="PDF follow-up packet generated successfully.",
                created_by=created_by,
            )
            return path

        log_intake_export(
            intake_id=intake_id,
            export_type="pdf",
            export_status="failed",
            file_path=None,
            message="Automatic PDF generation unavailable. Use Print packet → Save as PDF or install LibreOffice / pywin32 + Word.",
            created_by=created_by,
        )
        return None

    except Exception as exc:
        log_intake_export(
            intake_id=intake_id,
            export_type="pdf",
            export_status="error",
            file_path=None,
            message=str(exc),
            created_by=created_by,
        )
        return None


# -------------------------------------------------------------------
# INT-1N — Intake Export History Dashboard + Packet Versioning
# -------------------------------------------------------------------

def ensure_intake_export_version_columns():
    ensure_intake_export_log_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("PRAGMA table_info(intake_export_logs)")
    columns = {row[1] for row in cur.fetchall()}

    if "version_number" not in columns:
        cur.execute("ALTER TABLE intake_export_logs ADD COLUMN version_number INTEGER DEFAULT 1")

    if "packet_type" not in columns:
        cur.execute("ALTER TABLE intake_export_logs ADD COLUMN packet_type TEXT DEFAULT 'follow_up_packet'")

    conn.commit()
    conn.close()


def get_next_export_version(intake_id, export_type, packet_type="follow_up_packet"):
    ensure_intake_export_version_columns()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT COALESCE(MAX(version_number), 0)
        FROM intake_export_logs
        WHERE intake_id = ?
          AND export_type = ?
          AND packet_type = ?
          AND export_status IN ('success', 'failed', 'error')
    """, (intake_id, export_type, packet_type))

    row = cur.fetchone()
    conn.close()

    current = row[0] if row and row[0] is not None else 0
    return int(current) + 1


def log_intake_export_versioned(
    intake_id,
    export_type,
    export_status,
    file_path=None,
    message=None,
    created_by=None,
    packet_type="follow_up_packet",
    version_number=None
):
    ensure_intake_export_version_columns()

    if version_number is None:
        version_number = get_next_export_version(intake_id, export_type, packet_type)

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_export_logs (
            intake_id, firm_id, export_type, export_status,
            file_path, message, created_at, created_by,
            version_number, packet_type
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        firm_id,
        export_type,
        export_status,
        file_path,
        message,
        now,
        created_by,
        version_number,
        packet_type,
    ))

    conn.commit()
    conn.close()

    return {
        "intake_id": intake_id,
        "export_type": export_type,
        "export_status": export_status,
        "file_path": file_path,
        "message": message,
        "created_at": now,
        "created_by": created_by,
        "version_number": version_number,
        "packet_type": packet_type,
    }


def list_all_intake_export_logs(limit=100):
    ensure_intake_export_version_columns()
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, export_type, export_status, file_path, message,
               created_at, created_by, version_number, packet_type
        FROM intake_export_logs
        WHERE firm_id = ?
        ORDER BY id DESC
        LIMIT ?
    """, (firm_id, limit))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": row[0],
            "export_type": row[1],
            "export_status": row[2],
            "file_path": row[3],
            "message": row[4],
            "created_at": format_intake_timestamp(row[5]),
            "created_by": row[6] or "—",
            "version_number": row[7] or 1,
            "packet_type": row[8] or "follow_up_packet",
        }
        for row in rows
    ]


def list_intake_export_logs_versioned(intake_id):
    ensure_intake_export_version_columns()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT export_type, export_status, file_path, message,
               created_at, created_by, version_number, packet_type
        FROM intake_export_logs
        WHERE intake_id = ?
        ORDER BY id DESC
        LIMIT 100
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": intake_id,
            "export_type": row[0],
            "export_status": row[1],
            "file_path": row[2],
            "message": row[3],
            "created_at": format_intake_timestamp(row[4]),
            "created_by": row[5] or "—",
            "version_number": row[6] or 1,
            "packet_type": row[7] or "follow_up_packet",
        }
        for row in rows
    ]


def get_intake_export_summary(limit=100):
    logs = list_all_intake_export_logs(limit=limit)

    summary = {
        "total": len(logs),
        "success": 0,
        "failed": 0,
        "error": 0,
        "docx": 0,
        "pdf": 0,
        "other": 0,
    }

    for log in logs:
        status = log.get("export_status")
        export_type = log.get("export_type")

        if status in summary:
            summary[status] += 1

        if export_type == "docx":
            summary["docx"] += 1
        elif export_type == "pdf":
            summary["pdf"] += 1
        else:
            summary["other"] += 1

    return summary


def generate_followup_packet_docx_logged_versioned(intake_id, created_by=None):
    version = get_next_export_version(intake_id, "docx", "follow_up_packet")

    try:
        path = generate_followup_packet_docx(intake_id)
        if path:
            log_intake_export_versioned(
                intake_id=intake_id,
                export_type="docx",
                export_status="success",
                file_path=path,
                message=f"DOCX follow-up packet generated successfully. Version {version}.",
                created_by=created_by,
                packet_type="follow_up_packet",
                version_number=version,
            )
            return path

        log_intake_export_versioned(
            intake_id=intake_id,
            export_type="docx",
            export_status="failed",
            file_path=None,
            message=f"DOCX follow-up packet could not be generated. Version {version}.",
            created_by=created_by,
            packet_type="follow_up_packet",
            version_number=version,
        )
        return None

    except Exception as exc:
        log_intake_export_versioned(
            intake_id=intake_id,
            export_type="docx",
            export_status="error",
            file_path=None,
            message=f"{exc} Version {version}.",
            created_by=created_by,
            packet_type="follow_up_packet",
            version_number=version,
        )
        return None


def generate_followup_packet_pdf_logged_versioned(intake_id, created_by=None):
    version = get_next_export_version(intake_id, "pdf", "follow_up_packet")

    try:
        path = generate_followup_packet_pdf(intake_id)
        if path:
            log_intake_export_versioned(
                intake_id=intake_id,
                export_type="pdf",
                export_status="success",
                file_path=path,
                message=f"PDF follow-up packet generated successfully. Version {version}.",
                created_by=created_by,
                packet_type="follow_up_packet",
                version_number=version,
            )
            return path

        log_intake_export_versioned(
            intake_id=intake_id,
            export_type="pdf",
            export_status="failed",
            file_path=None,
            message=f"Automatic PDF generation unavailable. Use Print packet → Save as PDF or install LibreOffice / pywin32 + Word. Version {version}.",
            created_by=created_by,
            packet_type="follow_up_packet",
            version_number=version,
        )
        return None

    except Exception as exc:
        log_intake_export_versioned(
            intake_id=intake_id,
            export_type="pdf",
            export_status="error",
            file_path=None,
            message=f"{exc} Version {version}.",
            created_by=created_by,
            packet_type="follow_up_packet",
            version_number=version,
        )
        return None


# -------------------------------------------------------------------
# INT-1N-SCOPE — Export History Firm Scope Repair
# -------------------------------------------------------------------

def list_all_intake_export_logs_any_scope(limit=100):
    """
    Admin/local fallback: returns export logs across all firm scopes.
    Used only when the active firm-scoped export dashboard is empty.
    """
    ensure_intake_export_version_columns()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, export_type, export_status, file_path, message,
               created_at, created_by, version_number, packet_type, firm_id
        FROM intake_export_logs
        ORDER BY id DESC
        LIMIT ?
    """, (limit,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": row[0],
            "export_type": row[1],
            "export_status": row[2],
            "file_path": row[3],
            "message": row[4],
            "created_at": format_intake_timestamp(row[5]),
            "created_by": row[6] or "—",
            "version_number": row[7] or 1,
            "packet_type": row[8] or "follow_up_packet",
            "firm_id": row[9] or "—",
        }
        for row in rows
    ]


def get_intake_export_summary_from_logs(logs):
    summary = {
        "total": len(logs or []),
        "success": 0,
        "failed": 0,
        "error": 0,
        "docx": 0,
        "pdf": 0,
        "other": 0,
    }

    for log in logs or []:
        status = log.get("export_status")
        export_type = log.get("export_type")

        if status in summary:
            summary[status] += 1

        if export_type == "docx":
            summary["docx"] += 1
        elif export_type == "pdf":
            summary["pdf"] += 1
        else:
            summary["other"] += 1

    return summary


def list_intake_export_logs_dashboard(limit=200):
    """
    Firm-aware first, local-admin fallback second.
    Returns (logs, summary, scope_label).
    """
    scoped_logs = list_all_intake_export_logs(limit=limit)

    if scoped_logs:
        return scoped_logs, get_intake_export_summary_from_logs(scoped_logs), "Active Firm Scope"

    fallback_logs = list_all_intake_export_logs_any_scope(limit=limit)
    return fallback_logs, get_intake_export_summary_from_logs(fallback_logs), "All Local Export Logs"


# -------------------------------------------------------------------
# INT-1O — Intake Module Completion Ledger
# -------------------------------------------------------------------

VALID_MODULE_LEDGER_STATUSES = {
    "completed": "Completed",
    "locked": "Locked",
    "pending": "Pending",
    "fallback": "Fallback / Limited",
    "failed": "Failed",
    "skipped": "Skipped",
}

DEFAULT_INTAKE_MODULE_LEDGER = [
    ("INT-1A", "Intake Lane Map", "locked", "Classifies the user into the correct intake lane."),
    ("INT-1B", "Intake Translation Map", "locked", "Translates plain-language answers into categories, modules, document requests, and risk flags."),
    ("INT-1C", "Intake Scoring Engine", "locked", "Scores complexity, urgency, and readiness."),
    ("INT-1D", "Client-Facing Fiduciary Snapshot", "locked", "Creates calm client-facing intake snapshot."),
    ("INT-1E", "Save Snapshot + Intake Dashboard List", "locked", "Stores and lists intake records."),
    ("INT-1F", "Resume + Completion Status Controls", "locked", "Separates completed and incomplete intakes."),
    ("INT-1G", "Dashboard Polish + Print/Export Prep", "locked", "Adds dashboard filters, print view, and export-prep screen."),
    ("INT-1H", "Internal Review Notes", "locked", "Adds internal staff/fiduciary notes."),
    ("INT-1I", "Follow-Up Task Builder", "locked", "Creates structured follow-up tasks from intake outputs."),
    ("INT-1J", "Task Filters + Workflow Polish", "locked", "Groups and summarizes follow-up tasks."),
    ("INT-1K", "Follow-Up Packet / Checklist Export Prep", "locked", "Compiles packet-prep printable checklist view."),
    ("INT-1L-DOCX", "Follow-Up Packet DOCX Export", "locked", "Generates downloadable DOCX follow-up packet."),
    ("INT-1L-PDF", "Follow-Up Packet PDF Export", "fallback", "Automatic PDF export requires LibreOffice or Word/pywin32; browser print-to-PDF remains available."),
    ("INT-1M", "Export Gitignore + Audit Log", "locked", "Ignores generated exports and logs export attempts."),
    ("INT-1N", "Export History Dashboard + Packet Versioning", "locked", "Adds versioned export history."),
    ("INT-1N-SCOPE", "Export History Firm Scope Repair", "locked", "Adds local fallback for all-export dashboard scope."),
]


def ensure_intake_module_ledger_tables():
    ensure_intake_export_version_columns()
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_module_ledger (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            module_code TEXT UNIQUE NOT NULL,
            module_name TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            status_label TEXT,
            description TEXT,
            locked_at TEXT,
            updated_at TEXT,
            updated_by TEXT,
            notes TEXT
        )
    """)

    conn.commit()
    conn.close()


def upsert_intake_module_status(
    module_code,
    module_name,
    status="pending",
    description="",
    notes="",
    updated_by=None
):
    ensure_intake_module_ledger_tables()

    if status not in VALID_MODULE_LEDGER_STATUSES:
        status = "pending"

    status_label = VALID_MODULE_LEDGER_STATUSES.get(status, status)
    now = datetime.utcnow().isoformat(timespec="seconds")
    locked_at = now if status in ["locked", "completed"] else None

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_module_ledger (
            module_code, module_name, status, status_label,
            description, locked_at, updated_at, updated_by, notes
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(module_code) DO UPDATE SET
            module_name = excluded.module_name,
            status = excluded.status,
            status_label = excluded.status_label,
            description = excluded.description,
            locked_at = COALESCE(intake_module_ledger.locked_at, excluded.locked_at),
            updated_at = excluded.updated_at,
            updated_by = excluded.updated_by,
            notes = excluded.notes
    """, (
        module_code,
        module_name,
        status,
        status_label,
        description,
        locked_at,
        now,
        updated_by,
        notes,
    ))

    conn.commit()
    conn.close()


def seed_default_intake_module_ledger(updated_by="system"):
    ensure_intake_module_ledger_tables()

    for code, name, status, description in DEFAULT_INTAKE_MODULE_LEDGER:
        upsert_intake_module_status(
            module_code=code,
            module_name=name,
            status=status,
            description=description,
            notes="Seeded from confirmed INT-1 build sequence.",
            updated_by=updated_by,
        )


def list_intake_module_ledger():
    ensure_intake_module_ledger_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT module_code, module_name, status, status_label,
               description, locked_at, updated_at, updated_by, notes
        FROM intake_module_ledger
        ORDER BY id ASC
    """)

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "module_code": row[0],
            "module_name": row[1],
            "status": row[2],
            "status_label": row[3] or VALID_MODULE_LEDGER_STATUSES.get(row[2], row[2]),
            "description": row[4],
            "locked_at": format_intake_timestamp(row[5]) if row[5] else "",
            "updated_at": format_intake_timestamp(row[6]) if row[6] else "",
            "updated_by": row[7] or "—",
            "notes": row[8] or "",
        }
        for row in rows
    ]


def summarize_intake_module_ledger(modules):
    summary = {
        "total": len(modules or []),
        "locked": 0,
        "completed": 0,
        "pending": 0,
        "fallback": 0,
        "failed": 0,
        "skipped": 0,
    }

    for module in modules or []:
        status = module.get("status")
        if status in summary:
            summary[status] += 1

    return summary


# -------------------------------------------------------------------
# INT-2A — Intake-to-Document Recommendation Engine
# -------------------------------------------------------------------

DOCUMENT_WORKFLOW_CATALOG = {
    "foundational_estate_package": {
        "title": "Foundational Estate Planning Package",
        "description": "Start a basic planning package covering family structure, decision-makers, asset overview, and core estate documents.",
        "workflow_type": "planning",
        "default_priority": "high",
    },
    "document_audit": {
        "title": "Existing Document Audit",
        "description": "Review existing wills, trusts, powers of attorney, deeds, beneficiary forms, insurance policies, and related documents.",
        "workflow_type": "review",
        "default_priority": "high",
    },
    "trust_document_review": {
        "title": "Trust Document Review",
        "description": "Review existing trust documents, trustee authority, beneficiaries, amendments, funding status, and execution readiness.",
        "workflow_type": "review",
        "default_priority": "high",
    },
    "real_property_review": {
        "title": "Real Property Review",
        "description": "Review deeds, ownership, mortgages, insurance, tax bills, property type, transfer readiness, and funding risks.",
        "workflow_type": "asset",
        "default_priority": "high",
    },
    "business_continuity_packet": {
        "title": "Business Continuity Packet",
        "description": "Review entity documents, operating agreements, EIN records, business authority, succession planning, and liability exposure.",
        "workflow_type": "business",
        "default_priority": "high",
    },
    "fiduciary_authority_review": {
        "title": "Fiduciary Authority Review",
        "description": "Confirm trustee, executor, agent, or administrator authority before fiduciary action is taken.",
        "workflow_type": "fiduciary",
        "default_priority": "high",
    },
    "professional_review_checklist": {
        "title": "Professional Review Checklist",
        "description": "Prepare issues that may require legal, tax, court, creditor, or professional review before documents or transfers proceed.",
        "workflow_type": "professional_review",
        "default_priority": "urgent",
    },
    "beneficiary_guardian_planning": {
        "title": "Beneficiary / Guardian Planning",
        "description": "Review beneficiaries, minor children, guardianship considerations, special-needs concerns, and distribution objectives.",
        "workflow_type": "family",
        "default_priority": "high",
    },
    "asset_inventory_packet": {
        "title": "Asset Inventory Packet",
        "description": "Build a structured inventory of assets, documents, ownership records, account information, and supporting evidence.",
        "workflow_type": "inventory",
        "default_priority": "normal",
    },
    "next_session_agenda": {
        "title": "Next Session Agenda",
        "description": "Prepare a focused agenda for the next follow-up meeting based on the intake snapshot and task list.",
        "workflow_type": "session",
        "default_priority": "normal",
    },
}


def _recommendation_unique(recommendations):
    seen = set()
    output = []

    for item in recommendations:
        key = item.get("workflow_key")
        if not key or key in seen:
            continue
        seen.add(key)
        output.append(item)

    priority_order = {
        "urgent": 1,
        "high": 2,
        "normal": 3,
        "low": 4,
    }

    output.sort(key=lambda r: (priority_order.get(r.get("priority"), 5), -int(r.get("confidence", 0))))
    return output


def _make_document_recommendation(workflow_key, reason, priority=None, confidence=75, source="engine"):
    workflow = DOCUMENT_WORKFLOW_CATALOG.get(workflow_key)
    if not workflow:
        return None

    return {
        "workflow_key": workflow_key,
        "title": workflow["title"],
        "description": workflow["description"],
        "workflow_type": workflow["workflow_type"],
        "priority": priority or workflow["default_priority"],
        "confidence": confidence,
        "reason": reason,
        "source": source,
        "status": "recommended",
    }


def build_document_recommendations(intake_id):
    """
    Converts intake snapshot data into recommended next document workflows.
    This does not generate documents. It only recommends the next workflow path.
    """
    packet = build_intake_followup_packet(intake_id)
    if not packet:
        return None

    snapshot = packet.get("snapshot", {}) or {}
    result = packet.get("result", {}) or {}
    summary = result.get("summary", {}) or {}
    scores = result.get("scores", {}) or {}

    categories = set(summary.get("system_categories", []) or [])
    meanings = set(summary.get("system_meanings", []) or [])
    module_triggers = set(summary.get("module_triggers", []) or [])
    documents = set(summary.get("document_requests", []) or [])
    next_sessions = set(summary.get("next_sessions", []) or [])
    risk_flags = set(summary.get("risk_flags", []) or [])

    recommendations = []

    def add(key, reason, priority=None, confidence=75, source="engine"):
        rec = _make_document_recommendation(
            workflow_key=key,
            reason=reason,
            priority=priority,
            confidence=confidence,
            source=source,
        )
        if rec:
            recommendations.append(rec)

    # Professional review / risk-first logic
    high_risk_flags = {
        "urgent_or_legal_review_flag",
        "tax_review_flag",
        "creditor_pressure_flag",
        "incapacity_planning_needed",
        "family_conflict_risk",
        "business_liability_possible",
    }

    if risk_flags.intersection(high_risk_flags) or scores.get("urgency_level") == "High":
        add(
            "professional_review_checklist",
            "High urgency or review flags were detected during intake.",
            priority="urgent",
            confidence=95,
        )

    # Foundational planning
    if "DOCUMENT_STATUS" in categories and ("documentation_gap" in risk_flags or "no_documents_available" in meanings):
        add(
            "foundational_estate_package",
            "The intake indicates missing or limited documents and foundational planning may be needed.",
            priority="high",
            confidence=88,
        )

    if "foundational_planning_review" in next_sessions:
        add(
            "foundational_estate_package",
            "The recommended next session points toward foundational planning.",
            priority="high",
            confidence=85,
        )

    # Document audit
    if "document_audit" in module_triggers or "document_audit_session" in next_sessions or "review_documents" in meanings:
        add(
            "document_audit",
            "Existing documents or document review objectives were detected.",
            priority="high",
            confidence=88,
        )

    if documents.intersection({"will", "existing_will", "trust_document", "power_of_attorney", "health_directive", "beneficiary_forms"}):
        add(
            "document_audit",
            "Estate, authority, or beneficiary documents were listed for review.",
            priority="high",
            confidence=82,
        )

    # Trust review
    if "trust_exists" in meanings or "trust_audit" in module_triggers or "trust_document_review" in next_sessions:
        add(
            "trust_document_review",
            "Trust documents or trust review triggers were detected.",
            priority="high",
            confidence=90,
        )

    if "FIDUCIARY_CONTEXT" in categories or "authority_review_needed" in risk_flags:
        add(
            "fiduciary_authority_review",
            "The intake indicates fiduciary authority, trust/estate administration, or authority review concerns.",
            priority="high",
            confidence=90,
        )

    # Real property
    if "ASSET_PROFILE" in categories and (
        "real_property_review" in module_triggers
        or "real_property_deep_dive" in next_sessions
        or documents.intersection({"deed", "deeds", "mortgage_statement", "property_tax_bill", "tax_bill"})
    ):
        add(
            "real_property_review",
            "Real property assets or property documents were identified.",
            priority="high",
            confidence=92,
        )

    # Business continuity
    if "BUSINESS_PROFILE" in categories or "business_continuity" in module_triggers or "business_continuity_review" in next_sessions:
        add(
            "business_continuity_packet",
            "Business ownership, business assets, or business continuity concerns were identified.",
            priority="high",
            confidence=90,
        )

    # Beneficiary / guardian planning
    beneficiary_triggers = {
        "minor_children_flag",
        "special_needs_flag",
        "guardian_review",
        "minor_beneficiary_controls",
        "children_guardian_review",
        "beneficiary_planning_review",
    }

    if "BENEFICIARY_PROFILE" in categories or risk_flags.intersection(beneficiary_triggers) or module_triggers.intersection(beneficiary_triggers) or next_sessions.intersection(beneficiary_triggers):
        add(
            "beneficiary_guardian_planning",
            "Beneficiary, child, guardian, or family planning issues were detected.",
            priority="high",
            confidence=86,
        )

    # Asset inventory
    if "ASSET_PROFILE" in categories or "asset_inventory" in module_triggers or "asset_document_deep_dive" in next_sessions:
        add(
            "asset_inventory_packet",
            "Assets and supporting records need to be organized before deeper review.",
            priority="normal",
            confidence=80,
        )

    # Always recommend next-session agenda once a packet exists.
    if packet.get("recommended_next_session"):
        add(
            "next_session_agenda",
            "A recommended next session was generated from the intake snapshot.",
            priority="normal",
            confidence=78,
        )

    recommendations = _recommendation_unique(recommendations)

    return {
        "intake_id": intake_id,
        "packet": packet,
        "recommendations": recommendations,
        "recommendation_count": len(recommendations),
        "scores": scores,
        "summary": summary,
    }


def ensure_intake_document_recommendation_tables():
    ensure_intake_module_ledger_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_document_recommendations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            workflow_key TEXT NOT NULL,
            title TEXT,
            workflow_type TEXT,
            priority TEXT,
            confidence INTEGER,
            reason TEXT,
            source TEXT,
            status TEXT DEFAULT 'recommended',
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def save_document_recommendations(intake_id, recommendations, created_by=None):
    ensure_intake_document_recommendation_tables()

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    for rec in recommendations or []:
        workflow_key = rec.get("workflow_key")
        if not workflow_key:
            continue

        cur.execute("""
            SELECT id, status
            FROM intake_document_recommendations
            WHERE intake_id = ? AND workflow_key = ?
            LIMIT 1
        """, (intake_id, workflow_key))
        existing = cur.fetchone()

        if existing:
            existing_id = existing[0]
            existing_status = existing[1] or "recommended"

            preserved_status = existing_status if existing_status in {
                "accepted",
                "deferred",
                "rejected",
                "launch_prepared",
            } else rec.get("status", "recommended")

            cur.execute("""
                UPDATE intake_document_recommendations
                SET title = ?,
                    workflow_type = ?,
                    priority = ?,
                    confidence = ?,
                    reason = ?,
                    source = ?,
                    status = ?,
                    updated_at = ?,
                    created_by = ?
                WHERE id = ?
            """, (
                rec.get("title"),
                rec.get("workflow_type"),
                rec.get("priority"),
                int(rec.get("confidence", 0)),
                rec.get("reason"),
                rec.get("source"),
                preserved_status,
                now,
                created_by,
                existing_id,
            ))
        else:
            cur.execute("""
                INSERT INTO intake_document_recommendations (
                    intake_id, firm_id, workflow_key, title, workflow_type,
                    priority, confidence, reason, source, status,
                    created_at, updated_at, created_by
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                intake_id,
                firm_id,
                workflow_key,
                rec.get("title"),
                rec.get("workflow_type"),
                rec.get("priority"),
                int(rec.get("confidence", 0)),
                rec.get("reason"),
                rec.get("source"),
                rec.get("status", "recommended"),
                now,
                now,
                created_by,
            ))

    conn.commit()
    conn.close()

def list_saved_document_recommendations(intake_id):
    ensure_intake_document_recommendation_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT workflow_key, title, workflow_type, priority, confidence,
               reason, source, status, created_at, updated_at, created_by
        FROM intake_document_recommendations
        WHERE intake_id = ?
        ORDER BY
            CASE priority
                WHEN 'urgent' THEN 1
                WHEN 'high' THEN 2
                WHEN 'normal' THEN 3
                WHEN 'low' THEN 4
                ELSE 5
            END,
            confidence DESC,
            id ASC
    """, (intake_id,))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "workflow_key": row[0],
            "title": row[1],
            "workflow_type": row[2],
            "priority": row[3],
            "confidence": row[4],
            "reason": row[5],
            "source": row[6],
            "status": row[7],
            "created_at": format_intake_timestamp(row[8]),
            "updated_at": format_intake_timestamp(row[9]),
            "created_by": row[10] or "—",
        }
        for row in rows
    ]


# -------------------------------------------------------------------
# INT-2B — Recommendation Rule Expansion + Confidence Tuning
# -------------------------------------------------------------------

def _text_signal_contains(text_values, keywords):
    combined = " ".join(str(v or "").lower() for v in text_values)
    return any(keyword.lower() in combined for keyword in keywords)


def _collect_recommendation_signal_text(packet):
    values = []

    for key in ["documents", "priorities", "review_flags"]:
        values.extend(packet.get(key, []) or [])

    for task in packet.get("tasks", []) or []:
        values.append(task.get("title"))
        values.append(task.get("description"))
        values.append(task.get("task_type_label"))
        values.append(task.get("status_label"))

    for note in packet.get("notes", []) or []:
        values.append(note.get("note_body"))
        values.append(note.get("note_type_label"))

    snapshot = packet.get("snapshot", {}) or {}
    values.extend(snapshot.get("documents_to_gather", []) or [])
    values.extend(snapshot.get("top_priorities", []) or [])
    values.extend(snapshot.get("review_flags", []) or [])
    values.append(snapshot.get("recommended_next_session"))

    return values


def tune_recommendation_confidence(base_confidence, scores, signal_strength=0):
    confidence = int(base_confidence or 75)

    urgency = (scores or {}).get("urgency_level")
    complexity = (scores or {}).get("complexity_level")
    readiness = (scores or {}).get("readiness_level")

    if urgency == "High":
        confidence += 5
    elif urgency == "Medium":
        confidence += 2

    if complexity in ["Advanced", "Complex"]:
        confidence += 4
    elif complexity == "Moderate":
        confidence += 2

    if readiness == "Not Ready":
        confidence += 3
    elif readiness == "Partially Ready":
        confidence += 1

    confidence += int(signal_strength or 0)

    if confidence > 99:
        confidence = 99
    if confidence < 50:
        confidence = 50

    return confidence


def expand_document_recommendations(recommendation_result):
    """
    Adds additional recommendations based on task text, packet contents,
    document requests, review flags, and score tuning.
    """
    if not recommendation_result:
        return None

    intake_id = recommendation_result.get("intake_id")
    packet = recommendation_result.get("packet", {}) or {}
    summary = recommendation_result.get("summary", {}) or {}
    scores = recommendation_result.get("scores", {}) or {}

    categories = set(summary.get("system_categories", []) or [])
    meanings = set(summary.get("system_meanings", []) or [])
    module_triggers = set(summary.get("module_triggers", []) or [])
    documents = set(summary.get("document_requests", []) or [])
    next_sessions = set(summary.get("next_sessions", []) or [])
    risk_flags = set(summary.get("risk_flags", []) or [])

    signal_text = _collect_recommendation_signal_text(packet)

    recommendations = list(recommendation_result.get("recommendations", []) or [])

    def existing_keys():
        return {r.get("workflow_key") for r in recommendations}

    def add(key, reason, priority=None, confidence=75, signal_strength=0, source="engine_tuned"):
        if key in existing_keys():
            # Tune existing recommendation upward if this rule strengthens it.
            for rec in recommendations:
                if rec.get("workflow_key") == key:
                    tuned = tune_recommendation_confidence(confidence, scores, signal_strength=signal_strength)
                    if tuned > int(rec.get("confidence", 0)):
                        rec["confidence"] = tuned
                        rec["reason"] = rec.get("reason") + " Additional signal: " + reason
                        if priority == "urgent" or (priority == "high" and rec.get("priority") not in ["urgent"]):
                            rec["priority"] = priority
                    return

        rec = _make_document_recommendation(
            workflow_key=key,
            reason=reason,
            priority=priority,
            confidence=tune_recommendation_confidence(confidence, scores, signal_strength=signal_strength),
            source=source,
        )
        if rec:
            recommendations.append(rec)

    # ------------------------------------------------------------
    # Business continuity expansion
    # ------------------------------------------------------------
    business_keywords = [
        "business", "operating agreement", "ein", "business license",
        "liability", "partner", "partnership", "company", "entity",
        "bank authority", "succession"
    ]

    if (
        "BUSINESS_PROFILE" in categories
        or "business_liability_possible" in risk_flags
        or "business_continuity_needed" in risk_flags
        or "business_continuity" in module_triggers
        or "business_governance_review" in next_sessions
        or documents.intersection({"operating_agreement", "business_license", "ein_letter", "partnership_agreement", "bank_authority_records"})
        or _text_signal_contains(signal_text, business_keywords)
    ):
        add(
            "business_continuity_packet",
            "Business-related documents, liability signals, partner/entity references, or business continuity tasks were detected.",
            priority="high",
            confidence=88,
            signal_strength=4,
        )

    # ------------------------------------------------------------
    # Real property expansion
    # ------------------------------------------------------------
    property_keywords = [
        "deed", "title", "mortgage", "property", "tax bill", "lease",
        "land", "real property", "insurance policy"
    ]

    if (
        "real_property_review" in module_triggers
        or "real_property_deep_dive" in next_sessions
        or documents.intersection({"deed", "deeds", "mortgage_statement", "property_tax_bill", "tax_bill", "lease"})
        or _text_signal_contains(signal_text, property_keywords)
    ):
        add(
            "real_property_review",
            "Property documents or real-property task signals were detected.",
            priority="high",
            confidence=90,
            signal_strength=3,
        )

    # ------------------------------------------------------------
    # Tax/professional review expansion
    # ------------------------------------------------------------
    tax_keywords = [
        "tax", "irs", "notice", "filing", "court", "legal", "creditor",
        "lawsuit", "claim", "professional review"
    ]

    if (
        "tax_review_flag" in risk_flags
        or "urgent_or_legal_review_flag" in risk_flags
        or "tax_review_referral" in next_sessions
        or documents.intersection({"tax_notices", "tax_filings", "court_documents", "claim_letters_if_available"})
        or _text_signal_contains(signal_text, tax_keywords)
    ):
        add(
            "professional_review_checklist",
            "Tax, legal, court, creditor, or professional-review signals were detected.",
            priority="urgent",
            confidence=94,
            signal_strength=5,
        )

    # ------------------------------------------------------------
    # Document audit expansion
    # ------------------------------------------------------------
    doc_audit_keywords = [
        "will", "trust", "power of attorney", "health directive",
        "beneficiary", "insurance", "existing document", "document checklist",
        "document follow-up"
    ]

    if (
        "DOCUMENT_STATUS" in categories
        or "document_audit" in module_triggers
        or "document_collection_review" in next_sessions
        or "documentation_gap" in risk_flags
        or _text_signal_contains(signal_text, doc_audit_keywords)
    ):
        add(
            "document_audit",
            "Document gaps, existing document references, or document follow-up tasks were detected.",
            priority="high",
            confidence=84,
            signal_strength=3,
        )

    # ------------------------------------------------------------
    # Foundational estate package expansion
    # ------------------------------------------------------------
    foundational_keywords = [
        "not ready", "missing", "family structure", "decision-maker",
        "guardian", "beneficiary", "foundational", "initial planning"
    ]

    if (
        scores.get("readiness_level") == "Not Ready"
        or "documentation_gap" in risk_flags
        or "foundational_planning_review" in next_sessions
        or _text_signal_contains(signal_text, foundational_keywords)
    ):
        add(
            "foundational_estate_package",
            "Readiness is low or foundational planning/document gaps were detected.",
            priority="high",
            confidence=86,
            signal_strength=3,
        )

    # ------------------------------------------------------------
    # Fiduciary authority expansion
    # ------------------------------------------------------------
    fiduciary_keywords = [
        "trustee", "executor", "fiduciary", "authority", "poa",
        "power of attorney", "administrator", "letters testamentary"
    ]

    if (
        "FIDUCIARY_CONTEXT" in categories
        or "authority_review_needed" in risk_flags
        or "fiduciary_authority_review" in next_sessions
        or documents.intersection({"power_of_attorney", "letters_testamentary_or_authority", "estate_authority_document"})
        or _text_signal_contains(signal_text, fiduciary_keywords)
    ):
        add(
            "fiduciary_authority_review",
            "Fiduciary authority, trustee/executor, or power-of-attorney signals were detected.",
            priority="high",
            confidence=87,
            signal_strength=3,
        )

    # ------------------------------------------------------------
    # Beneficiary / guardian expansion
    # ------------------------------------------------------------
    family_keywords = [
        "minor child", "guardian", "children", "beneficiary",
        "special-needs", "dependent", "spouse", "family"
    ]

    if (
        "BENEFICIARY_PROFILE" in categories
        or "minor_children_flag" in risk_flags
        or "special_needs_flag" in risk_flags
        or "children_guardian_review" in next_sessions
        or "beneficiary_planning_review" in next_sessions
        or _text_signal_contains(signal_text, family_keywords)
    ):
        add(
            "beneficiary_guardian_planning",
            "Beneficiary, child, guardian, dependent, or family-planning signals were detected.",
            priority="high",
            confidence=86,
            signal_strength=3,
        )

    # ------------------------------------------------------------
    # Asset inventory expansion
    # ------------------------------------------------------------
    asset_keywords = [
        "asset", "inventory", "account", "bank", "vehicle", "insurance",
        "title", "statement", "property", "business"
    ]

    if (
        "ASSET_PROFILE" in categories
        or "asset_document_deep_dive" in next_sessions
        or "asset_inventory" in module_triggers
        or _text_signal_contains(signal_text, asset_keywords)
    ):
        add(
            "asset_inventory_packet",
            "Asset, account, title, statement, or inventory signals were detected.",
            priority="normal",
            confidence=82,
            signal_strength=2,
        )

    # Always keep next-session agenda if a recommended next session exists.
    if packet.get("recommended_next_session"):
        add(
            "next_session_agenda",
            "Recommended next session exists and should be converted into a working agenda.",
            priority="normal",
            confidence=78,
            signal_strength=1,
        )

    recommendations = _recommendation_unique(recommendations)

    recommendation_result["recommendations"] = recommendations
    recommendation_result["recommendation_count"] = len(recommendations)
    recommendation_result["tuning_applied"] = True

    return recommendation_result


def build_document_recommendations_tuned(intake_id):
    base = build_document_recommendations(intake_id)
    if not base:
        return None
    return expand_document_recommendations(base)


# -------------------------------------------------------------------
# INT-2C — Recommendation Selection + Workflow Launch Prep
# -------------------------------------------------------------------

VALID_RECOMMENDATION_STATUSES = {
    "recommended": "Recommended",
    "accepted": "Accepted",
    "deferred": "Deferred",
    "rejected": "Rejected",
    "launch_prepared": "Launch Prepared",
}


def update_document_recommendation_status(
    intake_id,
    workflow_key,
    status,
    updated_by=None
):
    ensure_intake_document_recommendation_tables()

    if status not in VALID_RECOMMENDATION_STATUSES:
        raise ValueError("Invalid recommendation status.")

    now = datetime.utcnow().isoformat(timespec="seconds")

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        UPDATE intake_document_recommendations
        SET status = ?, updated_at = ?, created_by = ?
        WHERE intake_id = ? AND workflow_key = ?
    """, (
        status,
        now,
        updated_by,
        intake_id,
        workflow_key,
    ))

    conn.commit()
    conn.close()


def get_document_recommendation(intake_id, workflow_key):
    ensure_intake_document_recommendation_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT workflow_key, title, workflow_type, priority, confidence,
               reason, source, status, created_at, updated_at, created_by
        FROM intake_document_recommendations
        WHERE intake_id = ? AND workflow_key = ?
        LIMIT 1
    """, (intake_id, workflow_key))

    row = cur.fetchone()
    conn.close()

    if not row:
        return None

    return {
        "workflow_key": row[0],
        "title": row[1],
        "workflow_type": row[2],
        "priority": row[3],
        "confidence": row[4],
        "reason": row[5],
        "source": row[6],
        "status": row[7],
        "created_at": format_intake_timestamp(row[8]),
        "updated_at": format_intake_timestamp(row[9]),
        "created_by": row[10] or "—",
        "status_label": VALID_RECOMMENDATION_STATUSES.get(row[7], row[7]),
    }


def build_workflow_launch_prep(intake_id, workflow_key):
    """
    Prepares the selected recommendation for workflow launch.
    This does not generate the final document.
    """
    recommendation = get_document_recommendation(intake_id, workflow_key)

    if not recommendation:
        # Try rebuilding and saving recommendations if the row does not exist yet.
        result = build_document_recommendations_tuned(intake_id)
        if result:
            save_document_recommendations(intake_id, result.get("recommendations", []), created_by="system")
            recommendation = get_document_recommendation(intake_id, workflow_key)

    if not recommendation:
        return None

    packet = build_intake_followup_packet(intake_id)
    if not packet:
        return None

    workflow_key = recommendation.get("workflow_key")

    workflow_inputs = []
    workflow_checks = []
    workflow_warnings = []

    documents = packet.get("documents", []) or []
    review_flags = packet.get("review_flags", []) or []
    tasks = packet.get("tasks", []) or []

    # Generic inputs
    workflow_inputs.append("Confirm client identity / planning party.")
    workflow_inputs.append("Review intake snapshot and follow-up packet.")
    workflow_inputs.append("Confirm whether client-facing or internal-only workflow is being launched.")

    # Workflow-specific prep
    if workflow_key == "professional_review_checklist":
        workflow_inputs.extend([
            "Collect tax notices, court notices, claim letters, creditor letters, or urgent correspondence.",
            "Separate legal, tax, and administrative issues before document drafting.",
            "Identify whether outside professional review is required before action.",
        ])
        workflow_checks.extend([
            "Do not finalize documents until high-risk issues are reviewed.",
            "Confirm whether the issue is legal, tax, creditor, court, or administrative.",
        ])

    elif workflow_key == "business_continuity_packet":
        workflow_inputs.extend([
            "Collect operating agreement, EIN letter, business license, insurance policies, and bank authority records.",
            "Identify owners, partners, managers, authorized signers, and successor operators.",
            "Review business liability and continuity risks.",
        ])
        workflow_checks.extend([
            "Confirm entity type and current authority before preparing business documents.",
            "Separate business assets from personal/family assets where needed.",
        ])

    elif workflow_key == "real_property_review":
        workflow_inputs.extend([
            "Collect deed, mortgage statement, tax bill, insurance declaration, lease, and title records.",
            "Confirm owner name, vesting, liens, co-owners, and transfer restrictions.",
            "Identify whether the property is personal residence, rental, inherited, business, or trust-related.",
        ])
        workflow_checks.extend([
            "Do not treat property as fundable until title and authority are verified.",
            "Check for mortgage, lien, co-owner, or tax restrictions.",
        ])

    elif workflow_key == "foundational_estate_package":
        workflow_inputs.extend([
            "Confirm planning parties, family structure, beneficiaries, decision-makers, and primary assets.",
            "Determine whether will, trust, POA, health directive, or beneficiary updates are needed.",
            "Collect any existing estate or authority documents.",
        ])
        workflow_checks.extend([
            "Confirm whether this is new planning or revision of existing documents.",
            "Do not skip authority and beneficiary review.",
        ])

    elif workflow_key == "document_audit":
        workflow_inputs.extend([
            "Collect all existing wills, trusts, POAs, health directives, deeds, beneficiary forms, insurance policies, and entity documents.",
            "Sort documents by type, date, signer, notary/witness status, and current relevance.",
            "Identify missing signatures, outdated provisions, or conflicting documents.",
        ])
        workflow_checks.extend([
            "Do not assume older documents are revoked or valid without review.",
            "Flag unsigned, undated, incomplete, or conflicting documents.",
        ])

    elif workflow_key == "beneficiary_guardian_planning":
        workflow_inputs.extend([
            "Confirm beneficiaries, minors, guardians, dependents, special-needs concerns, and distribution preferences.",
            "Collect beneficiary designation forms and any guardian/care documents.",
            "Identify potential conflict, dependency, or incapacity issues.",
        ])
        workflow_checks.extend([
            "Do not finalize beneficiary planning without confirming current family structure.",
            "Flag minor-child or special-needs matters for careful review.",
        ])

    elif workflow_key == "asset_inventory_packet":
        workflow_inputs.extend([
            "Build an asset list with ownership, estimated value, supporting documents, and transfer readiness.",
            "Collect statements, titles, deeds, insurance, business records, and account registration information.",
            "Separate assets by personal, trust, business, heritage, digital, and financial category.",
        ])
        workflow_checks.extend([
            "Do not recommend transfer until ownership and restrictions are verified.",
            "Track missing documents as follow-up tasks.",
        ])

    elif workflow_key == "fiduciary_authority_review":
        workflow_inputs.extend([
            "Collect trustee, executor, agent, POA, letters testamentary, or appointment documents.",
            "Confirm who currently has authority and in what capacity.",
            "Review whether authority is current, limited, revoked, or conditional.",
        ])
        workflow_checks.extend([
            "Do not act as fiduciary without authority evidence.",
            "Verify signatures, dates, scope, and applicable limitations.",
        ])

    elif workflow_key == "next_session_agenda":
        workflow_inputs.extend([
            "Use the recommended next session as the agenda anchor.",
            "Pull open tasks, document requests, review flags, and internal notes into the agenda.",
            "Set the next meeting purpose and required documents.",
        ])
        workflow_checks.extend([
            "Keep the next session focused and avoid expanding into final document drafting too early.",
        ])

    else:
        workflow_inputs.extend([
            "Review packet and select appropriate workflow questions.",
            "Gather required documents and open tasks.",
        ])

    if review_flags:
        workflow_warnings.extend(review_flags)

    open_tasks = [task for task in tasks if task.get("status") != "completed"]
    completed_tasks = [task for task in tasks if task.get("status") == "completed"]

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "recommendation": recommendation,
        "packet": packet,
        "documents": documents,
        "review_flags": review_flags,
        "open_tasks": open_tasks,
        "completed_tasks": completed_tasks,
        "workflow_inputs": workflow_inputs,
        "workflow_checks": workflow_checks,
        "workflow_warnings": workflow_warnings,
        "launch_status": "Prepared",
    }


# -------------------------------------------------------------------
# INT-2D — Workflow-Specific Intake Bridge
# -------------------------------------------------------------------

WORKFLOW_BRIDGE_QUESTION_BANK = {
    "professional_review_checklist": {
        "title": "Professional Review Checklist Bridge",
        "purpose": "Clarify high-risk legal, tax, creditor, court, or professional review issues before document drafting.",
        "questions": [
            {
                "key": "review_issue_type",
                "label": "What type of professional review issue is present?",
                "input_type": "checkbox",
                "options": {
                    "tax": "Tax / IRS / filing issue",
                    "legal": "Legal or court issue",
                    "creditor": "Creditor, claim, or collection issue",
                    "business_liability": "Business liability issue",
                    "family_conflict": "Family conflict issue",
                    "not_sure": "Not sure",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "deadline_pressure",
                "label": "Is there a deadline, notice date, court date, or urgent response window?",
                "input_type": "radio",
                "options": {
                    "yes": "Yes",
                    "no": "No",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "documents_available",
                "label": "Which urgent documents are available?",
                "input_type": "checkbox",
                "options": {
                    "tax_notice": "Tax notice",
                    "court_notice": "Court document",
                    "claim_letter": "Claim or creditor letter",
                    "termination_notice": "Termination or adverse-action notice",
                    "insurance_letter": "Insurance letter",
                    "none": "None yet",
                    "other": "Other / not listed",
                },
            },
        ],
    },

    "business_continuity_packet": {
        "title": "Business Continuity Packet Bridge",
        "purpose": "Clarify business authority, documents, ownership, liability, and succession before preparing business continuity documents.",
        "questions": [
            {
                "key": "business_structure",
                "label": "What type of business structure is involved?",
                "input_type": "radio",
                "options": {
                    "sole_prop": "Sole proprietorship / DBA",
                    "llc": "LLC",
                    "corporation": "Corporation",
                    "partnership": "Partnership",
                    "trust_owned": "Trust-owned or estate-connected business",
                    "not_sure": "Not sure",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "business_documents",
                "label": "Which business documents are available?",
                "input_type": "checkbox",
                "options": {
                    "ein_letter": "EIN letter",
                    "operating_agreement": "Operating agreement",
                    "business_license": "Business license / registration",
                    "bank_records": "Bank authority records",
                    "insurance": "Insurance policy",
                    "contracts": "Contracts",
                    "none": "None yet",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "continuity_risk",
                "label": "What continuity risks need review?",
                "input_type": "checkbox",
                "options": {
                    "owner_unavailable": "Owner/operator unavailable",
                    "successor_needed": "Successor manager needed",
                    "partner_dispute": "Partner/co-owner dispute",
                    "liability": "Liability or insurance concern",
                    "records_missing": "Records missing or disorganized",
                    "not_sure": "Not sure",
                    "other": "Other / not listed",
                },
            },
        ],
    },

    "real_property_review": {
        "title": "Real Property Review Bridge",
        "purpose": "Clarify title, ownership, liens, co-owners, insurance, and transfer readiness before property documents are prepared.",
        "questions": [
            {
                "key": "property_type",
                "label": "What type of property is involved?",
                "input_type": "radio",
                "options": {
                    "primary_residence": "Primary residence",
                    "rental": "Rental / investment property",
                    "inherited": "Inherited property",
                    "business": "Business property",
                    "land": "Land / vacant property",
                    "not_sure": "Not sure",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "property_documents",
                "label": "Which property documents are available?",
                "input_type": "checkbox",
                "options": {
                    "deed": "Deed",
                    "tax_bill": "Tax bill",
                    "mortgage_statement": "Mortgage statement",
                    "insurance": "Insurance declaration",
                    "survey": "Survey",
                    "lease": "Lease",
                    "none": "None yet",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "property_risks",
                "label": "Which property issues need review?",
                "input_type": "checkbox",
                "options": {
                    "mortgage": "Mortgage or lien",
                    "co_owner": "Co-owner",
                    "tax_issue": "Tax issue",
                    "insurance_issue": "Insurance issue",
                    "transfer_question": "Transfer/funding question",
                    "not_sure": "Not sure",
                    "other": "Other / not listed",
                },
            },
        ],
    },

    "foundational_estate_package": {
        "title": "Foundational Estate Planning Bridge",
        "purpose": "Clarify the core people, documents, authority roles, and planning purpose before foundational estate documents are prepared.",
        "questions": [
            {
                "key": "planning_documents_needed",
                "label": "Which foundational documents may be needed?",
                "input_type": "checkbox",
                "options": {
                    "will": "Will",
                    "trust": "Trust",
                    "poa": "Power of attorney",
                    "health_directive": "Health directive",
                    "beneficiary_update": "Beneficiary update",
                    "not_sure": "Not sure",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "decision_makers",
                "label": "Have decision-makers been identified?",
                "input_type": "radio",
                "options": {
                    "yes": "Yes",
                    "partial": "Partially",
                    "no": "No",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "family_complexity",
                "label": "Are there family circumstances that may affect planning?",
                "input_type": "checkbox",
                "options": {
                    "minor_children": "Minor children",
                    "blended_family": "Blended family",
                    "special_needs": "Special-needs concern",
                    "family_conflict": "Family conflict",
                    "elder_care": "Elder care concern",
                    "none": "None of these",
                    "not_sure": "Not sure",
                },
            },
        ],
    },

    "document_audit": {
        "title": "Existing Document Audit Bridge",
        "purpose": "Clarify which documents exist, whether they are signed/current, and which may conflict or require deeper review.",
        "questions": [
            {
                "key": "document_types",
                "label": "Which documents should be audited?",
                "input_type": "checkbox",
                "options": {
                    "will": "Will",
                    "trust": "Trust",
                    "poa": "Power of attorney",
                    "health_directive": "Health directive",
                    "deed": "Deed",
                    "beneficiary_form": "Beneficiary form",
                    "business_document": "Business document",
                    "court_document": "Court document",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "execution_status",
                "label": "Do the documents appear signed/witnessed/notarized where needed?",
                "input_type": "radio",
                "options": {
                    "yes": "Yes",
                    "some": "Some documents",
                    "no": "No",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "audit_concerns",
                "label": "What concerns should the audit look for?",
                "input_type": "checkbox",
                "options": {
                    "outdated": "Outdated document",
                    "conflict": "Conflicting documents",
                    "missing_signature": "Missing signature/notary/witness",
                    "wrong_name": "Wrong name or outdated party",
                    "missing_assets": "Assets not addressed",
                    "not_sure": "Not sure",
                },
            },
        ],
    },

    "beneficiary_guardian_planning": {
        "title": "Beneficiary / Guardian Planning Bridge",
        "purpose": "Clarify beneficiaries, guardian choices, dependents, and distribution concerns before beneficiary planning proceeds.",
        "questions": [
            {
                "key": "beneficiary_group",
                "label": "Who needs to be considered in beneficiary planning?",
                "input_type": "checkbox",
                "options": {
                    "spouse": "Spouse",
                    "children": "Children",
                    "minor_children": "Minor children",
                    "grandchildren": "Grandchildren",
                    "parents": "Parents",
                    "charity": "Charity/community purpose",
                    "business_successor": "Business successor",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "guardian_status",
                "label": "Has a guardian or backup caregiver been identified for minors/dependents?",
                "input_type": "radio",
                "options": {
                    "yes": "Yes",
                    "partial": "Partially",
                    "no": "No",
                    "not_applicable": "Not applicable",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "distribution_concerns",
                "label": "Are there distribution concerns?",
                "input_type": "checkbox",
                "options": {
                    "minor_controls": "Minor beneficiary controls",
                    "special_needs": "Special-needs planning",
                    "spendthrift": "Spendthrift or creditor concern",
                    "conflict": "Family conflict",
                    "equal_vs_custom": "Equal vs. custom distribution",
                    "none": "None of these",
                    "not_sure": "Not sure",
                },
            },
        ],
    },

    "asset_inventory_packet": {
        "title": "Asset Inventory Packet Bridge",
        "purpose": "Clarify assets, ownership records, supporting documents, and transfer readiness before inventory documents are prepared.",
        "questions": [
            {
                "key": "asset_categories",
                "label": "Which asset categories should be inventoried?",
                "input_type": "checkbox",
                "options": {
                    "real_property": "Real property",
                    "bank_accounts": "Bank accounts",
                    "vehicles": "Vehicles / titled property",
                    "business": "Business interests",
                    "insurance": "Insurance",
                    "retirement": "Retirement / investment accounts",
                    "digital": "Digital assets",
                    "heritage": "Heritage / legacy assets",
                    "other": "Other / not listed",
                },
            },
            {
                "key": "ownership_clarity",
                "label": "Is ownership clear for the main assets?",
                "input_type": "radio",
                "options": {
                    "yes": "Yes",
                    "some": "Some assets",
                    "no": "No",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "inventory_documents",
                "label": "Which inventory records are available?",
                "input_type": "checkbox",
                "options": {
                    "statements": "Statements",
                    "titles": "Titles",
                    "deeds": "Deeds",
                    "insurance": "Insurance records",
                    "registrations": "Account registrations",
                    "appraisals": "Appraisals/valuations",
                    "none": "None yet",
                    "other": "Other / not listed",
                },
            },
        ],
    },

    "next_session_agenda": {
        "title": "Next Session Agenda Bridge",
        "purpose": "Convert the recommended next session into a focused agenda with documents, decisions, and follow-up topics.",
        "questions": [
            {
                "key": "next_session_focus",
                "label": "What should the next session focus on first?",
                "input_type": "radio",
                "options": {
                    "documents": "Documents to gather/review",
                    "assets": "Asset organization",
                    "family_roles": "Family roles and beneficiaries",
                    "risk": "Risk/professional review",
                    "business": "Business continuity",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "session_attendees",
                "label": "Who should attend the next session?",
                "input_type": "checkbox",
                "options": {
                    "client": "Client/planning party",
                    "spouse": "Spouse/partner",
                    "trustee": "Trustee/fiduciary",
                    "business_partner": "Business partner",
                    "advisor": "Professional advisor",
                    "family_member": "Family member",
                    "not_sure": "Not sure",
                },
            },
            {
                "key": "session_readiness",
                "label": "What must be ready before the next session?",
                "input_type": "checkbox",
                "options": {
                    "documents": "Documents gathered",
                    "asset_list": "Asset list drafted",
                    "questions": "Questions prepared",
                    "decision_makers": "Decision-makers identified",
                    "professional_review": "Professional review item identified",
                    "not_sure": "Not sure",
                },
            },
        ],
    },
}


def get_workflow_bridge_definition(workflow_key):
    return WORKFLOW_BRIDGE_QUESTION_BANK.get(workflow_key)


def ensure_workflow_bridge_tables():
    ensure_intake_document_recommendation_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_workflow_bridge_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            question_key TEXT NOT NULL,
            answer_key TEXT,
            answer_label TEXT,
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def save_workflow_bridge_answers(intake_id, workflow_key, form_data, created_by=None):
    ensure_workflow_bridge_tables()

    definition = get_workflow_bridge_definition(workflow_key)
    if not definition:
        raise ValueError("Unknown workflow bridge.")

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    try:
        # Replace current bridge answers for this workflow/intake.
        cur.execute("""
            DELETE FROM intake_workflow_bridge_answers
            WHERE intake_id = ? AND workflow_key = ?
        """, (intake_id, workflow_key))

        for question in definition.get("questions", []):
            qkey = question.get("key")
            options = question.get("options", {})
            input_type = question.get("input_type")

            if input_type == "checkbox":
                values = form_data.getlist(qkey) if hasattr(form_data, "getlist") else []
            else:
                value = form_data.get(qkey) if hasattr(form_data, "get") else None
                values = [value] if value else []

            for answer_key in values:
                if not answer_key:
                    continue

                answer_label = options.get(answer_key, answer_key)

                cur.execute("""
                    INSERT INTO intake_workflow_bridge_answers (
                        intake_id, workflow_key, firm_id, question_key,
                        answer_key, answer_label, created_at, updated_at, created_by
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    intake_id,
                    workflow_key,
                    firm_id,
                    qkey,
                    answer_key,
                    answer_label,
                    now,
                    now,
                    created_by,
                ))

        # IMPORTANT:
        # Do not call update_document_recommendation_status() here.
        # That function opens another SQLite connection while this write transaction is open.
        # On Windows SQLite, that can trigger "database is locked."
        cur.execute("""
            UPDATE intake_document_recommendations
            SET status = ?,
                updated_at = ?,
                created_by = ?
            WHERE intake_id = ? AND workflow_key = ?
        """, (
            "launch_prepared",
            now,
            created_by,
            intake_id,
            workflow_key,
        ))

        conn.commit()

    finally:
        conn.close()

def list_workflow_bridge_answers(intake_id, workflow_key):
    ensure_workflow_bridge_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT question_key, answer_key, answer_label, created_at, updated_at, created_by
        FROM intake_workflow_bridge_answers
        WHERE intake_id = ? AND workflow_key = ?
        ORDER BY id ASC
    """, (intake_id, workflow_key))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "question_key": row[0],
            "answer_key": row[1],
            "answer_label": row[2],
            "created_at": format_intake_timestamp(row[3]),
            "updated_at": format_intake_timestamp(row[4]),
            "created_by": row[5] or "—",
        }
        for row in rows
    ]


def build_workflow_bridge_summary(intake_id, workflow_key):
    definition = get_workflow_bridge_definition(workflow_key)
    launch = build_workflow_launch_prep(intake_id, workflow_key)

    if not definition or not launch:
        return None

    answers = list_workflow_bridge_answers(intake_id, workflow_key)

    answers_by_question = {}
    for answer in answers:
        answers_by_question.setdefault(answer["question_key"], []).append(answer)

    question_summaries = []
    for question in definition.get("questions", []):
        qkey = question.get("key")
        question_summaries.append({
            "question_key": qkey,
            "label": question.get("label"),
            "answers": answers_by_question.get(qkey, []),
        })

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "definition": definition,
        "launch": launch,
        "answers": answers,
        "question_summaries": question_summaries,
        "bridge_status": "Prepared" if answers else "Not Completed",
    }


# -------------------------------------------------------------------
# INT-2E — Bridge-to-Draft Packet Generator
# -------------------------------------------------------------------

DRAFT_PACKET_TYPE_MAP = {
    "professional_review_checklist": "Professional Review Draft Packet",
    "business_continuity_packet": "Business Continuity Draft Packet",
    "real_property_review": "Real Property Review Draft Packet",
    "foundational_estate_package": "Foundational Estate Planning Draft Packet",
    "document_audit": "Existing Document Audit Draft Packet",
    "beneficiary_guardian_planning": "Beneficiary / Guardian Planning Draft Packet",
    "asset_inventory_packet": "Asset Inventory Draft Packet",
    "fiduciary_authority_review": "Fiduciary Authority Review Draft Packet",
    "next_session_agenda": "Next Session Agenda Draft Packet",
}


def _bridge_answer_labels(summary):
    labels = []
    for item in summary.get("question_summaries", []) or []:
        for answer in item.get("answers", []) or []:
            labels.append(str(answer.get("answer_label") or ""))
    return labels


def _bridge_answer_text(summary):
    return " ".join(_bridge_answer_labels(summary)).lower()


def build_draft_packet_questions(workflow_key, bridge_summary):
    text = _bridge_answer_text(bridge_summary)
    questions = []

    # Universal drafting questions
    questions.extend([
        "Who is the primary planning party or client for this workflow?",
        "Should this draft be client-facing, internal-only, or both?",
        "Which documents must be collected before a draft can be finalized?",
    ])

    if workflow_key == "business_continuity_packet":
        questions.extend([
            "What is the exact business name and DBA/entity name?",
            "Who currently has authority to operate or bind the business?",
            "Who should serve as backup operator, manager, or successor contact?",
            "Which business records prove authority, ownership, or continuity?",
        ])

        if "sole proprietorship" in text or "dba" in text:
            questions.append("Should the continuity packet treat this as a sole-proprietor/DBA continuity plan?")
        if "liability" in text or "insurance" in text:
            questions.append("What insurance, liability, or risk documents must be reviewed before drafting?")
        if "records missing" in text:
            questions.append("Which missing records must be requested before the packet is finalized?")

    elif workflow_key == "professional_review_checklist":
        questions.extend([
            "What professional review issue must be separated first: legal, tax, court, creditor, or administrative?",
            "Is there a response deadline or notice date?",
            "Which documents should be copied into the professional review packet?",
        ])

    elif workflow_key == "real_property_review":
        questions.extend([
            "What is the property address or parcel identifier?",
            "Who is listed on the deed or title record?",
            "Are there liens, mortgages, co-owners, or transfer restrictions?",
            "Which title/property records must be collected before any transfer recommendation?",
        ])

    elif workflow_key == "foundational_estate_package":
        questions.extend([
            "Which foundational documents are most likely needed first?",
            "Who are the proposed decision-makers and successors?",
            "Are there minor children, guardianship issues, blended-family issues, or special-needs considerations?",
        ])

    elif workflow_key == "document_audit":
        questions.extend([
            "Which documents should be audited first?",
            "Are any documents unsigned, undated, unnotarized, outdated, or inconsistent?",
            "Which documents should be marked as active, superseded, incomplete, or uncertain?",
        ])

    elif workflow_key == "beneficiary_guardian_planning":
        questions.extend([
            "Who are the intended beneficiaries?",
            "Are any beneficiaries minors, dependents, or special-needs beneficiaries?",
            "Should distributions be equal, customized, staged, restricted, or reviewed further?",
        ])

    elif workflow_key == "asset_inventory_packet":
        questions.extend([
            "Which asset category should be inventoried first?",
            "What ownership evidence exists for each asset?",
            "Which assets are ready for review and which require missing documents?",
        ])

    elif workflow_key == "fiduciary_authority_review":
        questions.extend([
            "What document grants fiduciary authority?",
            "Who is the fiduciary, agent, trustee, executor, or administrator?",
            "What limitations, dates, signatures, or conditions affect authority?",
        ])

    elif workflow_key == "next_session_agenda":
        questions.extend([
            "What is the primary goal of the next session?",
            "Who needs to attend?",
            "Which documents or decisions must be ready before that session?",
        ])

    return questions


def build_draft_packet_open_issues(workflow_key, bridge_summary):
    issues = []

    launch = bridge_summary.get("launch", {}) or {}
    packet = launch.get("packet", {}) or {}

    open_tasks = launch.get("open_tasks", []) or []
    review_flags = launch.get("review_flags", []) or []
    documents = launch.get("documents", []) or []

    if open_tasks:
        issues.append(f"{len(open_tasks)} open follow-up task(s) remain before final drafting.")

    if review_flags:
        for flag in review_flags:
            issues.append(f"Review flag: {flag}")

    if documents:
        issues.append("Document checklist must be confirmed before final document generation.")

    text = _bridge_answer_text(bridge_summary)

    if "not sure" in text:
        issues.append("One or more bridge answers indicate uncertainty that must be clarified.")

    if "none yet" in text:
        issues.append("One or more required document categories may not be available yet.")

    if workflow_key == "business_continuity_packet":
        if "liability" in text or "insurance" in text:
            issues.append("Business liability/insurance concern requires review before final drafting.")
        if "sole proprietorship" in text or "dba" in text:
            issues.append("Sole proprietor/DBA continuity authority should be handled carefully.")

    return issues


def build_draft_packet_readiness(bridge_summary):
    answers = bridge_summary.get("answers", []) or []
    open_issues = build_draft_packet_open_issues(
        bridge_summary.get("workflow_key"),
        bridge_summary,
    )

    if not answers:
        return "Bridge Not Completed"

    if len(open_issues) >= 4:
        return "Draft Prep With Major Open Issues"

    if open_issues:
        return "Draft Prep With Open Issues"

    return "Draft Prep Ready"


def build_workflow_draft_packet(intake_id, workflow_key):
    bridge_summary = build_workflow_bridge_summary(intake_id, workflow_key)
    if not bridge_summary:
        return None

    launch = bridge_summary.get("launch", {}) or {}
    recommendation = launch.get("recommendation", {}) or {}

    draft_packet_type = DRAFT_PACKET_TYPE_MAP.get(
        workflow_key,
        "Workflow Draft Packet",
    )

    drafting_questions = build_draft_packet_questions(workflow_key, bridge_summary)
    open_issues = build_draft_packet_open_issues(workflow_key, bridge_summary)
    readiness = build_draft_packet_readiness(bridge_summary)

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "draft_packet_type": draft_packet_type,
        "readiness": readiness,
        "bridge_summary": bridge_summary,
        "launch": launch,
        "recommendation": recommendation,
        "drafting_questions": drafting_questions,
        "open_issues": open_issues,
        "documents": launch.get("documents", []) or [],
        "open_tasks": launch.get("open_tasks", []) or [],
        "completed_tasks": launch.get("completed_tasks", []) or [],
        "question_summaries": bridge_summary.get("question_summaries", []) or [],
        "notice": "This draft packet is a preparation layer only. It does not generate final legal, fiduciary, tax, property, or business documents.",
    }


# -------------------------------------------------------------------
# INT-2F — Draft Packet Export + Draft Readiness Ledger
# -------------------------------------------------------------------

def ensure_draft_readiness_tables():
    ensure_intake_export_log_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_draft_readiness_ledger (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            draft_packet_type TEXT,
            readiness TEXT,
            open_issue_count INTEGER DEFAULT 0,
            open_task_count INTEGER DEFAULT 0,
            completed_task_count INTEGER DEFAULT 0,
            document_count INTEGER DEFAULT 0,
            drafting_question_count INTEGER DEFAULT 0,
            status TEXT DEFAULT 'prepared',
            created_at TEXT,
            updated_at TEXT,
            updated_by TEXT,
            notes TEXT,
            UNIQUE(intake_id, workflow_key)
        )
    """)

    conn.commit()
    conn.close()


def upsert_draft_readiness_record(intake_id, workflow_key, draft_packet, updated_by=None):
    ensure_draft_readiness_tables()

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    readiness = draft_packet.get("readiness")
    open_issue_count = len(draft_packet.get("open_issues", []) or [])
    open_task_count = len(draft_packet.get("open_tasks", []) or [])
    completed_task_count = len(draft_packet.get("completed_tasks", []) or [])
    document_count = len(draft_packet.get("documents", []) or [])
    drafting_question_count = len(draft_packet.get("drafting_questions", []) or [])

    status = "blocked" if "Major Open Issues" in str(readiness) else "prepared"
    if readiness == "Draft Prep Ready":
        status = "ready"

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_draft_readiness_ledger (
            intake_id, workflow_key, firm_id, draft_packet_type, readiness,
            open_issue_count, open_task_count, completed_task_count,
            document_count, drafting_question_count, status,
            created_at, updated_at, updated_by, notes
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(intake_id, workflow_key) DO UPDATE SET
            draft_packet_type = excluded.draft_packet_type,
            readiness = excluded.readiness,
            open_issue_count = excluded.open_issue_count,
            open_task_count = excluded.open_task_count,
            completed_task_count = excluded.completed_task_count,
            document_count = excluded.document_count,
            drafting_question_count = excluded.drafting_question_count,
            status = excluded.status,
            updated_at = excluded.updated_at,
            updated_by = excluded.updated_by,
            notes = excluded.notes
    """, (
        intake_id,
        workflow_key,
        firm_id,
        draft_packet.get("draft_packet_type"),
        readiness,
        open_issue_count,
        open_task_count,
        completed_task_count,
        document_count,
        drafting_question_count,
        status,
        now,
        now,
        updated_by,
        "Auto-updated from INT-2F draft packet readiness.",
    ))

    conn.commit()
    conn.close()


def list_draft_readiness_records(intake_id=None):
    ensure_draft_readiness_tables()

    conn = get_connection()
    cur = conn.cursor()

    if intake_id:
        cur.execute("""
            SELECT intake_id, workflow_key, draft_packet_type, readiness,
                   open_issue_count, open_task_count, completed_task_count,
                   document_count, drafting_question_count, status,
                   created_at, updated_at, updated_by, notes
            FROM intake_draft_readiness_ledger
            WHERE intake_id = ?
            ORDER BY updated_at DESC
        """, (intake_id,))
    else:
        cur.execute("""
            SELECT intake_id, workflow_key, draft_packet_type, readiness,
                   open_issue_count, open_task_count, completed_task_count,
                   document_count, drafting_question_count, status,
                   created_at, updated_at, updated_by, notes
            FROM intake_draft_readiness_ledger
            ORDER BY updated_at DESC
            LIMIT 200
        """)

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": row[0],
            "workflow_key": row[1],
            "draft_packet_type": row[2],
            "readiness": row[3],
            "open_issue_count": row[4],
            "open_task_count": row[5],
            "completed_task_count": row[6],
            "document_count": row[7],
            "drafting_question_count": row[8],
            "status": row[9],
            "created_at": format_intake_timestamp(row[10]) if row[10] else "",
            "updated_at": format_intake_timestamp(row[11]) if row[11] else "",
            "updated_by": row[12] or "—",
            "notes": row[13] or "",
        }
        for row in rows
    ]


def ensure_draft_packet_export_dir():
    from pathlib import Path

    export_dir = Path("exports/draft_packets")
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def generate_workflow_draft_packet_docx(intake_id, workflow_key, created_by=None):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    draft_packet = build_workflow_draft_packet(intake_id, workflow_key)
    if not draft_packet:
        return None

    upsert_draft_readiness_record(
        intake_id=intake_id,
        workflow_key=workflow_key,
        draft_packet=draft_packet,
        updated_by=created_by,
    )

    export_dir = ensure_draft_packet_export_dir()
    filename = f"{safe_export_filename(intake_id)}_{safe_export_filename(workflow_key)}_Draft_Packet.docx"
    out_path = export_dir / filename

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(draft_packet.get("draft_packet_type") or "Workflow Draft Packet")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Intake ID: {intake_id} | Workflow: {workflow_key}")

    doc.add_paragraph("")

    def add_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        return p

    def add_bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.add_run(f"- {text}")

    add_heading("Draft Packet Summary")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Readiness", draft_packet.get("readiness")),
        ("Recommendation", draft_packet.get("recommendation", {}).get("title")),
        ("Priority", draft_packet.get("recommendation", {}).get("priority")),
        ("Confidence", f"{draft_packet.get('recommendation', {}).get('confidence')}%"),
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value or "")

    doc.add_paragraph("")

    add_heading("Workflow Bridge Answers")
    for item in draft_packet.get("question_summaries", []) or []:
        add_bullet(item.get("label") or "")
        answers = item.get("answers", []) or []
        if answers:
            for answer in answers:
                add_bullet(f"    {answer.get('answer_label')}")
        else:
            add_bullet("    No answer recorded.")

    add_heading("Drafting Questions")
    for item in draft_packet.get("drafting_questions", []) or []:
        add_bullet(f"[ ] {item}")

    add_heading("Open Issues Before Drafting")
    for item in draft_packet.get("open_issues", []) or []:
        add_bullet(item)
    if not draft_packet.get("open_issues"):
        add_bullet("No open issues detected.")

    add_heading("Documents Needed")
    for item in draft_packet.get("documents", []) or []:
        add_bullet(f"[ ] {item}")

    add_heading("Open Tasks")
    for task in draft_packet.get("open_tasks", []) or []:
        add_bullet(f"[ ] {task.get('title')}")
        add_bullet(f"    {task.get('task_type_label')} | {task.get('status_label')} | Priority: {task.get('priority_label')}")

    add_heading("Completed Tasks")
    for task in draft_packet.get("completed_tasks", []) or []:
        add_bullet(f"[x] {task.get('title')}")
        add_bullet(f"    {task.get('task_type_label')} | {task.get('status_label')}")

    add_heading("Important Notice")
    doc.add_paragraph(draft_packet.get("notice") or "This is a draft-preparation packet only.")

    doc.save(out_path)

    try:
        log_intake_export_versioned(
            intake_id=intake_id,
            export_type="draft_docx",
            export_status="success",
            file_path=str(out_path),
            message=f"Draft packet DOCX generated for {workflow_key}.",
            created_by=created_by,
            packet_type="draft_packet",
        )
    except Exception:
        pass

    return str(out_path)


# -------------------------------------------------------------------
# INT-2G — Controlled Document Draft Questionnaire
# -------------------------------------------------------------------

DOCUMENT_DRAFT_TYPE_CATALOG = {
    "business_continuity_packet": [
        {
            "document_key": "business_continuity_memo",
            "title": "Business Continuity Memo",
            "description": "A controlled internal/client-facing memo organizing business authority, continuity risks, documents, and next actions.",
        },
        {
            "document_key": "business_authority_checklist",
            "title": "Business Authority Checklist",
            "description": "A checklist to confirm who has authority to operate, sign, access records, and continue operations.",
        },
        {
            "document_key": "business_records_request",
            "title": "Business Records Request List",
            "description": "A structured request list for business records, insurance, contracts, authority documents, and operating documents.",
        },
    ],
    "professional_review_checklist": [
        {
            "document_key": "professional_review_memo",
            "title": "Professional Review Memo",
            "description": "A memo separating tax, legal, creditor, court, and administrative review issues.",
        },
        {
            "document_key": "urgent_issue_checklist",
            "title": "Urgent Issue Checklist",
            "description": "A checklist for deadlines, notices, court dates, tax notices, claim letters, or creditor pressure.",
        },
    ],
    "real_property_review": [
        {
            "document_key": "property_review_memo",
            "title": "Real Property Review Memo",
            "description": "A memo organizing ownership, deed/title, liens, taxes, insurance, and transfer-readiness questions.",
        },
        {
            "document_key": "property_records_request",
            "title": "Property Records Request List",
            "description": "A document request list for deeds, tax bills, mortgage records, insurance, surveys, leases, and title records.",
        },
    ],
    "foundational_estate_package": [
        {
            "document_key": "foundational_planning_memo",
            "title": "Foundational Planning Memo",
            "description": "A memo organizing family structure, decision-makers, beneficiaries, documents, and open planning issues.",
        },
        {
            "document_key": "estate_document_request",
            "title": "Estate Document Request List",
            "description": "A request list for wills, trusts, POAs, health directives, deeds, beneficiary forms, and asset statements.",
        },
    ],
    "document_audit": [
        {
            "document_key": "document_audit_memo",
            "title": "Existing Document Audit Memo",
            "description": "A memo identifying documents to audit, execution concerns, conflicts, missing signatures, and outdated provisions.",
        },
        {
            "document_key": "document_audit_checklist",
            "title": "Document Audit Checklist",
            "description": "A checklist to classify documents as active, outdated, incomplete, conflicting, or requiring review.",
        },
    ],
    "beneficiary_guardian_planning": [
        {
            "document_key": "beneficiary_guardian_memo",
            "title": "Beneficiary / Guardian Planning Memo",
            "description": "A memo organizing beneficiaries, guardians, minors, dependents, and distribution concerns.",
        },
        {
            "document_key": "beneficiary_review_checklist",
            "title": "Beneficiary Review Checklist",
            "description": "A checklist for beneficiary designations, guardian choices, minor controls, and distribution review.",
        },
    ],
    "asset_inventory_packet": [
        {
            "document_key": "asset_inventory_memo",
            "title": "Asset Inventory Memo",
            "description": "A memo organizing asset categories, ownership clarity, missing records, and inventory priorities.",
        },
        {
            "document_key": "asset_records_request",
            "title": "Asset Records Request List",
            "description": "A request list for statements, titles, deeds, insurance records, appraisals, account registrations, and supporting records.",
        },
    ],
    "fiduciary_authority_review": [
        {
            "document_key": "fiduciary_authority_memo",
            "title": "Fiduciary Authority Review Memo",
            "description": "A memo organizing fiduciary role, authority document, limitations, evidence, and restrictions.",
        },
        {
            "document_key": "fiduciary_authority_checklist",
            "title": "Fiduciary Authority Checklist",
            "description": "A checklist for trustee, executor, agent, administrator, POA, or other authority verification.",
        },
    ],
    "next_session_agenda": [
        {
            "document_key": "next_session_agenda",
            "title": "Next Session Agenda",
            "description": "A structured agenda for the next client/internal review session.",
        },
    ],
}


DOCUMENT_DRAFT_QUESTION_BANK = {
    "business_continuity_memo": [
        {"key": "business_name", "label": "What is the exact business name or DBA?", "input_type": "text"},
        {"key": "current_operator", "label": "Who currently operates or controls the business?", "input_type": "text"},
        {"key": "successor_contact", "label": "Who should be contacted if the operator is unavailable?", "input_type": "text"},
        {"key": "continuity_priority", "label": "What is the highest continuity priority?", "input_type": "radio", "options": {
            "authority": "Confirm authority",
            "records": "Organize records",
            "insurance": "Review insurance/liability",
            "succession": "Prepare successor operator",
            "not_sure": "Not sure",
        }},
    ],
    "business_authority_checklist": [
        {"key": "authority_holder", "label": "Who currently has signing or operating authority?", "input_type": "text"},
        {"key": "authority_evidence", "label": "What evidence supports that authority?", "input_type": "checkbox", "options": {
            "ein": "EIN letter",
            "operating_agreement": "Operating agreement",
            "bank_records": "Bank records",
            "license": "Business license",
            "contracts": "Contracts",
            "other": "Other",
        }},
        {"key": "authority_gap", "label": "Is there an authority gap?", "input_type": "radio", "options": {
            "yes": "Yes",
            "no": "No",
            "not_sure": "Not sure",
        }},
    ],
    "business_records_request": [
        {"key": "records_needed", "label": "Which business records must be requested?", "input_type": "checkbox", "options": {
            "ein": "EIN letter",
            "operating_agreement": "Operating agreement",
            "license": "Business license/registration",
            "insurance": "Insurance policy",
            "contracts": "Contracts",
            "bank_records": "Bank authority records",
            "tax_records": "Tax records",
            "other": "Other",
        }},
        {"key": "request_recipient", "label": "Who should receive the records request?", "input_type": "text"},
        {"key": "urgency", "label": "How urgent is the records request?", "input_type": "radio", "options": {
            "low": "Low",
            "normal": "Normal",
            "high": "High",
            "urgent": "Urgent",
        }},
    ],
    "professional_review_memo": [
        {"key": "issue_summary", "label": "Briefly summarize the review issue.", "input_type": "textarea"},
        {"key": "review_category", "label": "What category best fits the issue?", "input_type": "checkbox", "options": {
            "tax": "Tax",
            "legal": "Legal",
            "court": "Court",
            "creditor": "Creditor/claim",
            "business": "Business liability",
            "family": "Family conflict",
            "other": "Other",
        }},
        {"key": "deadline", "label": "Is there a deadline or response date?", "input_type": "text"},
    ],
    "urgent_issue_checklist": [
        {"key": "urgent_document", "label": "Which urgent document triggered this?", "input_type": "text"},
        {"key": "response_needed", "label": "What response appears needed?", "input_type": "textarea"},
        {"key": "outside_review", "label": "Is outside professional review needed before action?", "input_type": "radio", "options": {
            "yes": "Yes",
            "no": "No",
            "not_sure": "Not sure",
        }},
    ],
    "property_review_memo": [
        {"key": "property_identifier", "label": "What is the property address or parcel identifier?", "input_type": "text"},
        {"key": "record_owner", "label": "Who appears on the deed/title record?", "input_type": "text"},
        {"key": "property_issue", "label": "What property issue needs review?", "input_type": "checkbox", "options": {
            "deed": "Deed/title",
            "mortgage": "Mortgage/lien",
            "tax": "Tax bill",
            "insurance": "Insurance",
            "co_owner": "Co-owner",
            "transfer": "Transfer/funding",
            "other": "Other",
        }},
    ],
    "property_records_request": [
        {"key": "property_records_needed", "label": "Which property records should be requested?", "input_type": "checkbox", "options": {
            "deed": "Deed",
            "tax_bill": "Tax bill",
            "mortgage": "Mortgage statement",
            "insurance": "Insurance declarations",
            "survey": "Survey",
            "lease": "Lease",
            "title": "Title records",
            "other": "Other",
        }},
        {"key": "property_request_recipient", "label": "Who should provide these records?", "input_type": "text"},
    ],
    "foundational_planning_memo": [
        {"key": "planning_party", "label": "Who is the planning party?", "input_type": "text"},
        {"key": "primary_goal", "label": "What is the primary planning goal?", "input_type": "textarea"},
        {"key": "document_priority", "label": "Which foundational document should be considered first?", "input_type": "radio", "options": {
            "will": "Will",
            "trust": "Trust",
            "poa": "Power of attorney",
            "health": "Health directive",
            "beneficiary": "Beneficiary update",
            "not_sure": "Not sure",
        }},
    ],
    "estate_document_request": [
        {"key": "estate_records_needed", "label": "Which estate planning records are needed?", "input_type": "checkbox", "options": {
            "will": "Will",
            "trust": "Trust",
            "poa": "Power of attorney",
            "health": "Health directive",
            "deed": "Deed",
            "beneficiary": "Beneficiary form",
            "asset_statement": "Asset statement",
            "other": "Other",
        }},
        {"key": "request_deadline", "label": "When should these be gathered?", "input_type": "text"},
    ],
    "document_audit_memo": [
        {"key": "audit_scope", "label": "What documents are in the audit scope?", "input_type": "textarea"},
        {"key": "audit_risk", "label": "What audit concerns are present?", "input_type": "checkbox", "options": {
            "outdated": "Outdated",
            "conflict": "Conflicting",
            "unsigned": "Unsigned/incomplete",
            "wrong_party": "Wrong party/name",
            "missing_asset": "Missing asset reference",
            "not_sure": "Not sure",
        }},
    ],
    "document_audit_checklist": [
        {"key": "classification_needed", "label": "How should documents be classified?", "input_type": "checkbox", "options": {
            "active": "Active",
            "superseded": "Superseded",
            "incomplete": "Incomplete",
            "conflicting": "Conflicting",
            "review_required": "Review required",
        }},
        {"key": "audit_owner", "label": "Who will gather or review the documents?", "input_type": "text"},
    ],
    "beneficiary_guardian_memo": [
        {"key": "beneficiary_notes", "label": "Summarize beneficiary/guardian planning concerns.", "input_type": "textarea"},
        {"key": "minor_or_dependent", "label": "Are minors or dependents involved?", "input_type": "radio", "options": {
            "yes": "Yes",
            "no": "No",
            "not_sure": "Not sure",
        }},
    ],
    "beneficiary_review_checklist": [
        {"key": "beneficiary_records", "label": "Which beneficiary records should be reviewed?", "input_type": "checkbox", "options": {
            "life_insurance": "Life insurance beneficiary",
            "retirement": "Retirement account beneficiary",
            "bank": "Bank/account beneficiary",
            "trust": "Trust beneficiary provision",
            "will": "Will distribution provision",
            "other": "Other",
        }},
        {"key": "guardian_candidate", "label": "Who is the proposed guardian/caregiver, if applicable?", "input_type": "text"},
    ],
    "asset_inventory_memo": [
        {"key": "inventory_scope", "label": "What assets should be inventoried first?", "input_type": "textarea"},
        {"key": "ownership_clarity", "label": "Is ownership clear?", "input_type": "radio", "options": {
            "yes": "Yes",
            "some": "Some assets",
            "no": "No",
            "not_sure": "Not sure",
        }},
    ],
    "asset_records_request": [
        {"key": "asset_records_needed", "label": "Which asset records are needed?", "input_type": "checkbox", "options": {
            "statements": "Statements",
            "titles": "Titles",
            "deeds": "Deeds",
            "insurance": "Insurance records",
            "registrations": "Account registrations",
            "appraisals": "Appraisals/valuations",
            "other": "Other",
        }},
        {"key": "asset_request_priority", "label": "Which asset record is most urgent?", "input_type": "text"},
    ],
    "fiduciary_authority_memo": [
        {"key": "fiduciary_name", "label": "Who is the fiduciary/agent/trustee?", "input_type": "text"},
        {"key": "authority_document", "label": "What document grants authority?", "input_type": "text"},
        {"key": "authority_limits", "label": "Are there known limits or conditions?", "input_type": "textarea"},
    ],
    "fiduciary_authority_checklist": [
        {"key": "authority_evidence_needed", "label": "Which authority evidence is needed?", "input_type": "checkbox", "options": {
            "trust": "Trust instrument",
            "poa": "Power of attorney",
            "letters": "Letters testamentary/administration",
            "resolution": "Resolution/minute",
            "id": "Identification",
            "other": "Other",
        }},
        {"key": "authority_status", "label": "Is authority currently usable?", "input_type": "radio", "options": {
            "yes": "Yes",
            "no": "No",
            "not_sure": "Not sure",
        }},
    ],
    "next_session_agenda": [
        {"key": "agenda_goal", "label": "What is the next session goal?", "input_type": "textarea"},
        {"key": "agenda_items", "label": "Which agenda items should be included?", "input_type": "checkbox", "options": {
            "documents": "Document review",
            "assets": "Asset inventory",
            "family": "Family/beneficiary planning",
            "business": "Business continuity",
            "risk": "Risk/professional review",
            "next_steps": "Next steps",
        }},
        {"key": "session_owner", "label": "Who is responsible for preparing the next session?", "input_type": "text"},
    ],
}


def get_document_draft_types_for_workflow(workflow_key):
    return DOCUMENT_DRAFT_TYPE_CATALOG.get(workflow_key, [])


def get_document_draft_type(workflow_key, document_key):
    for item in get_document_draft_types_for_workflow(workflow_key):
        if item.get("document_key") == document_key:
            return item
    return None


def get_document_draft_questions(document_key):
    return DOCUMENT_DRAFT_QUESTION_BANK.get(document_key, [])


def ensure_document_draft_questionnaire_tables():
    ensure_draft_readiness_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_document_draft_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            document_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            question_key TEXT NOT NULL,
            answer_key TEXT,
            answer_label TEXT,
            created_at TEXT,
            updated_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def save_document_draft_answers(intake_id, workflow_key, document_key, form_data, created_by=None):
    ensure_document_draft_questionnaire_tables()

    questions = get_document_draft_questions(document_key)
    if not questions:
        raise ValueError("Unknown document draft questionnaire.")

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    try:
        cur.execute("""
            DELETE FROM intake_document_draft_answers
            WHERE intake_id = ? AND workflow_key = ? AND document_key = ?
        """, (intake_id, workflow_key, document_key))

        for question in questions:
            qkey = question.get("key")
            input_type = question.get("input_type")
            options = question.get("options", {})

            if input_type == "checkbox":
                values = form_data.getlist(qkey) if hasattr(form_data, "getlist") else []
            else:
                value = form_data.get(qkey) if hasattr(form_data, "get") else None
                values = [value] if value else []

            for answer_key in values:
                if not answer_key:
                    continue

                answer_label = options.get(answer_key, answer_key)

                cur.execute("""
                    INSERT INTO intake_document_draft_answers (
                        intake_id, workflow_key, document_key, firm_id,
                        question_key, answer_key, answer_label,
                        created_at, updated_at, created_by
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    intake_id,
                    workflow_key,
                    document_key,
                    firm_id,
                    qkey,
                    answer_key,
                    answer_label,
                    now,
                    now,
                    created_by,
                ))

        conn.commit()

    finally:
        conn.close()


def list_document_draft_answers(intake_id, workflow_key, document_key):
    ensure_document_draft_questionnaire_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT question_key, answer_key, answer_label, created_at, updated_at, created_by
        FROM intake_document_draft_answers
        WHERE intake_id = ? AND workflow_key = ? AND document_key = ?
        ORDER BY id ASC
    """, (intake_id, workflow_key, document_key))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "question_key": row[0],
            "answer_key": row[1],
            "answer_label": row[2],
            "created_at": format_intake_timestamp(row[3]),
            "updated_at": format_intake_timestamp(row[4]),
            "created_by": row[5] or "—",
        }
        for row in rows
    ]


def build_document_draft_preview(intake_id, workflow_key, document_key):
    draft_packet = build_workflow_draft_packet(intake_id, workflow_key)
    document_type = get_document_draft_type(workflow_key, document_key)
    questions = get_document_draft_questions(document_key)

    if not draft_packet or not document_type or not questions:
        return None

    answers = list_document_draft_answers(intake_id, workflow_key, document_key)
    answers_by_question = {}

    for answer in answers:
        answers_by_question.setdefault(answer["question_key"], []).append(answer)

    question_summaries = []
    for question in questions:
        qkey = question.get("key")
        question_summaries.append({
            "question_key": qkey,
            "label": question.get("label"),
            "input_type": question.get("input_type"),
            "answers": answers_by_question.get(qkey, []),
        })

    answered_count = len([q for q in question_summaries if q["answers"]])
    total_count = len(question_summaries)

    if answered_count == 0:
        preview_status = "Questionnaire Not Completed"
    elif answered_count < total_count:
        preview_status = "Partial Draft Preview"
    else:
        preview_status = "Draft Preview Ready"

    outline_sections = [
        "Purpose / Scope",
        "Information Gathered",
        "Open Issues",
        "Required Documents",
        "Recommended Next Steps",
        "Review Notice",
    ]

    if workflow_key == "business_continuity_packet":
        outline_sections.insert(2, "Business Authority / Continuity Notes")
    elif workflow_key == "real_property_review":
        outline_sections.insert(2, "Property Ownership / Records Notes")
    elif workflow_key == "professional_review_checklist":
        outline_sections.insert(2, "Professional Review Issue Notes")
    elif workflow_key == "asset_inventory_packet":
        outline_sections.insert(2, "Asset Inventory Notes")

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "document_key": document_key,
        "document_type": document_type,
        "draft_packet": draft_packet,
        "questions": questions,
        "answers": answers,
        "question_summaries": question_summaries,
        "answered_count": answered_count,
        "total_count": total_count,
        "preview_status": preview_status,
        "outline_sections": outline_sections,
        "notice": "This is a non-final draft preview generated from controlled questionnaire answers. It is not a final legal, fiduciary, tax, property, or business document.",
    }


# -------------------------------------------------------------------
# INT-2H — Non-Final Draft Document Generator
# -------------------------------------------------------------------

def _answers_for_question(preview, question_key):
    for item in preview.get("question_summaries", []) or []:
        if item.get("question_key") == question_key:
            return [a.get("answer_label") for a in item.get("answers", []) or [] if a.get("answer_label")]
    return []


def _first_answer(preview, question_key, default="Not provided"):
    answers = _answers_for_question(preview, question_key)
    return answers[0] if answers else default


def build_nonfinal_draft_sections(preview):
    workflow_key = preview.get("workflow_key")
    document_key = preview.get("document_key")
    draft_packet = preview.get("draft_packet", {}) or {}

    sections = []

    sections.append({
        "heading": "Purpose / Scope",
        "body": (
            f"This non-final draft organizes information for {preview.get('document_type', {}).get('title', 'the selected document')} "
            f"under workflow {workflow_key}. It is prepared from intake data, bridge answers, and controlled draft-questionnaire responses."
        ),
    })

    gathered_lines = []
    for item in preview.get("question_summaries", []) or []:
        label = item.get("label")
        answers = [a.get("answer_label") for a in item.get("answers", []) or [] if a.get("answer_label")]
        if answers:
            gathered_lines.append(f"{label}: " + "; ".join(answers))
        else:
            gathered_lines.append(f"{label}: Not provided")

    sections.append({
        "heading": "Information Gathered",
        "body": "\n".join(gathered_lines) if gathered_lines else "No controlled questionnaire answers were recorded.",
    })

    if workflow_key == "business_continuity_packet":
        business_name = _first_answer(preview, "business_name")
        current_operator = _first_answer(preview, "current_operator")
        successor_contact = _first_answer(preview, "successor_contact")
        continuity_priority = _first_answer(preview, "continuity_priority")

        sections.append({
            "heading": "Business Authority / Continuity Notes",
            "body": (
                f"Business / DBA: {business_name}\n"
                f"Current Operator / Control Party: {current_operator}\n"
                f"Successor or Backup Contact: {successor_contact}\n"
                f"Continuity Priority: {continuity_priority}\n\n"
                "This section is a preliminary continuity summary only. Authority, ownership, insurance, and business records "
                "must be verified before this draft is used for final action."
            ),
        })

    elif workflow_key == "professional_review_checklist":
        sections.append({
            "heading": "Professional Review Issue Notes",
            "body": "This draft should separate legal, tax, court, creditor, business, or family-risk issues before any final action.",
        })

    elif workflow_key == "real_property_review":
        sections.append({
            "heading": "Property Ownership / Records Notes",
            "body": "This draft should be supported by deed/title records, property tax records, mortgage/lien information, and insurance documents.",
        })

    elif workflow_key == "asset_inventory_packet":
        sections.append({
            "heading": "Asset Inventory Notes",
            "body": "This draft should be supported by ownership records, asset statements, titles, deeds, insurance records, and registration evidence.",
        })

    open_issues = draft_packet.get("open_issues", []) or []
    sections.append({
        "heading": "Open Issues",
        "body": "\n".join(open_issues) if open_issues else "No open issues were detected from the draft packet.",
    })

    documents = draft_packet.get("documents", []) or []
    sections.append({
        "heading": "Required Documents",
        "body": "\n".join([f"[ ] {doc}" for doc in documents]) if documents else "No document checklist was available.",
    })

    next_steps = [
        "Review missing questionnaire answers.",
        "Confirm all required documents are collected.",
        "Resolve open issues before final drafting.",
        "Confirm whether outside professional review is needed.",
        "Move to final-draft preparation only after review gates are cleared.",
    ]

    if preview.get("preview_status") == "Draft Preview Ready":
        next_steps.insert(0, "Controlled questionnaire is complete; proceed to review-gate evaluation.")
    else:
        next_steps.insert(0, "Controlled questionnaire is incomplete; complete missing answers before final drafting.")

    sections.append({
        "heading": "Recommended Next Steps",
        "body": "\n".join([f"- {step}" for step in next_steps]),
    })

    sections.append({
        "heading": "Review Notice",
        "body": (
            "NON-FINAL DRAFT PREVIEW ONLY. This document is generated for intake, organization, and preparation. "
            "It is not legal advice, tax advice, fiduciary advice, property-transfer advice, or a final enforceable document. "
            "Review and approval gates must be completed before any formal use."
        ),
    })

    return sections


def build_nonfinal_draft_document(intake_id, workflow_key, document_key):
    preview = build_document_draft_preview(intake_id, workflow_key, document_key)
    if not preview:
        return None

    sections = build_nonfinal_draft_sections(preview)

    missing_answers = []
    for item in preview.get("question_summaries", []) or []:
        if not item.get("answers"):
            missing_answers.append(item.get("label"))

    if missing_answers:
        document_status = "Non-Final Draft — Incomplete Questionnaire"
    elif preview.get("draft_packet", {}).get("open_issues"):
        document_status = "Non-Final Draft — Review Required"
    else:
        document_status = "Non-Final Draft — Ready for Review Gate"

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "document_key": document_key,
        "title": preview.get("document_type", {}).get("title", "Non-Final Draft Document"),
        "description": preview.get("document_type", {}).get("description", ""),
        "preview": preview,
        "sections": sections,
        "missing_answers": missing_answers,
        "document_status": document_status,
        "readiness": preview.get("draft_packet", {}).get("readiness"),
        "notice": "This is a non-final draft document generated from controlled intake data. It is not a final document.",
    }


# -------------------------------------------------------------------
# INT-2I — Non-Final Draft DOCX Export + Review Gate Ledger
# -------------------------------------------------------------------

def ensure_review_gate_tables():
    ensure_document_draft_questionnaire_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_review_gate_ledger (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            document_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            gate_name TEXT,
            gate_status TEXT DEFAULT 'pending',
            gate_reason TEXT,
            missing_answer_count INTEGER DEFAULT 0,
            open_issue_count INTEGER DEFAULT 0,
            open_task_count INTEGER DEFAULT 0,
            document_status TEXT,
            created_at TEXT,
            updated_at TEXT,
            updated_by TEXT,
            notes TEXT,
            UNIQUE(intake_id, workflow_key, document_key, gate_name)
        )
    """)

    conn.commit()
    conn.close()


def evaluate_nonfinal_review_gate(document):
    missing_count = len(document.get("missing_answers", []) or [])
    preview = document.get("preview", {}) or {}
    draft_packet = preview.get("draft_packet", {}) or {}
    open_issue_count = len(draft_packet.get("open_issues", []) or [])
    open_task_count = len(draft_packet.get("open_tasks", []) or [])

    if missing_count > 0:
        return {
            "gate_status": "blocked",
            "gate_reason": "Controlled questionnaire has missing answers.",
            "missing_answer_count": missing_count,
            "open_issue_count": open_issue_count,
            "open_task_count": open_task_count,
        }

    if open_issue_count > 0 or open_task_count > 0:
        return {
            "gate_status": "review_required",
            "gate_reason": "Draft has open issues or open tasks requiring review.",
            "missing_answer_count": missing_count,
            "open_issue_count": open_issue_count,
            "open_task_count": open_task_count,
        }

    return {
        "gate_status": "ready_for_review",
        "gate_reason": "Controlled questionnaire is complete and no open issues/tasks were detected.",
        "missing_answer_count": missing_count,
        "open_issue_count": open_issue_count,
        "open_task_count": open_task_count,
    }


def upsert_review_gate_record(intake_id, workflow_key, document_key, document, updated_by=None):
    ensure_review_gate_tables()

    gate = evaluate_nonfinal_review_gate(document)
    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_review_gate_ledger (
            intake_id, workflow_key, document_key, firm_id,
            gate_name, gate_status, gate_reason,
            missing_answer_count, open_issue_count, open_task_count,
            document_status, created_at, updated_at, updated_by, notes
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(intake_id, workflow_key, document_key, gate_name) DO UPDATE SET
            gate_status = excluded.gate_status,
            gate_reason = excluded.gate_reason,
            missing_answer_count = excluded.missing_answer_count,
            open_issue_count = excluded.open_issue_count,
            open_task_count = excluded.open_task_count,
            document_status = excluded.document_status,
            updated_at = excluded.updated_at,
            updated_by = excluded.updated_by,
            notes = excluded.notes
    """, (
        intake_id,
        workflow_key,
        document_key,
        firm_id,
        "non_final_draft_review_gate",
        gate["gate_status"],
        gate["gate_reason"],
        gate["missing_answer_count"],
        gate["open_issue_count"],
        gate["open_task_count"],
        document.get("document_status"),
        now,
        now,
        updated_by,
        "Auto-updated from INT-2I non-final draft review gate.",
    ))

    conn.commit()
    conn.close()

    return gate


def list_review_gate_records(intake_id=None):
    ensure_review_gate_tables()

    conn = get_connection()
    cur = conn.cursor()

    if intake_id:
        cur.execute("""
            SELECT intake_id, workflow_key, document_key, gate_name,
                   gate_status, gate_reason, missing_answer_count,
                   open_issue_count, open_task_count, document_status,
                   created_at, updated_at, updated_by, notes
            FROM intake_review_gate_ledger
            WHERE intake_id = ?
            ORDER BY updated_at DESC
        """, (intake_id,))
    else:
        cur.execute("""
            SELECT intake_id, workflow_key, document_key, gate_name,
                   gate_status, gate_reason, missing_answer_count,
                   open_issue_count, open_task_count, document_status,
                   created_at, updated_at, updated_by, notes
            FROM intake_review_gate_ledger
            ORDER BY updated_at DESC
            LIMIT 200
        """)

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": row[0],
            "workflow_key": row[1],
            "document_key": row[2],
            "gate_name": row[3],
            "gate_status": row[4],
            "gate_reason": row[5],
            "missing_answer_count": row[6],
            "open_issue_count": row[7],
            "open_task_count": row[8],
            "document_status": row[9],
            "created_at": format_intake_timestamp(row[10]) if row[10] else "",
            "updated_at": format_intake_timestamp(row[11]) if row[11] else "",
            "updated_by": row[12] or "—",
            "notes": row[13] or "",
        }
        for row in rows
    ]


def ensure_nonfinal_draft_export_dir():
    from pathlib import Path

    export_dir = Path("exports/nonfinal_drafts")
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def generate_nonfinal_draft_docx(intake_id, workflow_key, document_key, created_by=None):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = build_nonfinal_draft_document(intake_id, workflow_key, document_key)
    if not document:
        return None

    gate = upsert_review_gate_record(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        document=document,
        updated_by=created_by,
    )

    export_dir = ensure_nonfinal_draft_export_dir()
    filename = (
        f"{safe_export_filename(intake_id)}_"
        f"{safe_export_filename(workflow_key)}_"
        f"{safe_export_filename(document_key)}_NON_FINAL_DRAFT.docx"
    )
    out_path = export_dir / filename

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(document.get("title") or "Non-Final Draft Document")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("NON-FINAL DRAFT — REVIEW REQUIRED")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Intake ID: {intake_id} | Workflow: {workflow_key} | Document: {document_key}"
    )

    doc.add_paragraph("")

    def add_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        return p

    def add_body(text):
        for line in str(text or "").splitlines():
            doc.add_paragraph(line if line else "")

    add_heading("Draft Status")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    status_rows = [
        ("Document Status", document.get("document_status")),
        ("Draft Readiness", document.get("readiness")),
        ("Preview Status", document.get("preview", {}).get("preview_status")),
        ("Answered", f"{document.get('preview', {}).get('answered_count')} / {document.get('preview', {}).get('total_count')}"),
        ("Review Gate", gate.get("gate_status")),
        ("Gate Reason", gate.get("gate_reason")),
    ]

    for label, value in status_rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value or "")

    doc.add_paragraph("")

    if document.get("missing_answers"):
        add_heading("Missing Questionnaire Answers")
        for item in document.get("missing_answers", []):
            doc.add_paragraph(f"- {item}")

    for section in document.get("sections", []) or []:
        add_heading(section.get("heading"))
        add_body(section.get("body"))

    add_heading("Non-Final Draft Notice")
    add_body(
        "This document is a non-final draft generated from controlled intake data. "
        "It is not legal advice, tax advice, fiduciary advice, property-transfer advice, "
        "or a final enforceable document. Review gates must be completed before formal use."
    )

    doc.save(out_path)

    try:
        log_intake_export_versioned(
            intake_id=intake_id,
            export_type="nonfinal_docx",
            export_status="success",
            file_path=str(out_path),
            message=f"Non-final draft DOCX generated for {workflow_key}/{document_key}. Gate={gate.get('gate_status')}.",
            created_by=created_by,
            packet_type="nonfinal_draft",
        )
    except Exception:
        pass

    return str(out_path)


# -------------------------------------------------------------------
# INT-2J — Review Gate Resolution Controls
# -------------------------------------------------------------------

VALID_REVIEW_GATE_ACTIONS = {
    "missing_answers_resolved": "Missing Answers Resolved",
    "open_issues_reviewed": "Open Issues Reviewed",
    "professional_review_required": "Professional Review Required",
    "approved_for_final_draft_prep": "Approved for Final-Draft Preparation",
    "blocked_hold": "Blocked / Hold",
}

VALID_REVIEW_GATE_STATUSES = {
    "blocked": "Blocked",
    "review_required": "Review Required",
    "professional_review_required": "Professional Review Required",
    "approved_for_final_draft_prep": "Approved for Final-Draft Preparation",
    "hold": "Hold",
    "ready_for_review": "Ready for Review",
}


def ensure_review_gate_resolution_tables():
    ensure_review_gate_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_review_gate_actions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            document_key TEXT NOT NULL,
            gate_name TEXT DEFAULT 'non_final_draft_review_gate',
            firm_id TEXT DEFAULT 'FIRM-001',
            action_key TEXT,
            action_label TEXT,
            resulting_status TEXT,
            note TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def get_review_gate_record(intake_id, workflow_key, document_key, gate_name="non_final_draft_review_gate"):
    ensure_review_gate_resolution_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, workflow_key, document_key, gate_name,
               gate_status, gate_reason, missing_answer_count,
               open_issue_count, open_task_count, document_status,
               created_at, updated_at, updated_by, notes
        FROM intake_review_gate_ledger
        WHERE intake_id = ?
          AND workflow_key = ?
          AND document_key = ?
          AND gate_name = ?
        LIMIT 1
    """, (intake_id, workflow_key, document_key, gate_name))

    row = cur.fetchone()
    conn.close()

    if not row:
        return None

    return {
        "intake_id": row[0],
        "workflow_key": row[1],
        "document_key": row[2],
        "gate_name": row[3],
        "gate_status": row[4],
        "gate_status_label": VALID_REVIEW_GATE_STATUSES.get(row[4], row[4]),
        "gate_reason": row[5],
        "missing_answer_count": row[6],
        "open_issue_count": row[7],
        "open_task_count": row[8],
        "document_status": row[9],
        "created_at": format_intake_timestamp(row[10]) if row[10] else "",
        "updated_at": format_intake_timestamp(row[11]) if row[11] else "",
        "updated_by": row[12] or "—",
        "notes": row[13] or "",
    }


def list_review_gate_actions(intake_id, workflow_key, document_key, gate_name="non_final_draft_review_gate"):
    ensure_review_gate_resolution_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT action_key, action_label, resulting_status, note, created_at, created_by
        FROM intake_review_gate_actions
        WHERE intake_id = ?
          AND workflow_key = ?
          AND document_key = ?
          AND gate_name = ?
        ORDER BY id DESC
    """, (intake_id, workflow_key, document_key, gate_name))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "action_key": row[0],
            "action_label": row[1],
            "resulting_status": row[2],
            "resulting_status_label": VALID_REVIEW_GATE_STATUSES.get(row[2], row[2]),
            "note": row[3] or "",
            "created_at": format_intake_timestamp(row[4]) if row[4] else "",
            "created_by": row[5] or "—",
        }
        for row in rows
    ]


def resolve_review_gate_action(
    intake_id,
    workflow_key,
    document_key,
    action_key,
    note="",
    created_by=None,
    gate_name="non_final_draft_review_gate"
):
    ensure_review_gate_resolution_tables()

    if action_key not in VALID_REVIEW_GATE_ACTIONS:
        raise ValueError("Invalid review gate action.")

    action_label = VALID_REVIEW_GATE_ACTIONS[action_key]

    if action_key == "missing_answers_resolved":
        resulting_status = "review_required"
        reason = "Missing answers marked resolved; remaining issues require review."
    elif action_key == "open_issues_reviewed":
        resulting_status = "ready_for_review"
        reason = "Open issues marked reviewed; gate is ready for review."
    elif action_key == "professional_review_required":
        resulting_status = "professional_review_required"
        reason = "Professional review required before final-draft preparation."
    elif action_key == "approved_for_final_draft_prep":
        resulting_status = "approved_for_final_draft_prep"
        reason = "Approved for final-draft preparation. Final execution is still not authorized."
    elif action_key == "blocked_hold":
        resulting_status = "hold"
        reason = "Gate placed on hold / blocked."
    else:
        resulting_status = "review_required"
        reason = "Gate action recorded."

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    try:
        cur.execute("""
            INSERT INTO intake_review_gate_actions (
                intake_id, workflow_key, document_key, gate_name, firm_id,
                action_key, action_label, resulting_status, note,
                created_at, created_by
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            intake_id,
            workflow_key,
            document_key,
            gate_name,
            firm_id,
            action_key,
            action_label,
            resulting_status,
            note,
            now,
            created_by,
        ))

        cur.execute("""
            UPDATE intake_review_gate_ledger
            SET gate_status = ?,
                gate_reason = ?,
                updated_at = ?,
                updated_by = ?,
                notes = ?
            WHERE intake_id = ?
              AND workflow_key = ?
              AND document_key = ?
              AND gate_name = ?
        """, (
            resulting_status,
            reason,
            now,
            created_by,
            note,
            intake_id,
            workflow_key,
            document_key,
            gate_name,
        ))

        conn.commit()

    finally:
        conn.close()

    return {
        "action_key": action_key,
        "action_label": action_label,
        "resulting_status": resulting_status,
        "resulting_status_label": VALID_REVIEW_GATE_STATUSES.get(resulting_status, resulting_status),
        "reason": reason,
        "note": note,
        "created_at": now,
        "created_by": created_by,
    }


# -------------------------------------------------------------------
# INT-2K — Final-Draft Preparation Gate
# -------------------------------------------------------------------

def ensure_final_draft_prep_gate_tables():
    ensure_review_gate_resolution_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_final_draft_prep_gate (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            document_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            gate_status TEXT DEFAULT 'blocked',
            gate_reason TEXT,
            questionnaire_complete INTEGER DEFAULT 0,
            open_issues_reviewed INTEGER DEFAULT 0,
            open_tasks_reviewed INTEGER DEFAULT 0,
            professional_review_recorded INTEGER DEFAULT 0,
            required_documents_acknowledged INTEGER DEFAULT 0,
            admin_approved INTEGER DEFAULT 0,
            approval_note TEXT,
            created_at TEXT,
            updated_at TEXT,
            updated_by TEXT,
            UNIQUE(intake_id, workflow_key, document_key)
        )
    """)

    conn.commit()
    conn.close()


def evaluate_final_draft_prep_gate(intake_id, workflow_key, document_key):
    document = build_nonfinal_draft_document(intake_id, workflow_key, document_key)
    review_gate = get_review_gate_record(intake_id, workflow_key, document_key)
    actions = list_review_gate_actions(intake_id, workflow_key, document_key)

    if not document or not review_gate:
        return None

    preview = document.get("preview", {}) or {}
    draft_packet = preview.get("draft_packet", {}) or {}

    missing_answers = document.get("missing_answers", []) or []
    open_issues = draft_packet.get("open_issues", []) or []
    open_tasks = draft_packet.get("open_tasks", []) or []
    documents = draft_packet.get("documents", []) or []

    action_keys = {a.get("action_key") for a in actions}
    resulting_statuses = {a.get("resulting_status") for a in actions}

    questionnaire_complete = 1 if len(missing_answers) == 0 else 0

    open_issues_reviewed = 1 if (
        len(open_issues) == 0
        or "open_issues_reviewed" in action_keys
        or "approved_for_final_draft_prep" in action_keys
    ) else 0

    open_tasks_reviewed = 1 if (
        len(open_tasks) == 0
        or "open_issues_reviewed" in action_keys
        or "approved_for_final_draft_prep" in action_keys
    ) else 0

    professional_review_recorded = 1 if (
        "professional_review_required" in action_keys
        or "approved_for_final_draft_prep" in action_keys
        or review_gate.get("gate_status") in {
            "professional_review_required",
            "approved_for_final_draft_prep",
            "ready_for_review",
        }
    ) else 0

    # Required documents are acknowledged if the document list is empty, or if
    # review-gate action history shows an issue review / approval action.
    required_documents_acknowledged = 1 if (
        len(documents) == 0
        or "open_issues_reviewed" in action_keys
        or "approved_for_final_draft_prep" in action_keys
    ) else 0

    hard_blocks = []

    if not questionnaire_complete:
        hard_blocks.append("Controlled questionnaire has missing answers.")

    if not open_issues_reviewed:
        hard_blocks.append("Open issues have not been reviewed or accepted.")

    if not open_tasks_reviewed:
        hard_blocks.append("Open tasks have not been reviewed or accepted.")

    if not professional_review_recorded:
        hard_blocks.append("Professional review status has not been recorded.")

    if not required_documents_acknowledged:
        hard_blocks.append("Required document checklist has not been acknowledged.")

    if hard_blocks:
        gate_status = "blocked"
        gate_reason = " ".join(hard_blocks)
    else:
        gate_status = "ready_for_admin_approval"
        gate_reason = "All pre-approval review conditions are satisfied. Admin approval is still required."

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "document_key": document_key,
        "document": document,
        "review_gate": review_gate,
        "actions": actions,
        "questionnaire_complete": questionnaire_complete,
        "open_issues_reviewed": open_issues_reviewed,
        "open_tasks_reviewed": open_tasks_reviewed,
        "professional_review_recorded": professional_review_recorded,
        "required_documents_acknowledged": required_documents_acknowledged,
        "admin_approved": 0,
        "gate_status": gate_status,
        "gate_reason": gate_reason,
        "missing_answers": missing_answers,
        "open_issues": open_issues,
        "open_tasks": open_tasks,
        "required_documents": documents,
    }


def upsert_final_draft_prep_gate(intake_id, workflow_key, document_key, updated_by=None):
    ensure_final_draft_prep_gate_tables()

    evaluation = evaluate_final_draft_prep_gate(intake_id, workflow_key, document_key)
    if not evaluation:
        return None

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    # Preserve admin approval if already approved, unless the evaluation is blocked.
    cur.execute("""
        SELECT admin_approved, approval_note
        FROM intake_final_draft_prep_gate
        WHERE intake_id = ? AND workflow_key = ? AND document_key = ?
        LIMIT 1
    """, (intake_id, workflow_key, document_key))
    existing = cur.fetchone()

    existing_admin_approved = int(existing[0]) if existing else 0
    existing_note = existing[1] if existing else ""

    admin_approved = existing_admin_approved if evaluation["gate_status"] != "blocked" else 0

    if evaluation["gate_status"] == "blocked":
        final_status = "blocked"
    elif admin_approved:
        final_status = "approved_for_final_draft_preparation"
    else:
        final_status = evaluation["gate_status"]

    final_reason = (
        "Approved for final-draft preparation. This still does not authorize signing, filing, execution, transfer, or legal finalization."
        if final_status == "approved_for_final_draft_preparation"
        else evaluation["gate_reason"]
    )

    cur.execute("""
        INSERT INTO intake_final_draft_prep_gate (
            intake_id, workflow_key, document_key, firm_id,
            gate_status, gate_reason,
            questionnaire_complete, open_issues_reviewed, open_tasks_reviewed,
            professional_review_recorded, required_documents_acknowledged,
            admin_approved, approval_note, created_at, updated_at, updated_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(intake_id, workflow_key, document_key) DO UPDATE SET
            gate_status = excluded.gate_status,
            gate_reason = excluded.gate_reason,
            questionnaire_complete = excluded.questionnaire_complete,
            open_issues_reviewed = excluded.open_issues_reviewed,
            open_tasks_reviewed = excluded.open_tasks_reviewed,
            professional_review_recorded = excluded.professional_review_recorded,
            required_documents_acknowledged = excluded.required_documents_acknowledged,
            admin_approved = excluded.admin_approved,
            approval_note = excluded.approval_note,
            updated_at = excluded.updated_at,
            updated_by = excluded.updated_by
    """, (
        intake_id,
        workflow_key,
        document_key,
        firm_id,
        final_status,
        final_reason,
        evaluation["questionnaire_complete"],
        evaluation["open_issues_reviewed"],
        evaluation["open_tasks_reviewed"],
        evaluation["professional_review_recorded"],
        evaluation["required_documents_acknowledged"],
        admin_approved,
        existing_note,
        now,
        now,
        updated_by,
    ))

    conn.commit()
    conn.close()

    return get_final_draft_prep_gate(intake_id, workflow_key, document_key)


def get_final_draft_prep_gate(intake_id, workflow_key, document_key):
    ensure_final_draft_prep_gate_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT intake_id, workflow_key, document_key,
               gate_status, gate_reason,
               questionnaire_complete, open_issues_reviewed, open_tasks_reviewed,
               professional_review_recorded, required_documents_acknowledged,
               admin_approved, approval_note,
               created_at, updated_at, updated_by
        FROM intake_final_draft_prep_gate
        WHERE intake_id = ? AND workflow_key = ? AND document_key = ?
        LIMIT 1
    """, (intake_id, workflow_key, document_key))

    row = cur.fetchone()
    conn.close()

    if not row:
        return None

    return {
        "intake_id": row[0],
        "workflow_key": row[1],
        "document_key": row[2],
        "gate_status": row[3],
        "gate_reason": row[4],
        "questionnaire_complete": int(row[5] or 0),
        "open_issues_reviewed": int(row[6] or 0),
        "open_tasks_reviewed": int(row[7] or 0),
        "professional_review_recorded": int(row[8] or 0),
        "required_documents_acknowledged": int(row[9] or 0),
        "admin_approved": int(row[10] or 0),
        "approval_note": row[11] or "",
        "created_at": format_intake_timestamp(row[12]) if row[12] else "",
        "updated_at": format_intake_timestamp(row[13]) if row[13] else "",
        "updated_by": row[14] or "—",
    }


def approve_final_draft_prep_gate(intake_id, workflow_key, document_key, approval_note="", approved_by=None):
    ensure_final_draft_prep_gate_tables()

    gate = upsert_final_draft_prep_gate(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        updated_by=approved_by,
    )

    if not gate:
        raise ValueError("Final-draft preparation gate could not be evaluated.")

    if gate.get("gate_status") == "blocked":
        raise ValueError("Gate is blocked. Resolve required conditions before approval.")

    now = datetime.utcnow().isoformat(timespec="seconds")

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        UPDATE intake_final_draft_prep_gate
        SET gate_status = ?,
            gate_reason = ?,
            admin_approved = 1,
            approval_note = ?,
            updated_at = ?,
            updated_by = ?
        WHERE intake_id = ? AND workflow_key = ? AND document_key = ?
    """, (
        "approved_for_final_draft_preparation",
        "Approved for final-draft preparation. This still does not authorize signing, filing, execution, transfer, or legal finalization.",
        approval_note,
        now,
        approved_by,
        intake_id,
        workflow_key,
        document_key,
    ))

    conn.commit()
    conn.close()

    return get_final_draft_prep_gate(intake_id, workflow_key, document_key)


def list_final_draft_prep_gates(intake_id=None):
    ensure_final_draft_prep_gate_tables()

    conn = get_connection()
    cur = conn.cursor()

    if intake_id:
        cur.execute("""
            SELECT intake_id, workflow_key, document_key,
                   gate_status, gate_reason,
                   questionnaire_complete, open_issues_reviewed, open_tasks_reviewed,
                   professional_review_recorded, required_documents_acknowledged,
                   admin_approved, approval_note, updated_at, updated_by
            FROM intake_final_draft_prep_gate
            WHERE intake_id = ?
            ORDER BY updated_at DESC
        """, (intake_id,))
    else:
        cur.execute("""
            SELECT intake_id, workflow_key, document_key,
                   gate_status, gate_reason,
                   questionnaire_complete, open_issues_reviewed, open_tasks_reviewed,
                   professional_review_recorded, required_documents_acknowledged,
                   admin_approved, approval_note, updated_at, updated_by
            FROM intake_final_draft_prep_gate
            ORDER BY updated_at DESC
            LIMIT 200
        """)

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": row[0],
            "workflow_key": row[1],
            "document_key": row[2],
            "gate_status": row[3],
            "gate_reason": row[4],
            "questionnaire_complete": int(row[5] or 0),
            "open_issues_reviewed": int(row[6] or 0),
            "open_tasks_reviewed": int(row[7] or 0),
            "professional_review_recorded": int(row[8] or 0),
            "required_documents_acknowledged": int(row[9] or 0),
            "admin_approved": int(row[10] or 0),
            "approval_note": row[11] or "",
            "updated_at": format_intake_timestamp(row[12]) if row[12] else "",
            "updated_by": row[13] or "—",
        }
        for row in rows
    ]


# -------------------------------------------------------------------
# INT-2L — Final-Draft Gate Resolution Workflow
# -------------------------------------------------------------------

FINAL_DRAFT_RESOLUTION_ACTIONS = {
    "missing_answers_acknowledged": "Missing Answers Reviewed / Acknowledged",
    "open_issues_reviewed": "Open Issues Reviewed / Accepted",
    "open_tasks_reviewed": "Open Tasks Reviewed / Accepted",
    "required_documents_acknowledged": "Required Documents Acknowledged",
    "professional_review_confirmed": "Professional Review Status Confirmed",
}


def ensure_final_draft_resolution_tables():
    ensure_final_draft_prep_gate_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_final_draft_gate_actions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            document_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            action_key TEXT,
            action_label TEXT,
            note TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def list_final_draft_resolution_actions(intake_id, workflow_key, document_key):
    ensure_final_draft_resolution_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT action_key, action_label, note, created_at, created_by
        FROM intake_final_draft_gate_actions
        WHERE intake_id = ?
          AND workflow_key = ?
          AND document_key = ?
        ORDER BY id DESC
    """, (intake_id, workflow_key, document_key))

    rows = cur.fetchall()
    conn.close()

    return [
        {
            "action_key": row[0],
            "action_label": row[1],
            "note": row[2] or "",
            "created_at": format_intake_timestamp(row[3]) if row[3] else "",
            "created_by": row[4] or "—",
        }
        for row in rows
    ]


def _final_draft_resolution_action_keys(intake_id, workflow_key, document_key):
    actions = list_final_draft_resolution_actions(intake_id, workflow_key, document_key)
    return {a.get("action_key") for a in actions}


def record_final_draft_resolution_actions(
    intake_id,
    workflow_key,
    document_key,
    action_keys,
    note="",
    created_by=None
):
    ensure_final_draft_resolution_tables()

    if not action_keys:
        return []

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()
    recorded = []

    conn = get_connection()
    cur = conn.cursor()

    try:
        for action_key in action_keys:
            if action_key not in FINAL_DRAFT_RESOLUTION_ACTIONS:
                continue

            action_label = FINAL_DRAFT_RESOLUTION_ACTIONS[action_key]

            cur.execute("""
                INSERT INTO intake_final_draft_gate_actions (
                    intake_id, workflow_key, document_key, firm_id,
                    action_key, action_label, note, created_at, created_by
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                intake_id,
                workflow_key,
                document_key,
                firm_id,
                action_key,
                action_label,
                note,
                now,
                created_by,
            ))

            recorded.append({
                "action_key": action_key,
                "action_label": action_label,
                "note": note,
                "created_at": now,
                "created_by": created_by,
            })

        conn.commit()

    finally:
        conn.close()

    # Re-evaluate gate after recording actions.
    upsert_final_draft_prep_gate(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        updated_by=created_by,
    )

    return recorded


def evaluate_final_draft_prep_gate_with_resolutions(intake_id, workflow_key, document_key):
    """
    Extends INT-2K evaluation by allowing deliberate INT-2L resolution actions
    to satisfy review/acknowledgment conditions without hiding the underlying facts.
    Missing questionnaire answers can be acknowledged for preparation, but the
    underlying missing count remains visible in the non-final draft/review gate.
    """
    base = evaluate_final_draft_prep_gate(intake_id, workflow_key, document_key)
    if not base:
        return None

    resolution_keys = _final_draft_resolution_action_keys(intake_id, workflow_key, document_key)

    # Apply deliberate resolution acknowledgments.
    if "missing_answers_acknowledged" in resolution_keys:
        base["questionnaire_complete"] = 1

    if "open_issues_reviewed" in resolution_keys:
        base["open_issues_reviewed"] = 1

    if "open_tasks_reviewed" in resolution_keys:
        base["open_tasks_reviewed"] = 1

    if "required_documents_acknowledged" in resolution_keys:
        base["required_documents_acknowledged"] = 1

    if "professional_review_confirmed" in resolution_keys:
        base["professional_review_recorded"] = 1

    hard_blocks = []

    if not base["questionnaire_complete"]:
        hard_blocks.append("Controlled questionnaire has missing answers or has not been acknowledged.")

    if not base["open_issues_reviewed"]:
        hard_blocks.append("Open issues have not been reviewed or accepted.")

    if not base["open_tasks_reviewed"]:
        hard_blocks.append("Open tasks have not been reviewed or accepted.")

    if not base["professional_review_recorded"]:
        hard_blocks.append("Professional review status has not been recorded or confirmed.")

    if not base["required_documents_acknowledged"]:
        hard_blocks.append("Required document checklist has not been acknowledged.")

    if hard_blocks:
        base["gate_status"] = "blocked"
        base["gate_reason"] = " ".join(hard_blocks)
    else:
        base["gate_status"] = "ready_for_admin_approval"
        base["gate_reason"] = "All final-draft preparation conditions are reviewed or acknowledged. Admin approval is still required."

    return base


def upsert_final_draft_prep_gate_with_resolutions(intake_id, workflow_key, document_key, updated_by=None):
    ensure_final_draft_prep_gate_tables()

    evaluation = evaluate_final_draft_prep_gate_with_resolutions(intake_id, workflow_key, document_key)
    if not evaluation:
        return None

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT admin_approved, approval_note
        FROM intake_final_draft_prep_gate
        WHERE intake_id = ? AND workflow_key = ? AND document_key = ?
        LIMIT 1
    """, (intake_id, workflow_key, document_key))
    existing = cur.fetchone()

    existing_admin_approved = int(existing[0]) if existing else 0
    existing_note = existing[1] if existing else ""

    admin_approved = existing_admin_approved if evaluation["gate_status"] != "blocked" else 0

    if evaluation["gate_status"] == "blocked":
        final_status = "blocked"
    elif admin_approved:
        final_status = "approved_for_final_draft_preparation"
    else:
        final_status = evaluation["gate_status"]

    final_reason = (
        "Approved for final-draft preparation. This still does not authorize signing, filing, execution, transfer, or legal finalization."
        if final_status == "approved_for_final_draft_preparation"
        else evaluation["gate_reason"]
    )

    cur.execute("""
        INSERT INTO intake_final_draft_prep_gate (
            intake_id, workflow_key, document_key, firm_id,
            gate_status, gate_reason,
            questionnaire_complete, open_issues_reviewed, open_tasks_reviewed,
            professional_review_recorded, required_documents_acknowledged,
            admin_approved, approval_note, created_at, updated_at, updated_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(intake_id, workflow_key, document_key) DO UPDATE SET
            gate_status = excluded.gate_status,
            gate_reason = excluded.gate_reason,
            questionnaire_complete = excluded.questionnaire_complete,
            open_issues_reviewed = excluded.open_issues_reviewed,
            open_tasks_reviewed = excluded.open_tasks_reviewed,
            professional_review_recorded = excluded.professional_review_recorded,
            required_documents_acknowledged = excluded.required_documents_acknowledged,
            admin_approved = excluded.admin_approved,
            approval_note = excluded.approval_note,
            updated_at = excluded.updated_at,
            updated_by = excluded.updated_by
    """, (
        intake_id,
        workflow_key,
        document_key,
        firm_id,
        final_status,
        final_reason,
        evaluation["questionnaire_complete"],
        evaluation["open_issues_reviewed"],
        evaluation["open_tasks_reviewed"],
        evaluation["professional_review_recorded"],
        evaluation["required_documents_acknowledged"],
        admin_approved,
        existing_note,
        now,
        now,
        updated_by,
    ))

    conn.commit()
    conn.close()

    return get_final_draft_prep_gate(intake_id, workflow_key, document_key)


def build_final_draft_resolution_context(intake_id, workflow_key, document_key):
    gate = upsert_final_draft_prep_gate_with_resolutions(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        updated_by=None,
    )

    if not gate:
        return None

    document = build_nonfinal_draft_document(intake_id, workflow_key, document_key)
    review_gate = get_review_gate_record(intake_id, workflow_key, document_key)
    actions = list_final_draft_resolution_actions(intake_id, workflow_key, document_key)

    return {
        "gate": gate,
        "document": document,
        "review_gate": review_gate,
        "actions": actions,
        "action_options": FINAL_DRAFT_RESOLUTION_ACTIONS,
        "missing_answers": document.get("missing_answers", []) if document else [],
        "open_issues": document.get("preview", {}).get("draft_packet", {}).get("open_issues", []) if document else [],
        "open_tasks": document.get("preview", {}).get("draft_packet", {}).get("open_tasks", []) if document else [],
        "required_documents": document.get("preview", {}).get("draft_packet", {}).get("documents", []) if document else [],
    }


# -------------------------------------------------------------------
# INT-2M — Admin Approval for Final-Draft Preparation
# -------------------------------------------------------------------

def ensure_final_draft_admin_approval_tables():
    ensure_final_draft_resolution_tables()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS intake_final_draft_admin_approvals (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            intake_id TEXT NOT NULL,
            workflow_key TEXT NOT NULL,
            document_key TEXT NOT NULL,
            firm_id TEXT DEFAULT 'FIRM-001',
            approval_status TEXT DEFAULT 'approved_for_final_draft_preparation',
            approval_note TEXT NOT NULL,
            gate_status_before TEXT,
            gate_status_after TEXT,
            created_at TEXT,
            created_by TEXT
        )
    """)

    conn.commit()
    conn.close()


def record_final_draft_admin_approval(
    intake_id,
    workflow_key,
    document_key,
    approval_note,
    approved_by=None
):
    ensure_final_draft_admin_approval_tables()

    approval_note = (approval_note or "").strip()
    if not approval_note:
        raise ValueError("Approval note is required.")

    gate_before = upsert_final_draft_prep_gate_with_resolutions(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        updated_by=approved_by,
    )

    if not gate_before:
        raise ValueError("Final-draft preparation gate could not be evaluated.")

    if gate_before.get("gate_status") == "blocked":
        raise ValueError("Gate is blocked. Resolve required conditions before admin approval.")

    approved_gate = approve_final_draft_prep_gate(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        approval_note=approval_note,
        approved_by=approved_by,
    )

    now = datetime.utcnow().isoformat(timespec="seconds")
    firm_id = get_current_firm_id()

    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO intake_final_draft_admin_approvals (
            intake_id, workflow_key, document_key, firm_id,
            approval_status, approval_note,
            gate_status_before, gate_status_after,
            created_at, created_by
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        intake_id,
        workflow_key,
        document_key,
        firm_id,
        "approved_for_final_draft_preparation",
        approval_note,
        gate_before.get("gate_status"),
        approved_gate.get("gate_status") if approved_gate else None,
        now,
        approved_by,
    ))

    conn.commit()
    conn.close()

    return {
        "intake_id": intake_id,
        "workflow_key": workflow_key,
        "document_key": document_key,
        "approval_status": "approved_for_final_draft_preparation",
        "approval_note": approval_note,
        "gate_status_before": gate_before.get("gate_status"),
        "gate_status_after": approved_gate.get("gate_status") if approved_gate else None,
        "created_at": now,
        "created_by": approved_by,
        "gate": approved_gate,
    }


def list_final_draft_admin_approvals(intake_id=None, workflow_key=None, document_key=None):
    ensure_final_draft_admin_approval_tables()

    conn = get_connection()
    cur = conn.cursor()

    query = """
        SELECT intake_id, workflow_key, document_key,
               approval_status, approval_note,
               gate_status_before, gate_status_after,
               created_at, created_by
        FROM intake_final_draft_admin_approvals
    """

    filters = []
    params = []

    if intake_id:
        filters.append("intake_id = ?")
        params.append(intake_id)
    if workflow_key:
        filters.append("workflow_key = ?")
        params.append(workflow_key)
    if document_key:
        filters.append("document_key = ?")
        params.append(document_key)

    if filters:
        query += " WHERE " + " AND ".join(filters)

    query += " ORDER BY id DESC LIMIT 200"

    cur.execute(query, tuple(params))
    rows = cur.fetchall()
    conn.close()

    return [
        {
            "intake_id": row[0],
            "workflow_key": row[1],
            "document_key": row[2],
            "approval_status": row[3],
            "approval_note": row[4],
            "gate_status_before": row[5],
            "gate_status_after": row[6],
            "created_at": format_intake_timestamp(row[7]) if row[7] else "",
            "created_by": row[8] or "—",
        }
        for row in rows
    ]


def build_final_draft_admin_approval_context(intake_id, workflow_key, document_key):
    gate = upsert_final_draft_prep_gate_with_resolutions(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
        updated_by=None,
    )

    if not gate:
        return None

    approvals = list_final_draft_admin_approvals(
        intake_id=intake_id,
        workflow_key=workflow_key,
        document_key=document_key,
    )

    # INT-2M persistence repair:
    # If an approval ledger record exists, the display/context must reflect approval,
    # even after the gate is re-evaluated.
    if approvals:
        latest = approvals[0]
        gate["gate_status"] = "approved_for_final_draft_preparation"
        gate["gate_reason"] = (
            "Approved for final-draft preparation. This still does not authorize "
            "signing, filing, execution, transfer, or legal finalization."
        )
        gate["admin_approved"] = 1
        gate["approval_note"] = latest.get("approval_note", "")
        gate["updated_at"] = latest.get("created_at", gate.get("updated_at", ""))
        gate["updated_by"] = latest.get("created_by", gate.get("updated_by", "—"))

    document = build_nonfinal_draft_document(intake_id, workflow_key, document_key)

    can_approve = (
        gate.get("gate_status") in {
            "ready_for_admin_approval",
            "approved_for_final_draft_preparation",
        }
        and not gate.get("admin_approved")
        and not approvals
    )

    return {
        "gate": gate,
        "document": document,
        "approvals": approvals,
        "can_approve": can_approve,
        "already_approved": bool(gate.get("admin_approved") or approvals),
        "notice": "Admin approval only authorizes final-draft preparation. It does not authorize signing, filing, execution, transfer, or final legal use.",
    }
